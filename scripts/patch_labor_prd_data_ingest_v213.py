# -*- coding: utf-8 -*-
"""人员实名制 PRD：补充双路径数据接入（一期 ROMA + 标准上报接口）及流程图。"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

DOC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)
BACKUP = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_备份_V2.1.2_20260815.docx"
)
FIG_DIR = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\docs\prd-shots\labor-data-ingest"
)


def set_para_text(p, value: str):
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str):
    if not cell.paragraphs:
        cell.text = value
        return
    set_para_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def format_table(table):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_heading_space_after(doc):
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def ensure_backup():
    if not BACKUP.exists():
        shutil.copy2(DOC, BACKUP)
    print("backup:", BACKUP, "exists=", BACKUP.exists())


def make_flow_charts():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    FIG_DIR.mkdir(parents=True, exist_ok=True)

    def box(ax, x, y, w, h, text, fc="#F5F7FA", ec="#409EFF"):
        p = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            linewidth=1.2,
            edgecolor=ec,
            facecolor=fc,
        )
        ax.add_patch(p)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9, wrap=True)

    def arrow(ax, x1, y1, x2, y2):
        ax.add_patch(
            FancyArrowPatch(
                (x1, y1),
                (x2, y2),
                arrowstyle="-|>",
                mutation_scale=12,
                linewidth=1.1,
                color="#606266",
            )
        )

    # 图1：总体双路径
    fig, ax = plt.subplots(figsize=(10.5, 5.2), dpi=160)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_title("人员实名制数据接入双路径（总体）", fontsize=12, pad=8)
    box(ax, 4.2, 5.0, 3.6, 0.7, "智慧工程建设管控一体化平台\n人员实名制数据", fc="#ECF5FF")
    box(ax, 0.4, 2.8, 4.2, 1.4, "路径A（现有，逻辑不变）\n已对接一期实名制的项目\n一期系统 → ROMA → 本平台", fc="#F0F9EB", ec="#67C23A")
    box(ax, 7.4, 2.8, 4.2, 1.4, "路径B（补充）\n一期未对接的项目\n闸机/实名制子系统\n→ 标准上报接口 → 本平台", fc="#FDF6EC", ec="#E6A23C")
    box(ax, 3.5, 0.5, 5.0, 1.2, "落库人员主档（统一口径）\n人员基础信息 / 单位信息\n特种作业信息 / 安全教育信息", fc="#FEF0F0", ec="#F56C6C")
    arrow(ax, 2.5, 4.2, 5.2, 5.0)
    arrow(ax, 9.5, 4.2, 6.8, 5.0)
    arrow(ax, 2.5, 2.8, 5.0, 1.7)
    arrow(ax, 9.5, 2.8, 7.0, 1.7)
    p1 = FIG_DIR / "fig-data-ingest-overview.png"
    fig.tight_layout()
    fig.savefig(p1, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # 图2：路径B上报流程
    fig, ax = plt.subplots(figsize=(10.5, 3.8), dpi=160)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    ax.set_title("路径B：标准接口上报流程（一期未对接项目）", fontsize=12, pad=8)
    steps = [
        (0.3, "项目现场闸机\n或实名制子系统"),
        (3.0, "调用本平台\n标准上报接口"),
        (5.7, "校验鉴权\n与字段完整性"),
        (8.4, "写入人员主档\n及关联信息"),
    ]
    for i, (x, text) in enumerate(steps):
        box(ax, x, 1.4, 2.4, 1.3, text, fc="#FDF6EC", ec="#E6A23C")
        if i < len(steps) - 1:
            arrow(ax, x + 2.4, 2.05, steps[i + 1][0], 2.05)
    box(ax, 3.5, 0.2, 5.0, 0.8, "上报字段：人员基础信息、单位信息、特种作业信息、安全教育信息", fc="#F5F7FA", ec="#909399")
    p2 = FIG_DIR / "fig-data-ingest-api.png"
    fig.tight_layout()
    fig.savefig(p2, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print("charts:", p1, p2)
    return p1, p2


def update_meta_revision(doc):
    meta = doc.tables[0]
    for row in meta.rows:
        k = row.cells[0].text.strip()
        if k == "版本号":
            set_cell_text(row.cells[1], "V2.1.3")
        if k == "更新日期":
            set_cell_text(row.cells[1], "2026-08-15")

    rev = doc.tables[1]
    for row in rev.rows[1:]:
        if row.cells[0].text.strip() == "V2.1.3":
            set_cell_text(
                row.cells[3],
                "补充人员实名制数据双路径接入：一期已对接项目仍经 ROMA（逻辑不变）；"
                "一期未对接项目提供标准上报接口（闸机/实名制子系统上报），并补充流程图。",
            )
            set_cell_text(row.cells[2], "2026-08-15")
            return
    tbl = rev._tbl
    tbl.append(deepcopy(tbl.tr_lst[-1]))
    row = rev.rows[-1]
    vals = [
        "V2.1.3",
        "—",
        "2026-08-15",
        "补充人员实名制数据双路径接入：一期已对接项目仍经 ROMA（逻辑不变）；"
        "一期未对接项目由本平台提供标准接口文档，现场闸机或实名制子系统按接口上报"
        "（人员基础信息、单位信息、特种作业信息、安全教育信息）；补充总体/路径B流程图。",
    ]
    for i, v in enumerate(vals):
        set_cell_text(row.cells[i], v)
    format_table(rev)


def patch_52(doc, overview_img: Path):
    """在 5.2 核心流程后补充双路径说明与总流程图（不改写原 1～6 条）。"""
    # 避免重复插入
    if any("人员实名制数据接入双路径" in p.text for p in doc.paragraphs):
        print("5.2 already patched")
        return
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("核心流程："):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("找不到 5.2 核心流程段落")

    cur = insert_paragraph_after(
        anchor,
        "数据接入补充（双路径，互斥按项目接入方式选择，同一项目只走一种主路径）：",
        style="Normal",
    )
    cur = insert_paragraph_after(
        cur,
        "路径A（现有逻辑不变）：项目已在一期系统对接实名制数据时，本平台继续经 ROMA 直接获取人员主档/特种证/考勤等数据。",
        style="List Bullet",
    )
    cur = insert_paragraph_after(
        cur,
        "路径B（补充）：一期尚未对接的项目，由本平台提供标准接口文档；项目现场闸机或实名制子系统按接口上报人员数据（含人员基础信息、单位信息、特种作业信息、安全教育信息）。本平台不直连闸机私有协议，仅提供标准上报接口。",
        style="List Bullet",
    )
    img_p = insert_paragraph_after(cur, "")
    run = img_p.add_run()
    run.add_picture(str(overview_img), width=Inches(6.2))
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = insert_paragraph_after(img_p, "图 5.2-1 人员实名制数据接入双路径（总体）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)


def patch_64(doc):
    """仅补充 6.4 与数据来源相关表述，避免大改内部字段。"""
    mapping = {
        "实名制基本身份/单位工种/特种证，通过ROMA节后（暂定）对接一期实名制数据；": (
            "实名制基本身份/单位工种/特种证：已对接一期的项目，经 ROMA 获取一期实名制数据（现有逻辑不变）；"
            "一期未对接的项目，经本平台标准上报接口由现场闸机或实名制子系统上报（含人员基础信息、单位信息、特种作业信息、安全教育信息）。"
        ),
        "读：ROMA 主档（基本身份/单位工种/特种证等）；本地三级安全教育不被 ROMA 回写覆盖；": (
            "读：按项目接入方式读取人员主档——路径A 经 ROMA（基本身份/单位工种/特种证等）；"
            "路径B 经标准上报接口落库（人员基础信息、单位信息、特种作业信息、安全教育信息）。"
            "本地三级安全教育补录不被 ROMA 回写覆盖；接口上报的安全教育信息按主档关联展示。"
        ),
        "二期不支持新建人员主档；人员仅经一期 → ROMA 同步进入；界面无新建入口；本地三级安全教育记录不被 ROMA 回写覆盖": (
            "二期界面不支持手工新建人员主档（无新建入口）；人员进入方式为："
            "已对接一期的项目经一期→ROMA 同步；一期未对接项目经标准上报接口写入。"
            "本地三级安全教育补录不被 ROMA 回写覆盖。"
        ),
    }
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in mapping:
            set_para_text(p, mapping[t])
            n += 1
    print("6.4 replacements:", n)


def patch_71(doc, api_img: Path):
    """7.1 补充标准上报接口说明、表行与流程图。"""
    if any("标准实名制数据上报接口" in p.text for p in doc.paragraphs):
        print("7.1 already patched")
        return

    h71 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("7.1 外部系统接口"):
            h71 = p
            break
    if h71 is None:
        raise RuntimeError("找不到 7.1")

    # 插在标题后、下节标题前：找标题后第一个段落作为锚点后追加
    cur = insert_paragraph_after(
        h71,
        "人员实名制外部数据接入采用双路径：路径A 为一期经 ROMA 入站（现有逻辑不变）；"
        "路径B 为本平台对外发布标准接口文档，供一期未对接项目的现场闸机或实名制子系统上报。"
        "路径B 上报字段至少包含：人员基础信息、单位信息、特种作业信息、安全教育信息。"
        "接口鉴权、报文样例与错误码以单独《人员实名制标准上报接口文档》为准（本文仅约定业务范围）。",
        style="Normal",
    )
    img_p = insert_paragraph_after(cur, "")
    run = img_p.add_run()
    run.add_picture(str(api_img), width=Inches(6.2))
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = insert_paragraph_after(img_p, "图 7.1-1 路径B：标准接口上报流程（一期未对接项目）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)

    # 更新 table 89（外部接口表）
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows:
            continue
        if rows[0][:4] == ["接口名称", "对接系统", "方向", "主要数据"]:
            # 已有则跳过
            if any("标准实名制数据上报" in r[0] for r in rows[1:]):
                break
            table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
            row = table.rows[-1]
            vals = [
                "标准实名制数据上报接口",
                "项目现场闸机 / 实名制子系统（一期未对接项目）",
                "入站",
                "人员基础信息、单位信息、特种作业信息、安全教育信息",
            ]
            for i, v in enumerate(vals):
                set_cell_text(row.cells[i], v)
            format_table(table)
            print("table89 row added")
            break

    # 上游依赖表补充一行
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows:
            continue
        if rows[0][:4] == ["依赖模块/能力", "依赖内容", "本模块用法", "若不可用时的降级/阻断"]:
            if any("标准上报接口" in r[0] for r in rows[1:]):
                break
            table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
            row = table.rows[-1]
            vals = [
                "标准实名制上报接口（路径B）",
                "闸机/实名制子系统按标准接口上报的人员主档及关联信息",
                "一期未对接项目的人员数据入站；台账/统计/预警同源消费",
                "该项目无上报数据则台账为空；不影响已走 ROMA 的其他项目",
            ]
            for i, v in enumerate(vals):
                set_cell_text(row.cells[i], v)
            format_table(table)
            print("upstream table row added")
            break


def main():
    ensure_backup()
    overview, api = make_flow_charts()
    doc = Document(str(DOC))
    update_meta_revision(doc)
    patch_52(doc, overview)
    patch_64(doc)
    patch_71(doc, api)
    set_heading_space_after(doc)
    doc.save(str(DOC))
    print("saved:", DOC)


if __name__ == "__main__":
    main()
