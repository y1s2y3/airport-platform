# -*- coding: utf-8 -*-
"""生成《安全管理（人员实名制、车辆管理）需求确认书》"""
from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT.parent / "调研记录" / "4、安全管理"
OUT_FILE = OUT_DIR / "安全管理需求确认.docx"
SHOT_DIR = ROOT / "docs" / "safety-req-shots"

DOC_NO = "SZAIR-SAF-REQ-20260704"
VERSION = "V1.0"
TODAY = date(2026, 7, 4).strftime("%Y年%m月%d日")


def set_heading_space_after(paragraph, pt=6):
    paragraph.paragraph_format.space_after = Pt(pt)


def format_table_rows(table, min_height_cm=0.6):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(min_height_cm)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    set_heading_space_after(p)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "宋体"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_meta_table(doc):
    rows = [
        ("文档编号", DOC_NO),
        ("版本", VERSION),
        ("编制日期", TODAY),
        ("适用范围", "安全管理 · 人员实名制管理、车辆管理"),
        ("确认状态", "待业务方确认"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True)
        set_cell_text(table.rows[i].cells[1], v)
    format_table_rows(table)
    doc.add_paragraph()


def add_function_table(doc, rows):
    headers = ["用户端", "层级", "功能模块", "子功能", "功能描述", "需求源"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val)
    format_table_rows(table)
    widths = [Cm(2.2), Cm(2.0), Cm(2.4), Cm(2.4), Cm(7.8), Cm(1.8)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w


def add_screenshot_section(doc, title, desc, image_path=None):
    add_heading(doc, title, level=3)
    add_para(doc, desc)
    if image_path and Path(image_path).exists():
        doc.add_paragraph()
        doc.add_picture(str(image_path), width=Inches(6.2))
    doc.add_paragraph()


def build_function_rows():
    client = "一体化平台"
    src_research = "调研结果"
    src_bid = "招投标"
    src_proto = "原型确认"

    rows = [
        # 通用
        (client, "通用模块", "项目范围", "指挥部/项目切换", "顶部项目选择器或左侧项目树切换数据范围；指挥部可查看各项目汇总，项目级仅查看当前项目数据。", src_research),
        (client, "通用模块", "项目范围", "项目树待办提示", "预警清单等项目树节点展示待处置数量，便于指挥部快速定位问题项目。", src_proto),
        # 人员实名制
        (client, "人员实名制管理", "劳务看板", "汇总指标", "展示在场人数、今日进场/出场、登记人数、待处置预警等核心指标。", src_research),
        (client, "人员实名制管理", "劳务看板", "年龄结构分析", "环形图展示各年龄段人员占比。", src_research),
        (client, "人员实名制管理", "劳务看板", "工种结构分析", "环形图展示各工种人员占比。", src_research),
        (client, "人员实名制管理", "劳务看板", "出勤趋势", "折线图展示近阶段出勤人数与出勤率变化趋势。", src_research),
        (client, "人员实名制管理", "劳务看板", "各项目统计", "指挥部视角展示各项目在场、登记、预警等对比表格。", src_research),
        (client, "人员实名制管理", "劳务看板", "最新预警动态", "展示近期触发的实名制预警摘要，支持跳转预警清单。", src_research),
        (client, "人员实名制管理", "人员实名制", "人员列表", "按项目维护劳务人员台账，展示人员编号、姓名、工种、单位/班组、入场状态、三级教育完成情况等。", src_bid),
        (client, "人员实名制管理", "人员实名制", "检索筛选", "支持按姓名/编号/身份证/手机号/单位关键字及工种、入场状态筛选。", src_proto),
        (client, "人员实名制管理", "人员实名制", "手机号脱敏", "列表默认脱敏展示手机号，点击查看需记录操作日志。", src_research),
        (client, "人员实名制管理", "人员实名制", "新增/编辑人员", "未对接现场实名制系统的项目支持手工维护人员档案；已对接项目禁止手工新增。", src_research),
        (client, "人员实名制管理", "人员实名制", "人员详情", "查看基本信息、单位信息、资质证书、安全教育、进退场记录、考勤摘要等。", src_bid),
        (client, "人员实名制管理", "人员实名制", "现场对接标识", "列表展示项目是否已对接现场实名制系统，区分数据来源。", src_proto),
        (client, "人员实名制管理", "考勤明细", "进出记录列表", "同步展示门禁刷卡产生的进场/出场明细，含闸机、方向、时间等。", src_bid),
        (client, "人员实名制管理", "考勤明细", "多维检索", "支持按人员、闸机、方向、日期范围检索考勤明细。", src_research),
        (client, "人员实名制管理", "考勤统计", "统计汇总", "按项目、单位、工种等维度统计出勤人数、出勤率。", src_research),
        (client, "人员实名制管理", "考勤统计", "图表分析", "柱状图/表格展示各单位、各工种出勤对比。", src_research),
        (client, "人员实名制管理", "劳务黑名单", "黑名单台账", "维护指挥部统一劳务黑名单，限制违规人员进入各项目。", src_research),
        (client, "人员实名制管理", "劳务黑名单", "拦截记录", "记录黑名单人员尝试刷卡/登记进场的时间、闸机、拦截原因。", src_research),
        (client, "人员实名制管理", "劳务黑名单", "新增/解除", "支持新增黑名单人员、解除黑名单并记录操作说明。", src_research),
        (client, "人员实名制管理", "预警清单", "预警列表", "汇总各类实名制规则触发的预警，含预警编号、人员、规则类型、状态、触发时间。", src_bid),
        (client, "人员实名制管理", "预警清单", "状态筛选", "支持按待处置/处置中/已关闭及处置方式（项目自处置/上报指挥部）筛选。", src_research),
        (client, "人员实名制管理", "预警清单", "预警详情", "查看预警详情、关联人员信息、处置记录与分级上报链路。", src_research),
        (client, "人员实名制管理", "预警清单", "预警处置", "项目方可填写处置说明并关闭；需上报的预警按配置推送上级责任人。", src_research),
        (client, "人员实名制管理", "实名制配置", "预警规则配置", "指挥部配置各类预警规则开关及阈值（如连续作业时长、年龄下限、未出勤天数等）。", src_research),
        (client, "人员实名制管理", "实名制配置", "分级管控配置", "配置总监理→项目经理→安全部主管→指挥长四级上报天数及接收人。", src_research),
        (client, "人员实名制管理", "实名制配置", "现场对接配置", "按项目配置是否对接现场实名制系统；对接后人员数据以现场系统为准。", src_proto),
        (client, "人员实名制管理", "实名制配置", "对接字段说明", "展示对接数据字段范围（基本信息、资质、考勤等），便于项目侧联调。", src_proto),
        (client, "人员实名制管理", "数据对接", "劳务系统标准接口", "提供统一接口，支持不同项目、不同厂家劳务/闸机系统推送人员基础信息、资质及考勤数据。", src_bid),
        (client, "人员实名制管理", "数据对接", "准入核验", "门禁刷卡时自动比对黑名单、证书有效性、三级教育等规则；不通过则触发预警（PER-02/03）。", src_research),
        # 车辆管理
        (client, "车辆管理", "车辆管理看板", "汇总指标", "展示今日进场/出场、在场车辆、在途车辆、登记车辆及异常预警合计。", src_research),
        (client, "车辆管理", "车辆管理看板", "各项目统计", "指挥部视角对比各项目进出场量、在场/在途数量及预警情况。", src_research),
        (client, "车辆管理", "车辆管理看板", "预警动态", "展示近期车辆监管预警摘要。", src_proto),
        (client, "车辆管理", "进出场记录", "进出记录列表", "记录车辆进场、出场信息（车牌、类型、道闸、时间）；对接工地自建车辆系统同步数据，不做准入限制。", src_proto),
        (client, "车辆管理", "进出场记录", "检索筛选", "支持按车牌/道闸/类型关键字及进出场方向筛选。", src_proto),
        (client, "车辆管理", "车辆轨迹监管", "GIS轨迹展示", "接入工地自建定位系统轨迹数据，在 GIS 地图展示车辆位置与历史轨迹。", src_bid),
        (client, "车辆管理", "车辆轨迹监管", "电子围栏", "支持创建禁行区域、允许区域围栏，地图圈选围栏范围。", src_research),
        (client, "车辆管理", "车辆轨迹监管", "轨迹报警", "对进入禁行区、偏离允许区域等行为触发报警并纳入预警清单。", src_research),
        (client, "车辆管理", "预警清单", "预警列表", "汇总车辆轨迹监管报警及车辆监管类预警（如黑名单、证件过期等），支持查询与处置跟踪。", src_proto),
        (client, "车辆管理", "预警清单", "来源筛选", "支持按「车辆监管」「轨迹监管」及处置状态筛选。", src_proto),
        (client, "车辆管理", "车辆管理", "车辆台账", "维护项目施工车辆信息：车牌、类型、所属单位、司机、准入证明、状态等。", src_research),
        (client, "车辆管理", "车辆管理", "增删改查", "支持新增、编辑、删除车辆台账；支持导入导出。", src_research),
        (client, "车辆管理", "车辆管理", "授权道闸", "为车辆配置可通行道闸；支持单车授权与批量授权。", src_proto),
        (client, "车辆管理", "设备管理", "设备台账", "管理轨迹监测设备、道闸车牌识别设备，含设备编号、类型、位置、在线状态。", src_research),
        (client, "车辆管理", "设备管理", "设备绑定", "维护设备与项目、车牌（轨迹终端）的绑定关系。", src_research),
        (client, "车辆管理", "数据对接", "车辆系统接口", "对接工地自建车辆系统，同步车辆资质、进出场记录及轨迹数据（VEH-01/03）。", src_bid),
    ]
    return rows


def build_screenshot_sections():
    sections = [
        ("4.1 劳务看板", "指挥部/项目视角展示人员在场、工种与年龄结构、出勤趋势及最新预警动态。", "01-labor-dashboard.png"),
        ("4.2 人员实名制", "左侧项目树 + 人员列表，支持检索、入场状态筛选；未对接项目可新增人员。", "02-labor-realname.png"),
        ("4.3 预警清单（人员）", "汇总实名制各类预警，项目树展示待处置数量，支持详情查看与处置。", "03-labor-warning-list.png"),
        ("4.4 实名制配置", "指挥部配置预警规则与分级管控；项目节点配置现场实名制对接开关。", "04-labor-warning-config.png"),
        ("4.5 车辆管理看板", "展示各项目车辆进出场量、在场/在途数量及异常预警统计。", "05-vehicle-dashboard.png"),
        ("4.6 进出场记录", "记录车辆进场、出场信息，不做准入限制，支持方向与关键字筛选。", "06-vehicle-access.png"),
        ("4.7 车辆轨迹监管", "GIS 地图展示车辆轨迹与电子围栏，支持创建禁行/允许区域。", "07-vehicle-track.png"),
        ("4.8 车辆管理（台账）", "维护车辆档案，支持授权道闸、批量授权及导入导出。", "08-vehicle-registry-warning.png"),
        ("4.9 预警清单（车辆）", "汇总轨迹监管及车辆监管类预警，支持来源与状态筛选。", "09-vehicle-warning-list.png"),
    ]
    return sections


def add_confirm_table(doc):
    doc.add_paragraph()
    headers = ["确认方", "姓名", "意见", "签字 / 日期"]
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i in range(1, 5):
        for j in range(4):
            set_cell_text(table.rows[i].cells[j], "")
    format_table_rows(table)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title = doc.add_heading("安全管理（人员实名制、车辆管理）- 需求确认书", level=1)
    set_heading_space_after(title)
    for run in title.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph("深圳机场扩建工程 · 智慧工程建设管控一体化平台")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph()
    add_meta_table(doc)

    add_heading(doc, "背景与目标", level=2)
    add_para(
        doc,
        "深圳机场扩建工程参建人员多、车辆流动性大，各项目现场普遍建设智慧工地劳务系统、车辆管理系统，但数据分散、标准不一，指挥部难以及时掌握人员准入、考勤异常及重点车辆监管情况。需在智慧工程建设管控一体化平台中建设人员实名制管理与车辆管理模块，实现数据汇聚、规则预警、分级处置与指挥调度联动。",
    )
    add_para(
        doc,
        "依托标准接口对接各项目劳务/闸机/车辆系统，建立人员准入核验与预警规则、车辆进出场记录与轨迹监管能力；对黑名单、证书过期、超时作业、轨迹越界等风险统一预警、跟踪闭环，支撑指挥部与项目方协同管控，提升扩建工程安全精细化管理水平。",
        bold_prefix="建设目标：",
    )

    add_heading(doc, "二、功能范围", level=2)
    add_para(
        doc,
        "本文档聚焦一体化平台「人员实名制管理」「车辆管理」两大模块，面向工程指挥部及各项目安全、劳务管理人员。人员实名制涵盖劳务看板、人员档案、考勤明细与统计、劳务黑名单、预警清单及实名制配置；车辆管理涵盖车辆看板、进出场记录、轨迹监管、预警清单、车辆台账与设备管理。本文档基于当前原型页面整理功能清单与界面说明，附核心页面截图供业务方评审确认。本期不含工资监管、机械设备监测、安全巡检等模块（另行确认）。",
    )

    add_heading(doc, "三、功能清单", level=2)
    add_function_table(doc, build_function_rows())

    add_heading(doc, "四、核心页面截图", level=2)
    add_para(doc, "以下截图为当前原型核心页面，供业务方核对布局与功能范围。")
    for title, desc, fname in build_screenshot_sections():
        img = SHOT_DIR / fname if (SHOT_DIR / fname).exists() else None
        add_screenshot_section(doc, title, desc, img)

    add_heading(doc, "五、确认意见", level=2)
    add_para(doc, "请业务方对以上功能范围、功能描述及原型截图进行确认。如有补充或调整意见，请在下方表格填写。")
    add_confirm_table(doc)

    doc.save(OUT_FILE)
    print(f"已生成: {OUT_FILE}")


if __name__ == "__main__":
    main()
