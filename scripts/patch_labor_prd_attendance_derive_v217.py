# -*- coding: utf-8 -*-
"""V2.1.7：考勤明细由进出场流水派生；备份并升版。"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
DOC = BASE / "人员实名制_PRD_评审通过.docx"
BACKUP = BASE / "人员实名制_PRD_评审通过_backup_20260821.docx"

PARA_EXACT = [
    (
        "项目级人员进出场明细只读（一期经 ROMA），支持按日期/关键词/进出场/在场/工种筛选与导出；身份证脱敏。无编辑、补录、设备管理。",
        "项目级考勤日汇总只读：一人一日一行。进出场流水由一期系统上报，或由施工现场实名制系统按标准接口上报；"
        "列表中的进出场、进/出场时间与闸机、工时、在场状态均由当日进出场记录按规则计算。"
        "支持按日期/关键词/进出场/在场/工种筛选与导出；身份证脱敏；操作列可查看进出场记录。"
        "无编辑、补录、设备管理；不直连闸机私有协议。",
    ),
    (
        "只读查询 ROMA 考勤流水。",
        "只读查询进出场记录（入站真源）及由其派生的考勤日汇总；支持按人按日查看进出场记录弹窗。",
    ),
    (
        "无独立状态机；展示派生在场与进出场枚举。",
        "无独立状态机。进出场、在场状态由进出场记录派生，不单独流转。",
    ),
    (
        "核心实体见本文第5章「考勤明细记录」。",
        "核心实体见本文第5章「进出场记录」「考勤明细（日汇总）」。",
    ),
    (
        "考勤明细记录（只读）",
        "进出场记录（入站真源，只读）；考勤明细日汇总（派生，只读）",
    ),
]

# 业务规则与边界 bullet 整段替换（若存在）
PARA_CONTAINS_REPLACE = [
    (
        "只读；不做手工补录、不做考勤机设备管理",
        "只读；不做手工补录、不做考勤机设备管理；不直连闸机私有协议\n"
        "进出场记录来源：一期系统上报，或施工现场实名制系统按标准接口上报\n"
        "日汇总派生规则（统计日记为 D）：\n"
        "① 进场时间/进场闸机：取日期=D 的进场记录中时间最早的一条；无则空（界面「—」）\n"
        "② 出场时间/出场闸机：优先取日期=D 的出场记录中时间最晚的一条；若当日无出场，则取「当日最早进场」之后时间上第一条出场（允许跨自然日补全）；仍无则空（「—」）\n"
        "③ 进出场：有出场→已出场；仅有进场无出场→已进场；仅有出场无进场→已出场\n"
        "④ 在场状态：有进场且尚无有效出场→在场；有出场→不在场；仅有出场无进场→不在场；当日无进出→「—」\n"
        "⑤ 工时：有进有出（含跨日补全）→ min(出场, D 23:59:59) − 进场；仅有出场无进场→ min(出场, D 23:59:59) − D 00:00:00；仅有进场尚无出场→「—」。工时不计入 D 日 23:59:59 之后的时段\n"
        "⑥ 列表以现场设备报送的进出场数据为准；花名册已离场人员若仍有报送流水，仍可出现在考勤明细\n"
        "出场时间允许为空（未出场显示「—」）\n"
        "导出：按列表字段全量导出当前筛选结果，本期不描述异步细节",
    ),
]

META = {
    "版本号": "V2.1.7",
    "更新日期": "2026-08-21",
}

REVISION_ROW = (
    "V2.1.7",
    "—",
    "2026-08-21",
    "考勤明细：明确进出场记录为入站真源（一期上报或现场实名制标准接口）；"
    "日汇总「进出场/进场时间/进场闸机/出场时间/出场闸机/工时/在场状态」按进出场记录计算"
    "（进场取当日最早进场；出场取当日最晚出场，无则跨日补全首条出场；工时截断至当日 23:59:59；"
    "仅出场无进场按 0 点起算且状态为已出场/不在场）；列表以设备报送为准；同步产品架构与开发辅助。",
)

CELL_EXACT = {
    ("中文实体", "考勤明细", "ATTENDANCE_RECORD", "进出场流水（只读）"): (
        None,
        None,
    ),  # handled specially
    (
        "一期实名制 + ROMA",
        "人员主档、特种证、考勤明细",
        "只读展示、统计、预警规则输入",
        "人员/考勤无数据，阻断台账/考勤/统计",
    ): (
        "一期实名制 + ROMA",
        "人员主档、特种证；进出场/考勤经 ROMA 或同源入站",
        "只读展示、统计、预警规则输入；考勤日汇总由进出场记录派生",
        "人员无数据阻断台账/统计；无进出场流水则考勤明细为空",
    ),
    (
        "标准实名制上报接口（路径B）",
        "闸机/实名制子系统按标准接口上报的人员主档及关联信息",
        "一期未对接项目的人员数据入站；台账/统计/预警同源消费",
        "该项目无上报数据则台账为空；不影响已走 ROMA 的其他项目",
    ): (
        "标准实名制上报接口（路径B）",
        "闸机/实名制子系统按标准接口上报的人员主档、进出场记录及关联信息",
        "一期未对接项目的人员与进出场入站；台账/考勤日汇总/统计/预警同源消费",
        "该项目无上报数据则台账/考勤为空；不影响已走 ROMA 的其他项目",
    ),
    (
        "F-LABOR-05",
        "考勤明细",
        "考勤明细",
        "进出场明细只读筛选与导出",
    ): (
        "F-LABOR-05",
        "考勤明细",
        "考勤明细",
        "进出场记录入站；日汇总派生只读筛选/导出；可查看进出场记录",
    ),
    (
        "考勤明细",
        "`/labor/attendance-detail`",
        "筛选 + 指标 + 只读列表",
        "项目 Web",
    ): (
        "考勤明细",
        "`/labor/attendance-detail`",
        "筛选 + 指标 + 日汇总列表 + 查看进出场记录",
        "项目 Web；列由进出场记录派生",
    ),
    (
        "entry_status（进出场）",
        "已进场（clock_in）/ 已出场（clock_out）",
        "当日刷脸结果摘要",
    ): (
        "entry_status（进出场）",
        "已进场（clock_in）/ 已出场（clock_out）",
        "由当日进出场记录派生的日摘要，非手工录入",
    ),
    (
        "AC-01",
        "只读可筛",
        "存在 ROMA 考勤",
        "打开考勤明细并筛选",
        "可筛且不可编辑",
    ): (
        "AC-01",
        "只读可筛",
        "存在进出场上报流水",
        "打开考勤明细并筛选",
        "日汇总可筛且不可编辑；进出场相关列与派生规则一致",
    ),
    (
        "考勤明细同步",
        "一期经 ROMA",
        "入站",
        "考勤明细记录",
    ): (
        "进出场记录上报",
        "一期系统或现场实名制标准接口",
        "入站",
        "进出场记录（考勤日汇总由其派生）",
    ),
    (
        "进出场",
        "当日经考勤机刷脸形成的进入、离开记录；与在岗状态不是同一概念",
    ): (
        "进出场",
        "经闸机/实名制设备形成的进入、离开流水事件；与在岗状态不是同一概念；"
        "考勤明细「进出场」列为当日流水派生摘要",
    ),
    (
        "在场",
        "针对在岗人员：当日已进场且尚未出场为在场；未进场或已出场为不在场；离场不适用",
    ): (
        "在场",
        "按进出场记录派生：当日已进场且尚未出场为在场；已出场或不在场；"
        "仅有出场无进场视为不在场；列表以设备报送为准（含花名册已离场但仍有流水者）",
    ),
}

FIELD_CELL_UPDATES = {
    # TABLE52 list fields — key by (字段名, 中文名) first two meaningful cells
    ("clock_in", "进场时间"): "派生：取统计日 D 最早一条进场记录的时间；无则「—」",
    ("clock_out", "出场时间"): "派生：优先取 D 最晚一条出场；若无则取最早进场之后首条出场（可跨日）；无则「—」",
    ("gate_in", "进场闸机"): "派生：与当日最早进场记录同条闸机",
    ("gate_out", "出场闸机"): "派生：与所采用出场记录同条闸机",
    ("entry_status", "进出场"): "派生：有出场→已出场；仅进无出→已进场；仅出无进→已出场",
    ("work_hours", "工时"): "派生：有进有出→min(出场,D 23:59:59)−进场；仅出→min(出场,D 23:59:59)−D 00:00:00；仅进无出→「—」",
    ("on_site_status", "在场状态"): "派生：进且无出→在场；有出→不在场；仅出无进→不在场；无流水→「—」",
}


def set_para_text(p, value: str) -> None:
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    first = cell.paragraphs[0]
    set_para_text(first, value)
    for p in cell.paragraphs[1:]:
        set_para_text(p, "")


def format_table(table) -> None:
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_heading_space_after(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def cell_key(row) -> tuple[str, ...]:
    return tuple(c.text.strip().replace("\n", " ") for c in row.cells)


def update_meta(doc: Document) -> None:
    t0 = doc.tables[0]
    for row in t0.rows:
        k = row.cells[0].text.strip()
        if k in META:
            set_cell_text(row.cells[1], META[k])


def insert_revision(doc: Document) -> None:
    t1 = doc.tables[1]
    # insert after header as new first data row visually = after R0
    # python-docx: add_row at end then move — simpler append after finding V2.1.6
    hdr = [c.text.strip() for c in t1.rows[0].cells]
    # insert new row at index 1 by cloning last and shifting — use add_row then reorder XML
    new_row = t1.add_row()
    for i, val in enumerate(REVISION_ROW):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], val)
    tbl = t1._tbl
    tr = new_row._tr
    tbl.remove(tr)
    # insert after header
    hdr_tr = t1.rows[0]._tr
    hdr_tr.addnext(tr)
    format_table(t1)


def patch_entity_table(doc: Document) -> None:
    """TABLE12: replace 考勤明细 row and insert 进出场记录."""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:3] != ["中文实体", "英文编码", "层级"]:
            continue
        for row in table.rows[1:]:
            if row.cells[0].text.strip() == "考勤明细":
                set_cell_text(row.cells[0], "考勤明细（日汇总）")
                set_cell_text(row.cells[1], "ATTENDANCE_DAY")
                set_cell_text(row.cells[2], "项目")
                set_cell_text(row.cells[3], "一人一日派生汇总（只读，由进出场记录计算）")
                # insert access record row before this one
                new_row = table.add_row()
                set_cell_text(new_row.cells[0], "进出场记录")
                set_cell_text(new_row.cells[1], "PERSON_ACCESS_RECORD")
                set_cell_text(new_row.cells[2], "项目")
                set_cell_text(new_row.cells[3], "闸机进出场流水（入站真源，只读）")
                tbl = table._tbl
                tr = new_row._tr
                tbl.remove(tr)
                row._tr.addprevious(tr)
                format_table(table)
                return


def patch_field_table52(doc: Document) -> None:
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:3] != ["字段名", "中文名", "类型"]:
            continue
        # attendance list fields table has clock_in
        names = {row.cells[0].text.strip() for row in table.rows[1:]}
        if "clock_in" not in names:
            continue
        explain_idx = 4 if len(h) > 4 else len(h) - 1
        for row in table.rows[1:]:
            fn = row.cells[0].text.strip()
            cn = row.cells[1].text.strip()
            key = (fn, cn)
            if key in FIELD_CELL_UPDATES and explain_idx < len(row.cells):
                set_cell_text(row.cells[explain_idx], FIELD_CELL_UPDATES[key])
        # ensure work_hours / on_site rows exist updates
        format_table(table)
        return


def patch_cells(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            for old, new in CELL_EXACT.items():
                if old[0] is None:
                    continue
                if len(cells) < len(old):
                    continue
                if tuple(cells[: len(old)]) == old:
                    for i, val in enumerate(new):
                        if i < len(row.cells) and val is not None:
                            set_cell_text(row.cells[i], val)
                    n += 1
                    format_table(table)
                    break
    return n


def patch_paragraphs(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in PARA_EXACT:
            if t == old:
                set_para_text(p, new)
                n += 1
                break
        else:
            for needle, new in PARA_CONTAINS_REPLACE:
                if needle in t and "进出场记录来源" not in t:
                    set_para_text(p, new)
                    n += 1
                    break
    return n


def patch_core_flow(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text
        if t.startswith("核心流程：") and "一期经 ROMA 同步人员主档" in t:
            new = (
                "核心流程：\n"
                "1）人员主档：一期经 ROMA，或现场实名制按标准接口上报 → 项目 Web 台账、看板与指挥部统计只读展示；\n"
                "2）进出场记录：一期系统上报，或施工现场实名制系统按标准接口上报 → 落库为进出场流水；"
                "考勤明细日汇总（进/出场时间与闸机、进出场、工时、在场状态）由其按规则计算；\n"
                "3）项目侧仅可补录三级安全教育（本地保存，ROMA 回写不覆盖）；\n"
                "4）按项目实名制配置（轨迹外链、默认接收人、分级管控、9 条规则）定时生成/关闭预警；通知类一次性生成不重生；\n"
                "5）处置任务类预警仅默认接收人可在 Web（个人中心·预警中心或预警清单）处置；自动关闭类次日定时关闭；APP 预警中心仅查看；\n"
                "6）黑名单全平台维护，日扫各项目在岗命中，每项目各生成 1 条预警，引导处理到位后次日自动关闭；\n"
                "7）人员轨迹仅外链跳转：指挥部有地址可跳（含停用），项目侧须启用且地址有效。"
            )
            set_para_text(p, new)
            return


def add_acceptance_rows(doc: Document) -> None:
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:3] != ["编号", "验收项", "Given"]:
            continue
        # find attendance AC table by AC-01 只读可筛
        body = " ".join(c.text for c in table.rows[1].cells) if len(table.rows) > 1 else ""
        if "只读可筛" not in body and "考勤" not in body:
            # check any row
            joined = " ".join(c.text for r in table.rows for c in r.cells)
            if "打开考勤明细" not in joined:
                continue
        existing = {r.cells[0].text.strip() for r in table.rows}
        extras = [
            (
                "AC-03",
                "日汇总派生",
                "同日多条进/出场流水",
                "打开考勤明细某日",
                "进场取最早进场；出场取最晚出场（无则跨日补全）；与弹窗流水一致",
            ),
            (
                "AC-04",
                "仅出场",
                "某日仅有出场无进场",
                "查看日汇总行",
                "进出场=已出场；在场=不在场；工时自 D 00:00:00 起算至出场（截断 23:59:59）",
            ),
            (
                "AC-05",
                "查看流水",
                "存在进出场记录",
                "点击查看进出场记录",
                "弹窗展示该人该日（可改日期）进出场流水；只读",
            ),
        ]
        for row_vals in extras:
            if row_vals[0] in existing:
                continue
            row = table.add_row()
            for i, v in enumerate(row_vals):
                if i < len(row.cells):
                    set_cell_text(row.cells[i], v)
        format_table(table)
        return


def main() -> None:
    if not DOC.exists():
        raise SystemExit(f"missing {DOC}")
    shutil.copy2(DOC, BACKUP)
    print("backup ->", BACKUP)

    doc = Document(str(DOC))
    update_meta(doc)
    insert_revision(doc)
    n_p = patch_paragraphs(doc)
    patch_core_flow(doc)
    n_c = patch_cells(doc)
    patch_entity_table(doc)
    patch_field_table52(doc)
    add_acceptance_rows(doc)
    set_heading_space_after(doc)

    doc.save(str(DOC))
    print(f"saved {DOC} paras={n_p} cells={n_c}")

    # verify
    d2 = Document(str(DOC))
    ver = d2.tables[0].rows[1].cells[1].text.strip()
    rev0 = d2.tables[1].rows[1].cells[0].text.strip()
    print("version", ver, "first_rev", rev0)
    assert ver == "V2.1.7", ver
    assert rev0 == "V2.1.7", rev0


if __name__ == "__main__":
    main()
