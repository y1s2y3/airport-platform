# -*- coding: utf-8 -*-
"""在 6.7.8「预警触发\\消除规则」后同级补充「预警中心展示说明文案」。"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

DOC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)
BACKUP = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_备份_V2.1.3_20260815.docx"
)

# 预警中心展示说明文案（9 类，与配置规则一一对应）
COPY_ROWS = [
    [
        "未进行入场三级教育预警",
        "处置任务（系统自动关闭）",
        "#姓名#未按要求完成入场三级安全教育，请前往人员实名制补录教育信息；存在教育记录后，次日将自动关闭预警。",
        "占位符：#姓名#",
    ],
    [
        "特种作业证书缺失/过期预警",
        "处置任务（系统自动关闭）",
        "#姓名#特种作业证书缺失或已过期，请联系闸机/一期系统更新证书及有效期；有效信息回写后，次日将自动关闭预警。",
        "占位符：#姓名#",
    ],
    [
        "连续工作超12小时预警",
        "处置任务（手动处理）",
        "#姓名#连续作业时长超过项目配置阈值（默认12小时），请核实并督促休息、规范进出场；请在预警详情填写处置说明并点击「处置并关闭」。",
        "占位符：#姓名#；阈值以实名制配置为准",
    ],
    [
        "实名制年龄低于16周岁预警",
        "处置任务（系统自动关闭）",
        "#姓名#年龄低于实名制年龄下限（默认16周岁），请联系闸机系统更正信息或办理退场；条件满足后，次日将自动关闭预警。",
        "占位符：#姓名#；下限以实名制配置为准",
    ],
    [
        "高龄提醒（男65/女60）",
        "通知",
        "#姓名#年龄超过高龄提醒阈值（默认男65岁/女60岁），请做好工人健康情况排查；本条为通知类，无需关闭、不参与分级上报。",
        "占位符：#姓名#；阈值以实名制配置为准",
    ],
    [
        "身份证过期提醒",
        "通知",
        "#姓名#身份证已过期，请督促换证并更新证件信息与有效期；本条为通知类，无需关闭、不参与分级上报。",
        "占位符：#姓名#",
    ],
    [
        "连续x天未出勤预警",
        "处置任务（手动处理）",
        "#姓名#连续#天数#天无考勤记录，请联系参建单位核实原因并督促整改；请在预警详情填写处置说明并点击「处置并关闭」。",
        "占位符：#姓名#、#天数#（取配置阈值，默认3）",
    ],
    [
        "管理人员考勤不达标，每月出勤少于x天预警",
        "处置任务（手动处理）",
        "#姓名#当月出勤天数少于配置阈值（默认少于20天），请督促落实管理人员出勤要求；请在预警详情填写处置说明并点击「处置并关闭」。",
        "占位符：#姓名#；阈值以实名制配置为准",
    ],
    [
        "黑名单人员进场预警",
        "处置任务（系统自动关闭）",
        "#姓名#，身份证号:#身份证号（脱敏）#，命中劳务黑名单且当前在岗，请联系闸机系统办理退场或确认是否误录入；引导处理到位后，次日将自动关闭预警。",
        "占位符：#姓名#、#身份证号（脱敏）#；平台不做闸机强制拦截",
    ],
]


def set_para_text(p, value: str):
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str):
    if not cell.paragraphs:
        cell.text = value
        return
    set_para_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        # 同级小节标题加粗
        if text in ("预警中心展示说明文案",):
            run.bold = True
    return new_para


def format_table(table):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_heading_space_after(doc):
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def ensure_backup():
    if not BACKUP.exists():
        shutil.copy2(DOC, BACKUP)
    print("backup:", BACKUP, BACKUP.exists(), BACKUP.stat().st_size)


def update_meta_revision(doc):
    meta = doc.tables[0]
    for row in meta.rows:
        k = row.cells[0].text.strip()
        if k == "版本号":
            set_cell_text(row.cells[1], "V2.1.4")
        if k == "更新日期":
            set_cell_text(row.cells[1], "2026-08-15")

    rev = doc.tables[1]
    for row in rev.rows[1:]:
        if row.cells[0].text.strip() == "V2.1.4":
            set_cell_text(
                row.cells[3],
                "实名制配置：在「预警触发\\消除规则」后同级补充「预警中心展示说明文案」，"
                "按 9 类预警给出预警中心展示文案规则。",
            )
            set_cell_text(row.cells[2], "2026-08-15")
            return
    rev._tbl.append(deepcopy(rev._tbl.tr_lst[-1]))
    row = rev.rows[-1]
    vals = [
        "V2.1.4",
        "—",
        "2026-08-15",
        "实名制配置：在「预警触发\\消除规则」后同级补充「预警中心展示说明文案」，"
        "针对每类预警给出预警中心展示说明文案规则（含占位符约定）。",
    ]
    for i, v in enumerate(vals):
        set_cell_text(row.cells[i], v)
    format_table(rev)


def find_anchor(doc):
    """定位「消除规则」补充段落后、6.7.9 之前的锚点。"""
    if any(p.text.strip() == "预警中心展示说明文案" for p in doc.paragraphs):
        return None
    supplement = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("补充：自动关闭类预警统一为"):
            supplement = p
        if t.startswith("6.7.9") and supplement is not None:
            return supplement
    # fallback：直接找 预警触发\消除规则
    for p in doc.paragraphs:
        if "预警触发" in p.text and "消除规则" in p.text:
            return p
    return None


def insert_copy_section(doc):
    anchor = find_anchor(doc)
    if anchor is None:
        if any(p.text.strip() == "预警中心展示说明文案" for p in doc.paragraphs):
            print("section already exists, skip insert")
            return
        raise RuntimeError("找不到插入锚点（预警触发\\消除规则 补充段）")

    cur = insert_paragraph_after(anchor, "")
    cur = insert_paragraph_after(cur, "预警中心展示说明文案", style="Normal")
    # 加粗标题
    if cur.runs:
        cur.runs[0].bold = True
    cur = insert_paragraph_after(
        cur,
        "每一类预警在个人中心「预警中心」详情/引导处需展示对应说明文案。"
        "文案按规则类型配置，推送或打开详情时用触发人员信息替换占位符后展示。"
        "通用占位符：#姓名#、#天数#、#身份证号（脱敏）#；带配置阈值的类型以实名制配置当前值为准。",
        style="Normal",
    )

    # 在正文后插入表格：先建空段，再把 table XML 挂到该段之后
    table_holder = insert_paragraph_after(cur, "")
    table = doc.add_table(rows=1 + len(COPY_ROWS), cols=4)
    table.style = "Table Grid"
    headers = ["预警类型", "预警中心类型", "预警中心展示说明文案", "备注"]
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h)
        for run in table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, row_data in enumerate(COPY_ROWS):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i + 1].cells[j], val)
    format_table(table)

    # 将 add_table 追加在文末的表移到 table_holder 之后
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    table_holder._p.addnext(tbl)
    print("inserted copy section + table rows=", 1 + len(COPY_ROWS))


def main():
    ensure_backup()
    doc = Document(str(DOC))
    update_meta_revision(doc)
    insert_copy_section(doc)
    set_heading_space_after(doc)
    doc.save(str(DOC))
    print("saved:", DOC)


if __name__ == "__main__":
    main()
