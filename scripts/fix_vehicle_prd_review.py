# -*- coding: utf-8 -*-
"""车辆 PRD：按评审决议修订（文本替换 + 标题重编号；不插表行）。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx")
LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_vehicle_fix_log.txt"
)
TODAY = "2026-08-13"
VER = "V2.2.3"
NOTE = (
    "评审修订：去掉项目看板表述；项目角色不含 F-VEHICLE-01；"
    "轨迹系统验收与「有 URL 可跳（含停用）」对齐；上报统一 project_id；"
    "章节顺延重编号；明确在场为项目内判定、指挥部跨项目相加；"
    "6.6 去掉两端看板；业务规则补幂等（无 external_id 时本期须带或拒绝）。"
)

# paragraph / cell text replacements (order matters)
REPLACEMENTS: list[tuple[str, str]] = [
    # L1 管/不管
    (
        "项目车辆管理看板、进出场记录、车辆轨迹监管外跳",
        "项目进出场记录、车辆轨迹监管外跳",
    ),
    # L4 / 文案
    ("第三方按项目编码上报", "第三方按项目 ID（project_id）上报"),
    ("按项目编码上报", "按 project_id 上报"),
    ("项目进出场记录与两端看板近实时可见", "指挥部看板与项目进出场记录近实时可见"),
    ("列表与看板近实时可见", "指挥部看板与项目进出场记录近实时可见"),
    ("指标口径见 6.1.10", "指标口径见 6.1.9"),
    ("项目编码无法识别", "项目 ID 无法识别"),
    ("缺车牌或项目编码无效", "缺车牌或项目 ID 无效"),
    ("项目编码，对应平台项目主数据", "平台项目 ID，对应项目主数据；决定数据归属项目"),
    ("字段/项目编码对齐", "字段/project_id 对齐"),
    ("项目编码解析", "项目 ID 校验"),
    ("上报 project_code", "上报 project_id"),
    ("上报 projectCode", "上报 project_id"),
    ("project_code", "project_id"),
    ("projectCode", "project_id"),
    # 在场口径
    (
        "在场车辆 = 按车牌取全量最新一条进出场记录，最新方向为「进场」的车牌数（同车牌只计 1；含跨夜）",
        "在场车辆 = 在「统计范围内」按车牌取该范围内最新一条进出场记录，最新方向为「进场」的车牌数（同车牌只计 1；含跨夜）。项目级统计范围=当前项目；指挥部跨项目汇总=各项目在场数相加（同一车牌可在多项目同时计入）",
    ),
    (
        "按车牌取全量最新一条进出场记录；方向为「进场」则计 1 辆（含跨夜）",
        "在统计范围内按车牌取最新一条进出场记录；方向为「进场」则计 1 辆（含跨夜）。项目级=单项目内；指挥部=各项目分别计算后再相加",
    ),
    (
        "该车牌全量最新进出场记录方向为进场",
        "该车牌在本项目内最新进出场记录方向为进场",
    ),
    (
        "无记录，或全量最新为出场",
        "本项目无记录，或本项目内最新为出场",
    ),
    # L3 related narrative already OK; AC fixed in table cells
    ("无审批；。", "无审批流。"),
    ("无审批流。。", "无审批流。"),
    ("对齐架构 §0.10。", "上报字段与规则见本节 6.6.8～6.6.11。"),
    ("对齐架构 §0.10", "上报字段与规则见本节 6.6.8～6.6.11"),
]

# Heading renumber map: apply longest keys first
RENUMBER = [
    ("6.6.12 ", "6.6.11 "),
    ("6.6.11 ", "6.6.10 "),
    ("6.6.10 ", "6.6.9 "),
    ("6.5.12 ", "6.5.11 "),
    ("6.5.11 ", "6.5.10 "),
    ("6.5.10 ", "6.5.9 "),
    ("6.2.12 ", "6.2.11 "),
    ("6.2.11 ", "6.2.10 "),
    ("6.1.12 ", "6.1.11 "),
]


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def apply_text(s: str, pairs: list[tuple[str, str]]) -> tuple[str, int]:
    n = 0
    new = s
    for old, repl in pairs:
        if old in new:
            c = new.count(old)
            new = new.replace(old, repl)
            n += c
    return new, n


def replace_paragraph(para, pairs: list[tuple[str, str]]) -> int:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    new, n = apply_text(full, pairs)
    if new != full:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = new
    return n


def bump_meta(doc: Document) -> None:
    meta = doc.tables[0]
    for row in meta.rows:
        key = row.cells[0].text.strip()
        if key == "版本号" and len(row.cells) > 1:
            set_cell_text(row.cells[1], VER)
        if key == "更新日期" and len(row.cells) > 1:
            set_cell_text(row.cells[1], TODAY)


def write_revision_via_word(path: Path, log: list[str]) -> None:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(path), ReadOnly=False)
        if doc.Tables.Count < 2:
            log.append("no rev table")
            doc.Close(False)
            return
        tbl = doc.Tables(2)
        first = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first == VER:
            log.append("rev already " + VER)
        else:
            tbl.Rows(2).Select()
            word.Selection.InsertRowsAbove(1)
            row = tbl.Rows(2)
            # 版本|日期|说明  (3 col) or 版本|修订人|日期|说明
            ncols = row.Cells.Count
            if ncols >= 4:
                vals = [VER, "—", TODAY, NOTE]
            else:
                vals = [VER, TODAY, NOTE]
            for i, v in enumerate(vals, start=1):
                if i <= ncols:
                    row.Cells(i).Range.Text = v
            log.append("inserted rev " + VER)
        doc.Save()
        doc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def sanitize_word(src: Path, dst: Path, log: list[str]) -> bool:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(src), ReadOnly=False)
        if dst.exists():
            dst.unlink()
        doc.SaveAs2(str(dst), FileFormat=12)
        doc.Close(False)
        log.append("sanitized -> " + dst.name)
        return True
    except Exception as e:
        log.append("sanitize fail " + str(e))
        return False
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def patch_tables(doc: Document, log: list[str]) -> None:
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            blob = " | ".join(cells)

            # L2 roles
            if "项目经理" in blob or "车辆相关项目角色" in blob:
                for ci, cell in enumerate(row.cells):
                    t = cell.text
                    if "F-VEHICLE-01" in t:
                        new = (
                            t.replace("F-VEHICLE-01、04、05", "F-VEHICLE-04、05")
                            .replace("F-VEHICLE-01、04", "F-VEHICLE-04")
                            .replace("F-VEHICLE-01、05", "F-VEHICLE-05")
                            .replace("F-VEHICLE-01，", "")
                            .replace("F-VEHICLE-01,", "")
                            .replace("、F-VEHICLE-01", "")
                            .replace("F-VEHICLE-01、", "")
                            .replace("F-VEHICLE-01", "")
                        )
                        # clean double顿号
                        while "、、" in new:
                            new = new.replace("、、", "、")
                        new = new.strip("、，, ")
                        if new != t:
                            set_cell_text(cell, new)
                            log.append(f"L2 role T{ti}R{ri}C{ci}")

            # L3 AC-01 track system: 启用入列 -> URL非空入列
            if cells and cells[0] == "AC-01" and "启用入列" in blob:
                # Given 项目配置合法 URL 且启用 -> URL 非空（不论启用）
                for ci, cell in enumerate(row.cells):
                    t = cell.text.strip()
                    if t == "启用入列":
                        set_cell_text(cell, "URL非空入列")
                        log.append(f"L3 AC title T{ti}R{ri}")
                    if "合法 URL 且启用" in t or ("启用" in t and "URL" in t and "Given" not in blob):
                        # second cell often 验收项, third Given
                        pass
                # rewrite by position: 编号|验收项|Given|When|Then
                if len(row.cells) >= 5:
                    set_cell_text(row.cells[1], "URL非空入列")
                    set_cell_text(row.cells[2], "项目配置系统地址（URL）非空（启用或停用均可）")
                    set_cell_text(row.cells[3], "打开车辆轨迹系统列表")
                    set_cell_text(row.cells[4], "出现该项目且可跳转")
                    log.append(f"L3 AC-01 rewritten T{ti}R{ri}")

            # L4 field table: project_id 中文名 项目编码 -> 项目ID
            if len(row.cells) >= 3 and "project_id" in row.cells[1].text and "上报报文" in row.cells[0].text:
                # 字段名|中文名 might be col1 col2
                for ci, cell in enumerate(row.cells):
                    if cell.text.strip() == "项目编码":
                        set_cell_text(cell, "项目ID")
                        log.append(f"L4 cn name T{ti}R{ri}")
                # also description
                last = row.cells[-1].text
                if "项目编码" in last:
                    set_cell_text(row.cells[-1], last.replace("项目编码", "项目 ID"))

            # generic cell replacements
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_paragraph(para, REPLACEMENTS)


def add_idempotent_rules(doc: Document, log: list[str]) -> None:
    """在 6.6.9/6.6.10 业务规则列表中补充幂等（标题可能已重编号）。"""
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if "业务规则与边界" in t and t.startswith("6.6"):
            target_idx = i
            break
    if target_idx is None:
        log.append("WARN: no 6.6 business rules heading")
        return

    # collect following list bullets until next heading
    bullets = []
    j = target_idx + 1
    while j < len(doc.paragraphs):
        p = doc.paragraphs[j]
        style = p.style.name if p.style else ""
        t = p.text.strip()
        if style.startswith("Heading"):
            break
        if t:
            bullets.append((j, t))
        j += 1

    has_idem = any("幂等" in t for _, t in bullets)
    if has_idem:
        log.append("idempotent already in rules")
        return

    # Append by cloning last list paragraph text into a new run after last bullet
    # Find last list bullet paragraph under section
    last_list = None
    for j in range(target_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            break
        if style.startswith("List") or p.text.strip().startswith("唯一写入") or p.text.strip():
            if p.text.strip():
                last_list = p

    if last_list is None:
        log.append("WARN: no list to append idempotent")
        return

    # Insert after last_list using XML clone
    from copy import deepcopy

    new_el = deepcopy(last_list._element)
    # clear text in new element
    for node in new_el.iter():
        if node.tag.endswith("}t"):
            node.text = ""
    last_list._element.addnext(new_el)

    # find the new paragraph object
    # After insert, python-docx paragraphs refresh on access
    # Set text on the element
    texts = [
        "幂等：同一 project_id + external_id（若提供）重复上报不产生重复流水。",
        "无 external_id：本期要求上报方提供 external_id；缺失则拒绝该条并返回业务错误（不落库）。",
    ]
    # We only inserted one node; insert second
    second = deepcopy(new_el)
    new_el.addnext(second)

    def set_el_text(el, text: str) -> None:
        # put all text into first w:t
        first = True
        for node in el.iter():
            if node.tag.endswith("}t"):
                if first:
                    node.text = text
                    first = False
                else:
                    node.text = ""

    set_el_text(new_el, texts[0])
    set_el_text(second, texts[1])
    log.append("added idempotent rules")


def main() -> None:
    log: list[str] = []
    bak = SRC.with_name(f"车辆管理_PRD_评审修订前_副本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    shutil.copy2(SRC, bak)
    log.append("backup " + bak.name)

    doc = Document(str(SRC))
    total = 0
    for para in doc.paragraphs:
        # renumber headings first on full text
        full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
        new = full
        for old, repl in RENUMBER:
            if new.startswith(old) or f" {old}" in new:
                # only replace at start of heading text
                if new.startswith(old):
                    new = repl + new[len(old) :]
                    total += 1
        if new != full:
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = new
        total += replace_paragraph(para, REPLACEMENTS)

    patch_tables(doc, log)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # renumber inside cells if any
                    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
                    new = full
                    for old, repl in RENUMBER:
                        if old in new:
                            new = new.replace(old, repl)
                    if new != full:
                        if para.runs:
                            para.runs[0].text = new
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.text = new
                    total += replace_paragraph(para, REPLACEMENTS)

    bump_meta(doc)
    add_idempotent_rules(doc, log)
    doc.save(str(SRC))
    log.append(f"python saved replacements~{total}")

    # Word revision + sanitize
    write_revision_via_word(SRC, log)
    tmp = SRC.with_name("_vehicle_prd_fix_tmp.docx")
    if sanitize_word(SRC, tmp, log):
        shutil.copy2(tmp, SRC)
        try:
            tmp.unlink()
        except Exception:
            pass

    # verify
    d2 = Document(str(SRC))
    meta = {r.cells[0].text.strip(): r.cells[1].text.strip() for r in d2.tables[0].rows if len(r.cells) > 1}
    log.append(f"meta {meta.get('版本号')} {meta.get('更新日期')}")
    issues = []
    for p in d2.paragraphs:
        t = p.text
        if "项目车辆管理看板" in t:
            issues.append("para still 项目看板: " + t[:80])
        if "两端看板" in t:
            issues.append("para still 两端看板: " + t[:80])
        if "project_code" in t or "projectCode" in t:
            issues.append("para still project_code: " + t[:80])
        if t.startswith("6.1.12") or t.startswith("6.2.12") or t.startswith("6.6.12") or t.startswith("6.5.12"):
            issues.append("old heading left: " + t[:40])
    for t in d2.tables:
        for row in t.rows:
            blob = " | ".join(c.text for c in row.cells)
            if "项目车辆管理看板" in blob:
                issues.append("table 项目看板")
            if "project_code" in blob:
                issues.append("table project_code: " + blob[:80])
            if "项目经理" in blob and "F-VEHICLE-01" in blob:
                issues.append("role still F-VEHICLE-01: " + blob[:100])
            if "启用入列" in blob:
                issues.append("AC still 启用入列")
    log.append("issues " + str(len(issues)))
    for x in issues:
        log.append(" - " + x)

    # headings check
    for p in d2.paragraphs:
        if p.style and p.style.name.startswith("Heading") and p.text.strip().startswith("6."):
            log.append("H: " + p.text.strip()[:60])

    LOG.write_text("\n".join(log), encoding="utf-8")
    print("done", LOG)


if __name__ == "__main__":
    main()
