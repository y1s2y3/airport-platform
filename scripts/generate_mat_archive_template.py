"""生成材料设备归档 Word 模板（与 matEntryArchiveExport.js 表格索引对齐）。"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "templates" / "材料设备归档模板.docx"

TABLE_SPECS = [
    (5, 2),   # 0: GD-C1-347 报审表
    (6, 7),   # 1: GD-C1-341 合格证收集整理表
    (17, 13), # 2: GD-C1-342 质量证明文件汇总核查表
    (17, 6),  # 3: GD-C1-344 开箱检查验收记录
    (15, 8),  # 4: 附表1 材料进场数量清单
    (10, 4),  # 5: 附表 设备清单
]

PARAGRAPH_TITLES = [
    "工程材料、 构配件、 设备报审表",
    "GD-C1-347",
    "施工物资产品合格证收集整理表",
    "GD-C1-341",
    "施工物资产品质量证明文件汇总核查表",
    "GD-C1-342",
    "重要施工物资进场（开箱）检查验收记录",
    "（建筑设备工程）",
    "GD-C1-344",
    "附表1：材料进场数量清单",
    "项目名称：",
    "附表2：进场验收照片",
    "附表3：出厂质量证明文件",
]

HEADER_ROWS = {
    4: ["序号", "材料/设备名称", "厂家/品牌", "规格型号", "单位", "数量", "用途", "进场日期"],
    5: ["项目名称", "", "", ""],
}


def add_table(doc: Document, rows: int, cols: int, table_index: int) -> None:
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    if table_index in HEADER_ROWS:
        headers = HEADER_ROWS[table_index]
        for col, text in enumerate(headers):
            if col < cols:
                table.rows[0].cells[col].text = text
        if table_index == 5 and rows > 2:
            sub = ["设备名称", "厂家品牌", "单位", "数量"]
            for col, text in enumerate(sub):
                if col < cols:
                    table.rows[2].cells[col].text = text
    if table_index == 0 and rows >= 1:
        table.rows[0].cells[0].text = "单位(子单位)工程名称"
    if table_index == 1 and rows >= 1:
        table.rows[0].cells[0].text = "单位（子单位）工程名称"
    if table_index == 2 and rows >= 3:
        table.rows[2].cells[0].text = "序 号"
    if table_index == 3 and rows >= 1:
        table.rows[0].cells[0].text = "单位（子单位）工程名称"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)

    title_idx = 0
    for i, (rows, cols) in enumerate(TABLE_SPECS):
        while title_idx < len(PARAGRAPH_TITLES):
            doc.add_paragraph(PARAGRAPH_TITLES[title_idx])
            title_idx += 1
            if i == 0 and title_idx >= 2:
                break
            if i == 1 and title_idx >= 4:
                break
            if i == 2 and title_idx >= 6:
                break
            if i == 3 and title_idx >= 9:
                break
            if i == 4 and title_idx >= 11:
                break
        add_table(doc, rows, cols, i)

    while title_idx < len(PARAGRAPH_TITLES):
        doc.add_paragraph(PARAGRAPH_TITLES[title_idx])
        title_idx += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
