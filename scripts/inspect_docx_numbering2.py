# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from lxml import etree
import zipfile

p = Path(r"C:\Users\13065\Desktop\机场平台2期\《智慧工地建设管控一体化平台需求规格说明书》_updated.docx")
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

with zipfile.ZipFile(p) as z:
    num_xml = z.read("word/numbering.xml")
root = etree.fromstring(num_xml)
for num in root.findall("w:num", NS):
    nid = num.get(qn("w:numId"))
    abs_ref = num.find("w:abstractNumId", NS)
    abs_id = abs_ref.get(qn("w:val")) if abs_ref is not None else "?"
    print(f"numId={nid} -> abstractNumId={abs_id}")

doc = Document(str(p))
for i, para in enumerate(doc.paragraphs):
    st = para.style.name if para.style else ""
    if st != "Heading 1":
        continue
    print("=== sample Heading 1 paragraph runs ===")
    for r in para.runs:
        print(repr(r.text), r.font.name, r.font.size, r.font.bold, r._element.xml[:300] if r._element is not None else "")
    break

for lvl in (1, 2, 3):
    for para in doc.paragraphs:
        if (para.style.name if para.style else "") == f"Heading {lvl}":
            print(f"--- Heading {lvl} first run ---")
            if para.runs:
                r = para.runs[0]
                print("text", repr(r.text), "size", r.font.size, "name", r.font.name)
            break

# styles.xml heading rPr sizes
with zipfile.ZipFile(p) as z:
    styles = etree.fromstring(z.read("word/styles.xml"))
for sid in ("1", "2", "3", "4"):
    for st in styles.findall("w:style", NS):
        if st.get(qn("w:styleId")) == f"Heading{sid}" or st.get(qn("w:styleId")) == f"{sid}":
            name_el = st.find("w:name", NS)
            rpr = st.find("w:rPr", NS)
            print("style", st.get(qn("w:styleId")), name_el.get(qn("w:val")) if name_el is not None else "")
            if rpr is not None:
                sz = rpr.find("w:sz", NS)
                rfonts = rpr.find("w:rFonts", NS)
                print("  sz", sz.get(qn("w:val")) if sz is not None else None)
                if rfonts is not None:
                    print("  fonts", {k.split('}')[-1]: v for k,v in rfonts.attrib.items()})
