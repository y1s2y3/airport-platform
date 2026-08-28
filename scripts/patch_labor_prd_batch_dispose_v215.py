# -*- coding: utf-8 -*-
"""V2.1.5：批量消除→批量处置（关单）；状态统一为待处理/已关闭/未读/已读。"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

DOC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)

# 全文精确替换（段落 + 表格单元格）
EXACT_PAIRS = [
    (
        "规则触发后投递至个人中心「预警中心」；展示处置任务与通知两类条目及详情处置记录。"
        "Web 端处置任务类仅默认接收人可处置并关闭；APP 仅查看。"
        "「批量消除」仅从预警中心列表软删，不调用业务侧处置并关闭；"
        "「批量已读」仅对通知类「未读」生效。",
        "规则触发后投递至个人中心「预警中心」；展示处置任务与通知两类条目及详情处置记录。"
        "Web 端处置任务类仅默认接收人可处置并关闭；APP 仅查看。"
        "「批量处置预警」：勾选后弹窗填报处置说明与附件（与详情处置一致），"
        "仅状态为「待处理」的处置任务关闭为「已关闭」（同步业务预警与预警中心）；"
        "「批量已读」仅对通知类「未读」生效。",
    ),
    (
        "业务预警状态仍与 F-LABOR-06 一致（待处理/已关闭/未读/已读）。"
        "预警中心列表映射：处置任务类→类型「处置任务」、状态「待处置/已关闭」；"
        "通知类→类型「通知」、状态「未读/已读」。",
        "预警清单与个人中心「预警中心」状态口径一致：待处理 / 已关闭 / 未读 / 已读。"
        "处置任务类→类型「处置任务」、状态「待处理/已关闭」；"
        "通知类→类型「通知」、状态「未读/已读」（不再使用「已通知」「待处置」「已处置」）。",
    ),
    (
        "批量已读：仅对状态为「未读」的通知条目生效；处置任务「待处置」不受影响。",
        "批量已读：仅对状态为「未读」的通知条目生效；处置任务「待处理」不受影响。",
    ),
    (
        "批量消除：从预警中心列表软删（消息消除），不关闭业务预警、不写处置关闭时间线。",
        "批量处置预警：弹窗填报处置说明、附件（与详情处置一致）；"
        "仅「待处理」的处置任务批量关闭为「已关闭」（同步业务预警状态与处置时间线）；"
        "通知类不走批量处置。",
    ),
]

# 单元格内局部替换（谨慎）
CELL_PAIRS = [
    ("批量已读/消除", "批量已读/处置"),
    ("消除≠业务关闭", "批量处置=业务关闭为已关闭"),
    ("批量已读、批量消除", "批量已读、批量处置预警"),
    ("消除后列表移除且业务预警未关闭", "待处理处置任务变为已关闭；未读通知变已读"),
    ("批量已读与消除", "批量已读与处置"),
    ("列表-状态|处置任务：待处置/已处置；通知：未读/已读", None),  # handled below
]

STATUS_CELL_EXACT = {
    "列表-状态": "处置任务：待处理/已关闭；通知：未读/已读",
}

# 通知类终态：状态枚举中的「已通知」→「未读/已读」（仅状态语境）
STATUS_ENUM_PAIRS = [
    ("待处理、已关闭、已通知", "待处理、已关闭、未读、已读"),
    ("待处理/已关闭/已通知", "待处理/已关闭/未读/已读"),
    ("任务类状态「待处理→已关闭」，通知类终态「已通知」", "处置任务「待处理→已关闭」，通知类「未读→已读」"),
    ("状态为已通知，无需关闭", "状态为未读/已读，无需关闭"),
    ("通知类状态为已通知、无需关闭", "通知类状态为未读/已读、无需关闭"),
    ("通知类状态为已通知", "通知类状态为未读/已读"),
    ("终态「已通知」", "状态「未读/已读」"),
    ("状态为「已通知」", "状态为「未读」或「已读」"),
    ("不含通知类「已通知」", "不含通知类（未读/已读）"),
    ("通知类/已通知", "通知类未读/已读"),
    ("（通知类不计）", "（通知类未读/已读不计）"),
]


def set_para_text(p, value: str):
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str):
    if not cell.paragraphs:
        cell.text = value
        return
    set_para_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    set_para_text(p, new)
    return 1


def format_table(table):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def update_meta_revision(doc):
    meta = doc.tables[0]
    for row in meta.rows:
        k = row.cells[0].text.strip()
        if k == "版本号":
            set_cell_text(row.cells[1], "V2.1.5")
        if k == "更新日期":
            set_cell_text(row.cells[1], "2026-08-15")

    rev = doc.tables[1]
    note = (
        "预警中心「批量消除」改为「批量处置预警」（弹窗填说明+附件，仅待处理处置任务关闭为已关闭）；"
        "预警清单与预警中心状态统一为待处理/已关闭/未读/已读（取消待处置/已处置及通知类「已通知」状态名）。"
    )
    for row in rev.rows[1:]:
        if row.cells[0].text.strip() == "V2.1.5":
            set_cell_text(row.cells[3], note)
            set_cell_text(row.cells[2], "2026-08-15")
            return
    rev._tbl.append(deepcopy(rev._tbl.tr_lst[-1]))
    row = rev.rows[-1]
    for i, v in enumerate(["V2.1.5", "—", "2026-08-15", note]):
        set_cell_text(row.cells[i], v)
    format_table(rev)


def patch_tables(doc):
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows:
            continue
        head = rows[0]

        # 交互说明表 · 批量处置行
        if head[:4] == ["操作", "触发方式", "系统响应", "异常处理"]:
            for i, r in enumerate(rows[1:], start=1):
                if "批量处置" in r[0] or "批量消除" in r[0]:
                    set_cell_text(table.rows[i].cells[0], "批量处置预警")
                    set_cell_text(table.rows[i].cells[1], "勾选后点「批量处置预警」")
                    set_cell_text(
                        table.rows[i].cells[2],
                        "弹窗填报处置说明、附件（与详情处置一致）；仅「待处理」的处置任务改为「已关闭」（同步业务预警）",
                    )
                    set_cell_text(table.rows[i].cells[3], "未勾选或无可处置待处理任务则提示")
            format_table(table)

        # 字段表 · 列表-状态
        if head[:2] == ["字段", "说明"] or (len(head) >= 2 and head[0] == "字段"):
            for i, r in enumerate(rows[1:], start=1):
                if r[0] == "列表-状态":
                    set_cell_text(
                        table.rows[i].cells[1],
                        "处置任务：待处理/已关闭；通知：未读/已读",
                    )

        # 验收 AC
        if head[:5] == ["编号", "验收项", "Given", "When", "Then"]:
            blob = "\n".join("|".join(r) for r in rows)
            if "AC-01" in blob and "预警中心" in blob:
                data = [
                    [
                        "AC-01",
                        "Web、APP 预警中心可达",
                        "用户为预警责任人/接收人",
                        "打开个人中心 → 预警中心",
                        "可见对应处置任务/通知并可打开详情；APP 仅查看不可处置",
                    ],
                    [
                        "AC-02",
                        "处置任务手动关闭",
                        "Web 预警中心进入详情且当前用户为默认接收人",
                        "提交处置并关闭",
                        "业务预警与预警中心对应条状态均为「已关闭」；APP 无处置入口",
                    ],
                    [
                        "AC-03",
                        "自动关闭类引导",
                        "打开预警中心中自动关闭类处置任务详情",
                        "Web 按提示前往实名制/闸机侧处理，次日定时关闭；APP 仅查看",
                        "条件满足后业务与预警中心状态均为「已关闭」；APP 不操作",
                    ],
                    [
                        "AC-04",
                        "不双挂待办/通知",
                        "存在人员预警",
                        "查看我的待办 / 通知信息（及 APP 流程中心/消息提醒）",
                        "不再出现同条人员预警；仅预警中心展示",
                    ],
                    [
                        "AC-05",
                        "批量已读与处置",
                        "预警中心存在未读通知与待处理处置任务",
                        "批量已读、批量处置预警（填说明+附件）",
                        "未读通知变已读；待处理处置任务变已关闭（业务预警同步关闭）",
                    ],
                    [
                        "AC-06",
                        "APP 无业务菜单",
                        "APP 登录",
                        "浏览业务菜单",
                        "无实名制看板/台账/预警清单等",
                    ],
                ]
                while len(table.rows) < 1 + len(data):
                    table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
                for i, row_data in enumerate(data):
                    for j, val in enumerate(row_data):
                        set_cell_text(table.rows[i + 1].cells[j], val)
                for i in range(1 + len(data), len(table.rows)):
                    for j in range(len(head)):
                        set_cell_text(table.rows[i].cells[j], "")
                format_table(table)

        # 原型页表
        if head[:4] == ["页面名称", "页面路径", "核心内容", "备注"]:
            for i, r in enumerate(rows[1:], start=1):
                if "预警中心列表" in r[0]:
                    set_cell_text(
                        table.rows[i].cells[2],
                        "处置任务/通知列表；筛选、分页、批量已读/处置",
                    )

        # 功能清单 F-LABOR-10
        if head and head[0] == "功能ID":
            for i, r in enumerate(rows[1:], start=1):
                if r[0] == "F-LABOR-10":
                    set_cell_text(
                        table.rows[i].cells[3],
                        "预警中心汇总处置任务/通知；筛选分页；批量已读/处置；"
                        "Web 处置任务类仅默认接收人可处置，APP 仅查看",
                    )
                    if len(r) > 6:
                        set_cell_text(
                            table.rows[i].cells[6],
                            "不进入待办/通知；批量处置将待处理关闭为已关闭",
                        )


def patch_status_enum_everywhere(doc):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, STATUS_ENUM_PAIRS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, STATUS_ENUM_PAIRS)
                # 状态清单表常见行
                t0 = row.cells[0].text.strip() if row.cells else ""
                if t0.startswith("已通知"):
                    # 改状态名行为未读/已读说明
                    blob = "|".join(c.text.strip() for c in row.cells)
                    if "通知类" in blob or "终态" in blob or "pending" in blob.lower() or "notify" in blob.lower() or "已阅" in blob or "无需关闭" in blob:
                        set_cell_text(row.cells[0], "未读 / 已读（notify）")
                        if len(row.cells) > 1 and "已通知" in row.cells[1].text:
                            set_cell_text(
                                row.cells[1],
                                row.cells[1].text.replace("已通知", "未读或已读"),
                            )
    return n


def patch_6_6_notify_paragraphs(doc):
    pairs = [
        (
            "自动/手动关闭后变为已关闭；通知类状态为已通知、无需关闭、不参与分级上报。",
            "自动/手动关闭后变为已关闭；通知类状态为未读/已读、无需关闭、不参与分级上报。",
        ),
        (
            "通知类触发后即为已通知、仅留触发记录、不升级不关闭",
            "通知类触发后为未读（查阅后为已读）、仅留触发记录、不升级不关闭",
        ),
        (
            "通知类（高龄提醒、身份证过期提醒）：状态为已通知，无需关闭，不参与分级上报与超期升级；仍投递个人中心预警中心；详情只读、无处置并关闭；时间线仅保留触发记录。同一人同一规则在同一项目仅生成一条，已通知后即使条件仍成立也不重生。待处理/待人工处置指标不含通知类；处置方式筛选含「通知」。",
            "通知类（高龄提醒、身份证过期提醒）：状态为未读/已读，无需关闭，不参与分级上报与超期升级；仍投递个人中心预警中心；详情只读、无处置并关闭；时间线仅保留触发记录。同一人同一规则在同一项目仅生成一条，已读后即使条件仍成立也不重生。待处理/待人工处置指标不含通知类；处置方式筛选含「通知」。",
        ),
        (
            "已通知后即使条件仍成立也不重生",
            "已读后即使条件仍成立也不重生",
        ),
        (
            "通知类：同一人同一规则同一项目仅一条，已通知后不重生",
            "通知类：同一人同一规则同一项目仅一条，生成后不因重复条件重生（已读后亦不重生）",
        ),
    ]
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, pairs)
    return n


def apply_exact_and_cell(doc):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, EXACT_PAIRS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, EXACT_PAIRS)
                    # skip historical V2.1.2 revision wording intentionally? still replace 消除 in feature tables
                    text = p.text
                    new = text
                    for a, b in [
                        ("批量已读/消除", "批量已读/处置"),
                        ("消除≠业务关闭", "批量处置=业务关闭为已关闭"),
                        ("批量已读、批量消除", "批量已读、批量处置预警"),
                        ("消除后列表移除且业务预警未关闭", "待处理处置任务变为已关闭；未读通知变已读"),
                        ("批量已读与消除", "批量已读与处置"),
                        ("预警中心对应条状态变为「已处置」", "预警中心对应条状态变为「已关闭」"),
                        ("预警中心条变为已处置", "预警中心条变为已关闭"),
                        ("处置任务：待处置/已处置；通知：未读/已读", "处置任务：待处理/已关闭；通知：未读/已读"),
                        ("仅「待处置」任务改为已关闭", "仅「待处理」的处置任务改为已关闭"),
                        ("仅「待处置」任务改为已处置", "仅「待处理」的处置任务改为已关闭"),
                    ]:
                        if a in new:
                            new = new.replace(a, b)
                    if new != text:
                        set_para_text(p, new)
                        n += 1
    return n


def set_heading_space_after(doc):
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def main():
    # stop if word open is caller's responsibility
    doc = Document(str(DOC))
    update_meta_revision(doc)
    n1 = apply_exact_and_cell(doc)
    patch_tables(doc)
    n2 = patch_status_enum_everywhere(doc)
    n3 = patch_6_6_notify_paragraphs(doc)
    set_heading_space_after(doc)
    doc.save(str(DOC))
    print("saved V2.1.5", "exact/cell", n1, "enum", n2, "6.6", n3)


if __name__ == "__main__":
    main()
