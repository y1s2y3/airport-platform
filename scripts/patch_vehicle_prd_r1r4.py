# -*- coding: utf-8 -*-
"""优化车辆 PRD：R1 external_id AC、R2 车牌筛选 AC、R3 AC-06 文案、R4 车牌规则去「幂等」。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx")
LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_vehicle_r1r4_log.txt"
)
VER = "V2.2.5"
TODAY = "2026-08-13"
NOTE = "复审优化：补无 external_id 拒收验收；进出场筛选 AC 含车牌；修正看板空数据 Then 文案；车牌归一化不再关联幂等表述。"


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def replace_para_text(para, old: str, new: str) -> bool:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    if old not in full:
        return False
    full2 = full.replace(old, new)
    if para.runs:
        para.runs[0].text = full2
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = full2
    return True


def bump_meta(doc: Document) -> None:
    for row in doc.tables[0].rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], VER)
        if key == "更新日期":
            set_cell_text(row.cells[1], TODAY)


def main() -> None:
    log: list[str] = []
    bak = SRC.with_name(f"车辆管理_PRD_复审优化前_副本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    shutil.copy2(SRC, bak)
    log.append("backup " + bak.name)

    doc = Document(str(SRC))

    # R4: plate rule wording
    old_plate = "车牌比对/去重前归一化：去除首尾空白与中间空格后转大写，再参与在场判定与幂等相关比对"
    new_plate = "车牌比对/去重前归一化：去除首尾空白与中间空格后转大写，再参与在场判定等车牌去重比对（幂等仍以 project_id+external_id 为准）"
    n = 0
    for para in doc.paragraphs:
        if replace_para_text(para, old_plate, new_plate):
            n += 1
        # shorter variant if already partially edited
        if replace_para_text(
            para,
            "再参与在场判定与幂等相关比对",
            "再参与在场判定等车牌去重比对（幂等仍以 project_id+external_id 为准）",
        ):
            n += 1
    log.append(f"R4 plate wording x{n}")

    bump_meta(doc)
    doc.save(str(SRC))
    log.append("phase1 saved")

    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        wdoc = word.Documents.Open(str(SRC), ReadOnly=False)

        # revision
        tbl = wdoc.Tables(2)
        first = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first != VER:
            tbl.Rows(2).Select()
            word.Selection.InsertRowsAbove(1)
            row = tbl.Rows(2)
            ncols = row.Cells.Count
            vals = [VER, TODAY, NOTE] if ncols == 3 else [VER, "—", TODAY, NOTE]
            for i, v in enumerate(vals, start=1):
                if i <= ncols:
                    row.Cells(i).Range.Text = v
            log.append("rev " + VER)

        def find_table(needle: str) -> int | None:
            for i in range(1, wdoc.Tables.Count + 1):
                if needle in wdoc.Tables(i).Range.Text:
                    return i
            return None

        def cell_text(table_i: int, r: int, c: int) -> str:
            return (
                wdoc.Tables(table_i)
                .Cell(r, c)
                .Range.Text.strip()
                .replace("\r", "")
                .replace("\x07", "")
            )

        def set_cells(table_i: int, r: int, values: list[str]) -> None:
            row = wdoc.Tables(table_i).Rows(r)
            for c, val in enumerate(values, start=1):
                if c <= row.Cells.Count:
                    row.Cells(c).Range.Text = val

        def append_row(table_i: int, values: list[str], skip_if_id: str) -> None:
            tblx = wdoc.Tables(table_i)
            for r in range(2, tblx.Rows.Count + 1):
                if cell_text(table_i, r, 1) == skip_if_id:
                    log.append(f"skip existing {skip_if_id}")
                    return
            tblx.Rows(tblx.Rows.Count).Select()
            word.Selection.InsertRowsBelow(1)
            set_cells(table_i, tblx.Rows.Count, values)
            log.append(f"add {skip_if_id}")

        # R3: 6.1 AC-06 Then
        i61 = find_table("三项统计指标")
        if i61:
            for r in range(2, wdoc.Tables(i61).Rows.Count + 1):
                if cell_text(i61, r, 1) == "AC-06":
                    set_cells(
                        i61,
                        r,
                        [
                            "AC-06",
                            "无数据为0",
                            "统计日无任何进出场上报",
                            "打开看板",
                            "三项指标均显示 0，不报错",
                        ],
                    )
                    log.append("R3 AC-06 Then fixed")
                    break

        # R2: 6.4 AC-03
        i64 = find_table("无手动填报")
        if i64 is None:
            i64 = find_table("按方向筛选")
        if i64:
            for r in range(2, wdoc.Tables(i64).Rows.Count + 1):
                if cell_text(i64, r, 1) == "AC-03":
                    set_cells(
                        i64,
                        r,
                        [
                            "AC-03",
                            "车牌/方向筛选",
                            "存在不同车牌及进场与出场记录",
                            "分别按车牌、按方向筛选",
                            "列表结果符合筛选条件",
                        ],
                    )
                    log.append("R2 AC-03 plate+direction")
                    break

        # R1: 6.6 expand AC-02 and/or add AC-05
        i66 = find_table("project_id+external_id")
        if i66 is None:
            i66 = find_table("缺车牌或项目 ID")
        if i66:
            for r in range(2, wdoc.Tables(i66).Rows.Count + 1):
                if cell_text(i66, r, 1) == "AC-02":
                    set_cells(
                        i66,
                        r,
                        [
                            "AC-02",
                            "拒非法",
                            "缺车牌、或项目 ID 无效、或缺 external_id",
                            "调用上报",
                            "拒绝写入；列表无脏数据",
                        ],
                    )
                    log.append("R1 AC-02 includes external_id")
                    break
            append_row(
                i66,
                [
                    "AC-05",
                    "无 external_id 拒收",
                    "其他必填齐全但未传 external_id",
                    "调用上报",
                    "拒绝该条；不落库",
                ],
                "AC-05",
            )

        tmp = SRC.with_name("_vehicle_r1r4_tmp.docx")
        if tmp.exists():
            tmp.unlink()
        wdoc.SaveAs2(str(tmp), FileFormat=12)
        wdoc.Close(False)
        shutil.copy2(tmp, SRC)
        tmp.unlink()

        wdoc = word.Documents.Open(str(SRC), ReadOnly=True)
        log.append(f"WORD_OK {wdoc.Paragraphs.Count}")
        wdoc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    # verify
    d2 = Document(str(SRC))
    blob = "\n".join(p.text for p in d2.paragraphs)
    for t in d2.tables:
        for row in t.rows:
            blob += "\n" + " | ".join(c.text for c in row.cells)
    for needle in [
        "三项指标均显示 0，不报错",
        "车牌/方向筛选",
        "缺 external_id",
        "无 external_id 拒收",
        "车牌去重比对",
        "与幂等相关比对",
    ]:
        log.append(f"has[{needle}]={needle in blob}")

    LOG.write_text("\n".join(log), encoding="utf-8")
    print("done")


if __name__ == "__main__":
    main()
