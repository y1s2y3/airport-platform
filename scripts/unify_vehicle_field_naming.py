# -*- coding: utf-8 -*-
"""统一车辆管理 PRD 字段命名为 snake_case（与人员实名制 / 车辆架构一致）。"""
import re
import shutil
from pathlib import Path

from docx import Document

DIR = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
PRD = DIR / "车辆管理_PRD.docx"
BAK = DIR / "车辆管理_PRD_字段命名统一前_副本_20260813.docx"
ARCH = DIR / "product-architecture-车辆进出场管理-v1.md"

# 长键优先，避免部分替换
RENAME = [
    ("todayInCount", "today_in_count"),
    ("todayOutCount", "today_out_count"),
    ("onSiteCount", "on_site_count"),
    ("projectName", "project_name"),
    ("projectId", "project_id"),
    ("projectCode", "project_code"),
    ("systemName", "system_name"),
    ("updatedAt", "updated_at"),
    ("plateNo", "plate_no"),
    ("vehicleType", "vehicle_type"),
    ("gateName", "gate_name"),
    ("recordTime", "record_time"),
    ("externalId", "external_id"),
]


def replace_text(s: str) -> str:
    if not s:
        return s
    out = s
    for old, new in RENAME:
        out = out.replace(old, new)
    return out


def patch_paragraph(p):
    text = p.text
    if not text:
        return 0
    new = replace_text(text)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def patch_docx(path: Path) -> int:
    doc = Document(str(path))
    n = 0
    for p in doc.paragraphs:
        n += patch_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += patch_paragraph(p)
    doc.save(str(path))
    return n


def verify_prd(path: Path):
    doc = Document(str(path))
    fields = []
    for table in doc.tables:
        if not table.rows:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if "字段名" not in h:
            continue
        i = h.index("字段名")
        for row in table.rows[1:]:
            name = row.cells[i].text.strip()
            if name and name not in ("操作", "—", "-"):
                fields.append(name)
    camel = [f for f in fields if re.search(r"[a-z][A-Z]", f)]
    return fields, camel


def patch_arch(path: Path) -> int:
    text = path.read_text(encoding="utf-8")
    orig = text
    # architecture already snake_case; only fix any camel leftovers + revision note
    for old, new in RENAME:
        text = text.replace(f"`{old}`", f"`{new}`")
        text = text.replace(old, new)  # rare bare
    # avoid double-mangling if already snake - RENAME olds are camel only so OK

    if "V1.30 修订" in text and "字段命名" not in text.split("V1.30 修订", 1)[1][:200]:
        # bump to V1.31 for naming unify
        pass

    # Version bump
    if "| 版本号 | V1.30 |" in text:
        text = text.replace("| 版本号 | V1.30 |", "| 版本号 | V1.31 |", 1)
        text = text.replace("| 更新日期 | 2026-08-13 |", "| 更新日期 | 2026-08-13 |", 1)
        # insert revision row after version block line of V1.30
        marker = "| V1.30 修订 |"
        if marker in text and "| V1.31 修订 |" not in text:
            insert = (
                "| V1.31 修订 | 字段命名与人员实名制及本架构统一为 **snake_case**"
                "（如 `project_id`、`plate_no`、`today_in_count`）；车辆 PRD 同步。 |\n"
            )
            # place after the V1.30 revision line
            idx = text.find(marker)
            end = text.find("\n", idx)
            text = text[: end + 1] + insert + text[end + 1 :]
    elif "| 版本号 | V1.29 |" in text:
        text = text.replace("| 版本号 | V1.29 |", "| 版本号 | V1.31 |", 1)

    if text != orig:
        path.write_text(text, encoding="utf-8")
        return 1
    return 0


def main():
    shutil.copy2(PRD, BAK)
    print("backup:", BAK.name)

    hits = patch_docx(PRD)
    print("prd replace hits:", hits)

    fields, camel = verify_prd(PRD)
    print("prd fields:", sorted(set(fields)))
    print("prd remaining camel:", camel)

    arch_hits = patch_arch(ARCH)
    print("arch patched:", arch_hits)

    # verify arch has no camel field identifiers in backticks
    arch = ARCH.read_text(encoding="utf-8")
    cams = sorted(set(re.findall(r"`([a-z]+[A-Z][a-zA-Z0-9]*)`", arch)))
    print("arch remaining camel backticks:", cams)

    # sample labor style reminder
    labor = DIR / "人员实名制_PRD_待评审.docx"
    if labor.exists():
        lf, lc = verify_prd(labor)
        print("labor camel leftover:", lc[:10], "snake samples:", [x for x in lf if "_" in x][:8])


if __name__ == "__main__":
    main()
