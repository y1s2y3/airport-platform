# -*- coding: utf-8 -*-
"""V1.3 patch on 副本_20260812.docx"""
from docx import Document
from pathlib import Path

PATH = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_副本_20260812.docx"
)
ALSO = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_V1.3.docx"
)


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def set_cell_text(cell, value):
    if not cell.paragraphs:
        cell.text = value
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def apply_all(doc, pairs):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, pairs)
    return n


def main():
    doc = Document(str(PATH))

    pairs = [
        # --- dispose: only default recipient everywhere ---
        (
            "项目 Web 可处置角色，或 Web 个人中心默认接收人；APP 不可处置",
            "仅默认接收人（Web 个人中心或项目 Web 预警清单）；APP 不可处置",
        ),
        (
            "项目侧劳务相关角色\n指挥部层级分级推送用户",
            "项目侧授权用户可查看；处置仅默认接收人\n指挥部层级分级推送用户可查看通知",
        ),
        (
            "与预警清单与处置（功能ID：F-LABOR-06），功能相同",
            "Web 端且当前用户为默认接收人时，处置能力与 F-LABOR-06 预警清单对齐；APP/非默认接收人仅只读查看",
        ),
        # --- default recipient: separately configured ---
        (
            "指挥部按项目维护三区块：人员轨迹外链、分级管控、预警规则（9 条）。保存后影响指挥部轨迹系统列表入列、项目侧栏跳转与预警行为。",
            "指挥部按项目维护：人员轨迹外链、默认接收人、分级管控、预警规则（9 条）。默认接收人在配置中单独指定（不等于分级第1级，除非人工选成同一人）。保存后影响轨迹列表入列、侧栏跳转与预警处置权限。",
        ),
        (
            "可读写人员轨迹跳转配置、分级上报配置与预警规则配置；某项目「系统地址」非空时，即派生进入「人员轨迹系统」（F-LABOR-02）列表。",
            "可读写人员轨迹跳转配置、默认接收人、分级上报配置与预警规则配置；某项目「系统地址」非空时，即派生进入「人员轨迹系统」（F-LABOR-02）列表。",
        ),
        (
            "轨迹 + 分级 + 规则",
            "轨迹 + 默认接收人 + 分级 + 规则",
        ),
        (
            "处置权限：仅配置的默认接收人可处置；其他分级管控用户仅接收通知、只读查看；APP 端一律只读",
            "处置权限：仅「实名制配置」中单独指定的默认接收人可处置（Web）；其他分级上报接收人仅通知/只读；APP 一律只读。默认接收人≠自动等于分级第1级。",
        ),
        (
            "升级上报/分级管控，系统仅做通知，详情仅可查看，仅默认接收人可处置",
            "升级上报/分级管控，系统仅做通知；详情对非默认接收人只读，仅配置的默认接收人可处置",
        ),
        # terminology tip in 成功指标 already says 仅默认接收人 - OK
        # --- education: 有数据即关 + 次日 ---
        (
            "定时任务每日检测，三级教育记录有数据后，系统自动关闭该预警。",
            "次日定时任务检测：只要存在三级安全教育记录（有数据即关，不校验是否合格/类型），系统自动关闭该预警。",
        ),
        (
            "编辑三级教育并保存\n存在对应自动关闭类预警且条件满足\n预警可被系统关闭",
            "编辑三级教育并保存（有记录即可）\n存在对应自动关闭类预警\n次日定时任务自动关闭该预警",
        ),
        (
            "自动关闭类且已补录合格三级教育\n系统检测\n自动关闭",
            "自动关闭类且已补录三级教育记录（有数据即可，不校验合格）\n次日定时任务检测\n自动关闭",
        ),
        (
            "定时任务每日检测，补传有效特种作业证书并维护有效期后，系统校验通过即自动关闭该预警。",
            "次日定时任务检测：ROMA/一期回写有效特种作业证书及有效期后，系统校验通过即自动关闭该预警（二期界面不可改特种证）。",
        ),
        (
            "定时任务每日检测，年龄信息正确，或人员退场后，自动关闭。",
            "次日定时任务检测：年龄信息正确（经 ROMA 回写）或人员退场后，自动关闭。",
        ),
        (
            "定时任务每日检测，人员移出黑名单，或人员退场后，自动关闭。",
            "次日定时任务检测：人员移出黑名单或人员退场（离场）后，自动关闭。",
        ),
        (
            "#姓名#未按要求录入三级教育,请前往人员实名制页面补录信息, 完成后，次日将自动关闭预警。。",
            "#姓名#未按要求录入三级教育，请前往人员实名制页面补录信息；有教育记录后，次日将自动关闭预警。",
        ),
        (
            "#姓名#特种作业证书缺失/过期,请联系闸机系统更新人员信息，完成后，次日将自动关闭预警。",
            "#姓名#特种作业证书缺失/过期，请联系闸机系统更新人员信息；信息回写后，次日将自动关闭预警。",
        ),
        (
            "#姓名#年龄低于16周岁，请联系闸机系统更新人员信息,或办理退场，完成后，次日将自动关闭预警。",
            "#姓名#年龄低于16周岁，请联系闸机系统更新人员信息或办理退场；条件满足后，次日将自动关闭预警。",
        ),
        (
            "#姓名#，身份证号:#身份证号（脱敏）#，属于指挥部规定黑名单人员，请联系闸机系统办理人员退场，或联系项目经理确认是否错误录入。完成后，次日将自动关闭预警。",
            "#姓名#，身份证号:#身份证号（脱敏）#，属于指挥部规定黑名单人员，请联系闸机系统办理退场或确认是否误录入；引导处理到位后，次日将自动关闭预警。",
        ),
        # unify auto close note in 6.7 supplement
        (
            "补充：以上定时任务可在凌晨完成计算，每日9:00推送任务/通知，避免后续APP消息、短信提醒在凌晨打扰用户。",
            "补充：自动关闭类预警统一为「次日定时任务」关闭（凌晨计算）。每日9:00推送任务/通知，避免凌晨打扰。短信提醒本期不做。",
        ),
        # --- blacklist 6.8.1 ---
        (
            "指挥部维护全平台共用（不按项目拆分）劳务黑名单；每日定时扫描各项目在岗人员，命中黑名单且规则开启时生成「黑名单人员进场预警」；二期不做闸机强制拦截，仅预警并由责任人处置。",
            "指挥部维护全平台共用（不按项目拆分）劳务黑名单；每日定时扫描各项目在岗人员，命中黑名单且规则开启时生成「黑名单人员进场预警」；二期不做闸机强制拦截。预警为系统自动关闭类：默认接收人可在 Web 查看并引导处理，处理到位（退场/移出黑名单等）后次日自动关闭。",
        ),
        (
            "黑名单日扫在岗命中生成「黑名单人员进场预警」（系统自动关闭类）；默认接收人可在 Web 查看/引导处理",
            "黑名单日扫在岗命中生成「黑名单人员进场预警」（系统自动关闭类）；默认接收人可在 Web 查看并引导处理，处理到位后次日自动关闭",
        ),
        # --- SMS out of scope ---
        (
            "车辆预警；多地同时在场预警；工资管理；按班组考勤统计",
            "车辆预警；多地同时在场预警；工资管理；按班组考勤统计；短信提醒（本期不做）",
        ),
    ]

    print("hits:", apply_all(doc, pairs))

    # Fix AC cells that may be split across cells (education)
    for table in doc.tables:
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells]
            if not vals:
                continue
            # AC education
            if vals[0] == "AC-03" and "教育" in vals[1]:
                if len(row.cells) >= 5:
                    set_cell_text(row.cells[2], "编辑三级教育并保存（有记录即可，不校验合格）")
                    set_cell_text(row.cells[3], "存在对应自动关闭类预警")
                    set_cell_text(row.cells[4], "次日定时任务自动关闭该预警")
                    print("fixed education AC-03")
            if vals[0] == "AC-02" and "自动关闭" in vals[1]:
                if len(row.cells) >= 5:
                    set_cell_text(
                        row.cells[2],
                        "自动关闭类且已补录三级教育记录（有数据即可）",
                    )
                    set_cell_text(row.cells[3], "次日定时任务检测")
                    set_cell_text(row.cells[4], "自动关闭")
                    print("fixed warning AC-02 auto close")

    # Insert 默认接收人 field description near 分级管控 if possible:
    # Find table with reportDays / level headers and prepend a note paragraph... 
    # Better: add row to a simple config fields table, or patch 默认示例 paragraph.
    for p in doc.paragraphs:
        if p.text.strip().startswith("默认示例：安全员"):
            # expand paragraph to include default recipient rule
            old = p.text
            add = (
                "默认接收人：在实名制配置中按项目单独指定 1 人（必填）；"
                "仅该账号可在 Web 处置任务类预警；与分级管控各级责任人分开配置。"
            )
            if "默认接收人：在实名制配置" not in old:
                new = add + "\n" + old
                if p.runs:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(new)
                print("added default recipient note near 分级示例")
            break

    # Add field rows into 分级管控 table? Find table with level/recipientId
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        # look for field tables containing projectId enabled for track - skip
        # Find table that has first data row level
        texts = ["|".join(c.text.strip() for c in r.cells) for r in table.rows[:3]]
        if any(t.startswith("level|") or "|level|" in f"|{t}|" for t in texts) or (
            len(table.rows) > 1 and table.rows[1].cells[0].text.strip() == "level"
        ):
            # check if defaultRecipient already exists
            alltxt = "\n".join(c.text for r in table.rows for c in r.cells)
            if "defaultRecipientId" in alltxt or "默认接收人" in alltxt and "defaultRecipient" in alltxt:
                break
            # Insert after header a new row by cloning last row pattern - add_row at end then we can't easily reorder.
            # Instead append row documenting default recipient as separate entity note via add_row at top after header is hard.
            # Add at end of this table:
            row = table.add_row()
            vals = [
                "defaultRecipientId",
                "默认接收人",
                "ref",
                "是",
                "按项目在实名制配置中单独指定；仅此人可 Web 处置任务类预警；与分级各级责任人分开配置，不自动等同第1级",
            ]
            for i, v in enumerate(vals):
                if i < len(row.cells):
                    set_cell_text(row.cells[i], v)
            print("added defaultRecipientId row in table", ti)
            break

    # Meta version
    for table in doc.tables[:6]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            if cells[0] == "版本号":
                set_cell_text(row.cells[1], "V1.3")
            if cells[0] == "更新日期":
                set_cell_text(row.cells[1], "2026-08-12")
            if cells[0] == "状态":
                set_cell_text(row.cells[1], "待评审")

    for table in doc.tables[:10]:
        headers = [c.text.strip() for c in table.rows[0].cells]
        if not headers or headers[0] != "版本":
            continue
        existing = [r.cells[0].text.strip() for r in table.rows]
        if "V1.3" not in existing:
            new_row = table.add_row()
            mapping = {
                "版本": "V1.3",
                "修订人": "",
                "日期": "2026-08-12",
                "说明": (
                    "评审修订：处置仅默认接收人（含项目预警清单）；默认接收人在配置中单独指定；"
                    "三级教育有数据即关；自动关闭类统一次日关闭；黑名单为引导处理后次日自动关闭；短信本期不做"
                ),
            }
            for i, h in enumerate(headers):
                if i < len(new_row.cells):
                    set_cell_text(new_row.cells[i], mapping.get(h, ""))
            print("revision V1.3 added")
        break

    for out in (PATH, ALSO):
        try:
            doc.save(str(out))
            print("saved", out)
        except PermissionError:
            print("LOCKED", out)


if __name__ == "__main__":
    main()
