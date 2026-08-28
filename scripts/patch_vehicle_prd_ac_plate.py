# -*- coding: utf-8 -*-
"""补 6.1/6.3 AC、车牌归一化、清单筛选文案（AC 行用 Word COM 插入，避免 OOXML 损坏）。"""
from __future__ import annotations

import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx")
LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_vehicle_ac_log.txt"
)
TODAY = "2026-08-13"
VER = "V2.2.4"
NOTE = "补齐看板/轨迹配置验收；车牌比对前 trim+去空格+大写；进出场清单筛选改为车牌、方向。"
PLATE_RULE = (
    "车牌比对/去重前归一化：去除首尾空白与中间空格后转大写，再参与在场判定与幂等相关比对"
)


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def replace_in_doc(doc: Document, old: str, new: str, log: list[str]) -> None:
    n = 0
    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
        if old not in full:
            continue
        n += full.count(old)
        full2 = full.replace(old, new)
        if para.runs:
            para.runs[0].text = full2
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = full2
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
                    if old not in full:
                        continue
                    n += full.count(old)
                    full2 = full.replace(old, new)
                    if para.runs:
                        para.runs[0].text = full2
                        for r in para.runs[1:]:
                            r.text = ""
                    else:
                        para.text = full2
    if n:
        log.append(f"replace {old!r}->{new!r} x{n}")


def add_bullet_under(doc: Document, heading_prefix: str, text: str, log: list[str]) -> None:
    if any(text[:16] in (p.text or "") for p in doc.paragraphs):
        log.append(f"skip existing: {text[:16]}")
        return
    target = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(heading_prefix):
            target = i
            break
    if target is None:
        log.append(f"WARN no heading {heading_prefix}")
        return
    last = None
    for j in range(target + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            break
        if p.text.strip():
            last = p
    if last is None:
        log.append(f"WARN no body under {heading_prefix}")
        return
    new_el = deepcopy(last._element)
    for node in new_el.iter():
        if node.tag.endswith("}t"):
            node.text = ""
    last._element.addnext(new_el)
    first = True
    for node in new_el.iter():
        if node.tag.endswith("}t"):
            if first:
                node.text = text
                first = False
            else:
                node.text = ""
    log.append(f"bullet under {heading_prefix}")


def bump_meta(doc: Document) -> None:
    for row in doc.tables[0].rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], VER)
        if key == "更新日期":
            set_cell_text(row.cells[1], TODAY)


def word_find_table_index(doc, needle: str) -> int | None:
    """1-based Word table index containing needle."""
    for i in range(1, doc.Tables.Count + 1):
        try:
            txt = doc.Tables(i).Range.Text
        except Exception:
            continue
        if needle in txt:
            return i
    return None


def word_append_ac_rows(doc, table_index: int, rows: list[list[str]], log: list[str], label: str) -> None:
    tbl = doc.Tables(table_index)
    existing = set()
    for r in range(2, tbl.Rows.Count + 1):
        try:
            v = tbl.Cell(r, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
            existing.add(v)
        except Exception:
            pass
    for vals in rows:
        if vals[0] in existing:
            log.append(f"{label} skip {vals[0]}")
            continue
        # insert below last row
        tbl.Rows(tbl.Rows.Count).Select()
        doc.Application.Selection.InsertRowsBelow(1)
        row = tbl.Rows(tbl.Rows.Count)
        for c, val in enumerate(vals, start=1):
            if c <= row.Cells.Count:
                row.Cells(c).Range.Text = val
        log.append(f"{label} add {vals[0]}")


def main() -> None:
    log: list[str] = []
    bak = SRC.with_name(f"车辆管理_PRD_AC补齐前_副本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    shutil.copy2(SRC, bak)
    log.append("backup " + bak.name)

    # Phase 1: python-docx text + bullets + meta
    doc = Document(str(SRC))
    replace_in_doc(doc, "关键词/方向筛选", "车牌、方向筛选", log)
    replace_in_doc(doc, "关键词、方向筛选", "车牌、方向筛选", log)
    add_bullet_under(doc, "6.1.9", PLATE_RULE, log)
    add_bullet_under(doc, "6.4.9", "车牌归一化：与 6.1.9 相同（trim、去空格、转大写后比对）", log)
    add_bullet_under(doc, "6.6.9", "车牌归一化：落库展示可保留原文；在场/去重比对按 6.1.9 规则", log)
    bump_meta(doc)
    doc.save(str(SRC))
    log.append("phase1 saved")

    # Phase 2: Word COM — AC rows + revision + sanitize
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        wdoc = word.Documents.Open(str(SRC), ReadOnly=False)

        # revision
        tbl = wdoc.Tables(2)
        first = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first != VER:
            tbl.Rows(2).Select()
            word.Selection.InsertRowsAbove(1)
            row = tbl.Rows(2)
            ncols = row.Cells.Count
            vals = [VER, TODAY, NOTE] if ncols == 3 else [VER, "—", TODAY, NOTE]
            for i, v in enumerate(vals, start=1):
                if i <= ncols:
                    row.Cells(i).Range.Text = v
            log.append("rev " + VER)

        i61 = word_find_table_index(wdoc, "三项统计指标")
        if i61 is None:
            i61 = word_find_table_index(wdoc, "菜单位置")
        log.append(f"word table 6.1={i61}")
        if i61:
            word_append_ac_rows(
                wdoc,
                i61,
                [
                    ["AC-03", "跨夜在场", "昨夜进场且今日尚未出场", "打开指挥部看板", "在场车辆计入该车牌（按项目内最新记录）"],
                    ["AC-04", "项目明细", "多项目均有上报流水", "查看指挥部看板项目明细", "各项目三项指标分别展示且可对上流水"],
                    ["AC-05", "进入进出场", "指挥部项目明细有数据", "点击行操作「查看项目详情」", "进入该项目进出场记录页"],
                    ["AC-06", "无数据为0", "统计日无任何进出场上报", "打开看板", "三项指标均显示 0（非空/报错）"],
                    ["AC-07", "无 APP", "APP 用户", "查找车辆管理看板菜单", "不可见"],
                    ["AC-08", "车牌归一化", "同项目上报「粤B 12345」与「粤B12345」且最新为进场", "查看在场车辆", "按归一化后同一车牌只计 1"],
                ],
                log,
                "6.1",
            )

        i63 = word_find_table_index(wdoc, "启用必填 URL")
        if i63 is None:
            i63 = word_find_table_index(wdoc, "启用必填")
        log.append(f"word table 6.3={i63}")
        if i63:
            word_append_ac_rows(
                wdoc,
                i63,
                [
                    ["AC-04", "停用保留 URL", "项目已填 URL 后改为停用并保存", "打开车辆轨迹系统列表", "仍出现该项目；状态为停用；仍可跳转"],
                    ["AC-05", "清空 URL 退出列表", "项目原有 URL，保存时清空 URL", "打开车辆轨迹系统列表", "该项目不再出现"],
                    ["AC-06", "停用后项目侧不可跳", "配置为停用但 URL 非空", "项目用户点击「车辆轨迹监管」", "不可外跳，提示至指挥部维护配置"],
                ],
                log,
                "6.3",
            )

        # sanitize save
        tmp = SRC.with_name("_vehicle_ac_tmp.docx")
        if tmp.exists():
            tmp.unlink()
        wdoc.SaveAs2(str(tmp), FileFormat=12)
        wdoc.Close(False)
        shutil.copy2(tmp, SRC)
        tmp.unlink()

        # verify open
        wdoc = word.Documents.Open(str(SRC), ReadOnly=True)
        log.append(f"WORD_OK {wdoc.Paragraphs.Count}")
        wdoc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    d2 = Document(str(SRC))
    checks = [
        "车牌、方向筛选",
        "关键词/方向",
        "车牌比对/去重前归一化",
        "AC-08",
        "停用保留 URL",
        "AC-06",
    ]
    for needle in checks:
        hit = any(needle in p.text for p in d2.paragraphs)
        if not hit:
            hit = any(needle in c.text for t in d2.tables for r in t.rows for c in r.cells)
        log.append(f"has[{needle}]={hit}")

    LOG.write_text("\n".join(log), encoding="utf-8")
    print("done", LOG)


if __name__ == "__main__":
    main()
