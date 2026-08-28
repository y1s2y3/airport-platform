# -*- coding: utf-8 -*-
"""车辆 PRD：导出不做；一凭证一项目；越权 project_id 拒收。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx")
LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_vehicle_auth_export_log.txt"
)
VER = "V2.2.6"
TODAY = "2026-08-13"
NOTE = (
    "明确本期进出场列表不做导出；对接凭证一证一项目；"
    "上报 project_id 不在凭证授权范围则拒收。"
)

REPLACEMENTS = [
    (
        "进出场列表若平台统一提供导出则启用；无则不做特例",
        "本期进出场列表不做导出（不提供导出按钮/能力）",
    ),
    (
        "若平台统一提供导出则启用；无则不做特例",
        "本期进出场列表不做导出",
    ),
    (
        "经平台授权的第三方对接凭证 | 按项目授权范围写入",
        "经平台授权的第三方对接凭证（一证一项目） | 仅可写入该凭证绑定的唯一项目；上报 project_id 与绑定项目不一致则拒收",
    ),
]


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def replace_in_text(full: str) -> tuple[str, int]:
    n = 0
    new = full
    for old, repl in REPLACEMENTS:
        if old in new:
            c = new.count(old)
            new = new.replace(old, repl)
            n += c
    # softer patterns for export rows
    if "导出" in new and "本期" not in new and ("若平台" in new or "统一提供" in new):
        pass
    return new, n


def apply_para(para) -> int:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    new, n = replace_in_text(full)
    if new != full:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = new
    return n


def bump_meta(doc: Document) -> None:
    for row in doc.tables[0].rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], VER)
        if key == "更新日期":
            set_cell_text(row.cells[1], TODAY)


def main() -> None:
    log: list[str] = []
    bak = SRC.with_name(
        f"车辆管理_PRD_导出凭证修订前_副本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    shutil.copy2(SRC, bak)
    log.append("backup " + bak.name)

    doc = Document(str(SRC))
    total = 0
    for para in doc.paragraphs:
        total += apply_para(para)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            blob = " | ".join(cells)

            # 2.3 列表导出 row
            if cells and cells[0].strip() == "列表导出":
                if len(row.cells) >= 3:
                    set_cell_text(row.cells[2], "本期不做导出")
                    log.append("T8 列表导出 -> 本期不做")
                total += 1

            # 6.4 interaction 导出 row
            if cells and cells[0].strip() == "导出":
                # 操作|触发|响应|异常
                if len(row.cells) >= 4:
                    set_cell_text(row.cells[0], "导出")
                    set_cell_text(row.cells[1], "—")
                    set_cell_text(row.cells[2], "本期不提供")
                    set_cell_text(row.cells[3], "无导出按钮")
                    log.append("6.4 导出行改为不提供")
                total += 1

            # 6.6.10 权限表
            if "对接凭证" in blob or ("调用上报接口" in blob and "授权" in blob):
                if len(row.cells) >= 3:
                    # 操作 | 可访问角色 | 说明
                    if "调用上报" in row.cells[0].text:
                        set_cell_text(row.cells[1], "经平台授权的第三方对接凭证（一证一项目）")
                        set_cell_text(
                            row.cells[2],
                            "凭证仅绑定一个项目；仅可写入该项目；上报 project_id 与绑定项目不一致则拒收",
                        )
                        log.append("6.6.10 权限一证一项目")
                        total += 1

            # NFR 安全行补一句（不改性能）
            if cells and cells[0].strip() == "安全" and len(row.cells) > 1:
                cur = row.cells[1].text.strip()
                add = "对接凭证一证一项目；project_id 超出绑定项目则拒收。"
                if "一证一项目" not in cur:
                    set_cell_text(row.cells[1], cur.rstrip("；。") + "；" + add)
                    log.append("NFR 安全补充一证一项目")
                    total += 1

            for cell in row.cells:
                for para in cell.paragraphs:
                    total += apply_para(para)

    # 6.6.9 业务规则 bullets
    from copy import deepcopy

    rules_added = False
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("6.6.9") and "业务规则" in para.text:
            # find last bullet
            last = None
            for j in range(i + 1, len(doc.paragraphs)):
                p = doc.paragraphs[j]
                style = p.style.name if p.style else ""
                if style.startswith("Heading"):
                    break
                if p.text.strip():
                    last = p
            if last and "一证一项目" not in "\n".join(
                p.text for p in doc.paragraphs[i : i + 20]
            ):
                texts = [
                    "对接凭证一证一项目：一张凭证仅绑定一个项目。",
                    "授权校验：上报 project_id 必须等于该凭证绑定项目，否则拒绝该条（不落库）。",
                    "本期不做进出场列表导出。",
                ]
                anchor = last._element
                for text in texts:
                    new_el = deepcopy(last._element)
                    for node in new_el.iter():
                        if node.tag.endswith("}t"):
                            node.text = ""
                    anchor.addnext(new_el)
                    first = True
                    for node in new_el.iter():
                        if node.tag.endswith("}t"):
                            if first:
                                node.text = text
                                first = False
                            else:
                                node.text = ""
                    anchor = new_el
                rules_added = True
                log.append("added 6.6.9 bullets")
            break
    if not rules_added:
        log.append("WARN: 6.6.9 bullets maybe already present or heading missing")

    bump_meta(doc)
    doc.save(str(SRC))
    log.append(f"python saved total~{total}")

    # Word: revision + AC for unauthorized project_id + sanitize
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        wdoc = word.Documents.Open(str(SRC), ReadOnly=False)
        tbl = wdoc.Tables(2)
        first = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first != VER:
            tbl.Rows(2).Select()
            word.Selection.InsertRowsAbove(1)
            row = tbl.Rows(2)
            ncols = row.Cells.Count
            vals = [VER, TODAY, NOTE] if ncols == 3 else [VER, "—", TODAY, NOTE]
            for i, v in enumerate(vals, start=1):
                if i <= ncols:
                    row.Cells(i).Range.Text = v
            log.append("rev " + VER)

        # find 6.6 AC table
        def find_table(needle: str) -> int | None:
            for i in range(1, wdoc.Tables.Count + 1):
                if needle in wdoc.Tables(i).Range.Text:
                    return i
            return None

        i66 = find_table("无 external_id 拒收") or find_table("project_id+external_id")
        if i66:
            tblx = wdoc.Tables(i66)
            exists = False
            for r in range(2, tblx.Rows.Count + 1):
                v = tblx.Cell(r, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
                title = tblx.Cell(r, 2).Range.Text.strip().replace("\r", "").replace("\x07", "")
                if v == "AC-06" or "授权范围" in title or "绑定项目" in title:
                    exists = True
            if not exists:
                tblx.Rows(tblx.Rows.Count).Select()
                word.Selection.InsertRowsBelow(1)
                r = tblx.Rows.Count
                vals = [
                    "AC-06",
                    "越权项目拒收",
                    "凭证绑定项目 A，上报 project_id=B",
                    "调用上报",
                    "拒绝该条；不落库",
                ]
                for c, val in enumerate(vals, start=1):
                    tblx.Cell(r, c).Range.Text = val
                log.append("added AC-06 越权项目拒收")
            else:
                log.append("AC-06 already exists")

        # 6.4 AC add 无导出 if useful
        i64 = find_table("车牌/方向筛选") or find_table("无手动填报")
        if i64:
            tblx = wdoc.Tables(i64)
            exists = False
            for r in range(2, tblx.Rows.Count + 1):
                title = tblx.Cell(r, 2).Range.Text.strip().replace("\r", "").replace("\x07", "")
                if "导出" in title:
                    exists = True
            if not exists:
                tblx.Rows(tblx.Rows.Count).Select()
                word.Selection.InsertRowsBelow(1)
                r = tblx.Rows.Count
                vals = [
                    "AC-05",
                    "无导出",
                    "项目用户打开进出场记录",
                    "查找导出",
                    "无导出按钮/入口",
                ]
                for c, val in enumerate(vals, start=1):
                    if c <= tblx.Rows(r).Cells.Count:
                        tblx.Cell(r, c).Range.Text = val
                log.append("added 6.4 AC-05 无导出")

        tmp = SRC.with_name("_vehicle_auth_tmp.docx")
        if tmp.exists():
            tmp.unlink()
        wdoc.SaveAs2(str(tmp), FileFormat=12)
        wdoc.Close(False)
        shutil.copy2(tmp, SRC)
        tmp.unlink()

        wdoc = word.Documents.Open(str(SRC), ReadOnly=True)
        log.append(f"WORD_OK {wdoc.Paragraphs.Count}")
        wdoc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    d2 = Document(str(SRC))
    blob = "\n".join(p.text for p in d2.paragraphs)
    for t in d2.tables:
        for row in t.rows:
            blob += "\n" + "|".join(c.text for c in row.cells)
    for needle in [
        "本期不做导出",
        "一证一项目",
        "越权项目拒收",
        "若平台统一提供导出",
        "无导出按钮",
    ]:
        log.append(f"has[{needle}]={needle in blob}")

    LOG.write_text("\n".join(log), encoding="utf-8")
    print("done")


if __name__ == "__main__":
    main()
