# -*- coding: utf-8 -*-
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
doc = Document(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx")
out = []
for row in doc.tables[0].rows:
    out.append("META: " + row.cells[0].text.strip() + " = " + row.cells[1].text.strip()[:100])
for row in doc.tables[1].rows:
    out.append("REV: " + " | ".join(c.text.strip()[:70] for c in row.cells))

for p in doc.paragraphs:
    t = p.text.strip()
    if any(k in t for k in ("6.10", "预警中心", "处置任务", "图 6.10", "批量消除", "批量已读")):
        out.append(f"P|{p.style.name}|{t[:160]}")

draw = 0
for p in doc.paragraphs:
    if p._element.xpath(".//*[local-name()='drawing']"):
        draw += 1
out.append(f"DRAWINGS={draw}")

# leftover old wording
old_hits = []
for p in doc.paragraphs:
    t = p.text
    if "预警消息与详情" in t or "我的代办" in t and "预警" in t:
        old_hits.append(t[:120])
for table in doc.tables:
    for row in table.rows:
        blob = "|".join(c.text for c in row.cells)
        if "个人中心-我的代办" in blob or "个人中心预警消息" in blob:
            old_hits.append(blob[:120])
out.append("OLD_HITS=" + str(len(old_hits)))
for h in old_hits[:20]:
    out.append("OLD: " + h)

for ti, table in enumerate(doc.tables):
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        joined = "|".join(cells)
        if "F-LABOR-10" in joined or "AC-01" in joined or "预警中心列表" in joined or "批量已读" in joined:
            out.append(f"T{ti}: " + " | ".join(c[:45] for c in cells))

Path(r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_prd_v212_verify.txt").write_text(
    "\n".join(out), encoding="utf-8"
)
print("written", len(out), "drawings", draw)
