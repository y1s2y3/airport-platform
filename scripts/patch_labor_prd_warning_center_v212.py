# -*- coding: utf-8 -*-
"""同步个人中心·预警中心正式工程口径至人员实名制 PRD（V2.1.2）。"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph

DOC_PATH = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)
SHOT_DIR = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\docs\prd-shots\labor-warning-center"
)
SHOTS = [
    ("01-web-warning-center-list.png", "图 6.10-1 Web 个人中心 · 预警中心列表"),
    ("02-web-warning-detail.png", "图 6.10-2 Web 预警详情（自预警中心进入）"),
    ("03-app-warning-center.png", "图 6.10-3 APP 个人中心 · 预警中心"),
]


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def set_para_text(p, value):
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value):
    if not cell.paragraphs:
        cell.text = value
        return
    set_para_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def apply_pairs(doc, pairs):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, pairs)
    return n


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc, exact=None, startswith=None, contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if startswith is not None and t.startswith(startswith):
            return p
        if contains is not None and contains in t:
            return p
    return None


def format_table(table):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_heading_space_after(doc):
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def update_meta_and_revision(doc):
    meta = doc.tables[0]
    for row in meta.rows:
        if row.cells[0].text.strip() == "版本号":
            set_cell_text(row.cells[1], "V2.1.2")
        if row.cells[0].text.strip() == "更新日期":
            set_cell_text(row.cells[1], "2026-08-14")

    rev = doc.tables[1]
    # 若已有 V2.1.2 则跳过
    for row in rev.rows[1:]:
        if row.cells[0].text.strip() == "V2.1.2":
            set_cell_text(
                row.cells[3],
                "个人中心人员预警迁入「预警中心」；类型「处置任务/通知」；待办/通知不再承载；"
                "批量已读/消除；列表字段调整；补充 Web/APP 原型截图。",
            )
            set_cell_text(row.cells[2], "2026-08-14")
            return
    # 复制表头样式新增一行
    tbl = rev._tbl
    new_tr = deepcopy(tbl.tr_lst[-1])
    tbl.append(new_tr)
    row = rev.rows[-1]
    vals = [
        "V2.1.2",
        "—",
        "2026-08-14",
        "个人中心人员预警迁入「预警中心」；类型列「处置任务/通知」；待办与通知信息不再挂人员预警；"
        "支持批量已读（仅未读通知）与批量消除（列表软删、不关闭业务预警）；"
        "列表增加项目简称、处理人姓名；预警描述不含人员编号；补充 Web/APP 原型截图。",
    ]
    for i, v in enumerate(vals):
        set_cell_text(row.cells[i], v)
    format_table(rev)


def rewrite_610_section(doc):
    """重写 6.10 关键关键段落与关键表。"""
    mapping = {
        "6.10 个人中心预警消息与详情（功能ID：F-LABOR-10）": "6.10 个人中心·预警中心与详情（功能ID：F-LABOR-10）",
        "在Web、APP个人中心，提供人员预警的消息通知与详情。": (
            "在 Web、APP 个人中心提供「预警中心」：汇总人员实名制处置任务与通知类预警，支持列表筛选、分页、"
            "详情查看；Web 端默认接收人可对处置任务类进行处置并关闭，APP 仅查看。"
        ),
        "任务类预警：Web 个人中心（仅默认接收人）或项目 Web「预警清单/详情」可处置；APP 个人中心仅查看、不做处置；通知类仅查阅。口径与 F-LABOR-06 一致。": (
            "人员预警统一在「预警中心」展示，不再进入「我的待办 / 通知信息」。"
            "处置任务类：Web 预警中心（仅默认接收人）或项目 Web「预警清单/详情」可处置；APP 预警中心仅查看。"
            "通知类：预警中心内查阅与已读，无业务关闭。口径与 F-LABOR-06 一致。"
        ),
        "展示人员预警及其处置记录；Web 端任务类仅默认接收人可在个人中心处置并关闭；APP 端仅查看；消息投递遵循个人中心公共能力。": (
            "规则触发后投递至个人中心「预警中心」；展示处置任务与通知两类条目及详情处置记录。"
            "Web 端处置任务类仅默认接收人可处置并关闭；APP 仅查看。"
            "「批量消除」仅从预警中心列表软删，不调用业务侧处置并关闭；「批量已读」仅对通知类「未读」生效。"
        ),
        "预警状态口径与 F-LABOR-06 一致（待处理/已关闭/已通知）；": (
            "业务预警状态仍与 F-LABOR-06 一致（待处理/已关闭/已通知）。"
            "预警中心列表映射：处置任务类→类型「处置任务」、状态「待处置/已处置」；"
            "通知类→类型「通知」、状态「未读/已读」。"
        ),
        "通知类无处置入口，已阅即可。": "通知类无业务关闭入口；可批量/单条标为已读。",
        "代办": "预警中心列表",
        "手动处理，提示文字": "处置任务类详情提示与自动关闭引导文案（见下表）；列表字段见本节字段表。",
        "各Tab页新增所属模块字段，标记代办通知数据来源；\n历史数据同步一遍，或假数据可删除。": (
            "个人中心各公共 Tab（待办/已办/发起/抄送/通知）可保留「所属模块」字段标记数据来源；"
            "人员实名制预警专属数据仅出现在「预警中心」，历史待办/通知中的人员预警数据应迁移或清理，避免双挂。"
        ),
        "6.10.12个人中心-补充需求": "6.10.12 个人中心补充说明",
    }

    in_rules = False
    rules_patched = False
    # 快照段落 XML 元素，避免插入后迭代集合变化
    paras = list(doc.paragraphs)
    for p in paras:
        t = p.text.strip()
        style_name = p.style.name if p.style else ""
        if t.startswith("6.10.9"):
            in_rules = True
        elif style_name.startswith("Heading") and t.startswith("6.10."):
            in_rules = False
        elif style_name.startswith("Heading") and not t.startswith("6.10"):
            in_rules = False

        if t in mapping:
            set_para_text(p, mapping[t])
            continue

        if in_rules and not rules_patched and t in ("无", "") and "List" in style_name:
            set_para_text(
                p,
                "人员预警仅在「预警中心」展示；「我的待办 / 通知信息 / 流程中心 / 消息提醒」不再重复挂人员预警。",
            )
            cur = p
            extra = [
                "预警中心类型：「处置任务」「通知」。处置任务对应业务任务类（手动/自动关闭）；通知对应通知类。",
                "批量已读：仅对状态为「未读」的通知条目生效；处置任务「待处置」不受影响。",
                "批量消除：从预警中心列表软删（消息消除），不关闭业务预警、不写处置关闭时间线。",
                "预警描述展示规则名称+人员姓名+触发原因摘要，不含人员编号；项目名称用简称；处理人展示姓名。",
                "APP 预警中心仅查看详情，无处置并关闭入口。",
            ]
            for text in extra:
                cur = insert_paragraph_after(cur, text, style=p.style)
            rules_patched = True


def update_610_tables(doc):
    """更新原型页表、交互表、验收表等。"""
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows:
            continue
        head = rows[0]

        # 原型页面表
        if head[:4] == ["页面名称", "页面路径", "核心内容", "备注"]:
            if any("个人中心" in r[0] for r in rows[1:]):
                data = [
                    ["个人中心·预警中心列表", "个人中心 → 预警中心", "处置任务/通知列表；筛选、分页、批量已读/消除", "Web/APP 同源口径"],
                    ["个人中心·预警详情（Web）", "预警中心 → 详情", "预警信息 + 时间线；默认接收人可处置并关闭", "回跳 tab=warning-center"],
                    ["个人中心·预警详情（APP）", "预警中心 → 详情", "预警信息 + 时间线（只读）", "不可处置"],
                ]
                # 保证行数
                while len(table.rows) < 1 + len(data):
                    table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
                for i, row_data in enumerate(data):
                    for j, val in enumerate(row_data):
                        set_cell_text(table.rows[i + 1].cells[j], val)
                # 清空多余行
                for i in range(1 + len(data), len(table.rows)):
                    for j in range(len(head)):
                        set_cell_text(table.rows[i].cells[j], "")
                format_table(table)

        # 交互说明表（接收/查看消息）
        if head[:4] == ["操作", "触发方式", "系统响应", "异常处理"]:
            blob = "\n".join("|".join(r) for r in rows)
            if "我的代办" in blob or "去处置" in blob:
                data = [
                    ["接收/查看", "个人中心 → 预警中心", "展示人员预警处置任务与通知", "无消息空态"],
                    ["筛选/分页", "类型、状态、关键词；分页器", "刷新当前页列表", "无结果空态"],
                    ["批量已读", "勾选后点「批量已读」", "仅「未读」通知改为已读", "无未读则提示"],
                    ["批量消除预警", "勾选后点「批量消除预警」", "列表软删所选条目；不关闭业务预警", "未勾选提示"],
                    [
                        "打开详情",
                        "点击「详情」",
                        "Web：展示详情，默认接收人对处置任务可处置；APP/非默认接收人只读",
                        "预警已删除/无权限提示",
                    ],
                ]
                while len(table.rows) < 1 + len(data):
                    table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
                for i, row_data in enumerate(data):
                    for j, val in enumerate(row_data):
                        set_cell_text(table.rows[i + 1].cells[j], val)
                for i in range(1 + len(data), len(table.rows)):
                    for j in range(len(head)):
                        set_cell_text(table.rows[i].cells[j], "")
                format_table(table)

        # 验收表 AC
        if head[:5] == ["编号", "验收项", "Given", "When", "Then"]:
            blob = "\n".join("|".join(r) for r in rows)
            if "AC-01" in blob and "个人中心" in blob:
                data = [
                    [
                        "AC-01",
                        "Web、APP 预警中心可达",
                        "用户为预警责任人/接收人",
                        "打开个人中心 → 预警中心",
                        "可见对应处置任务/通知并可打开详情；APP 仅查看不可处置",
                    ],
                    [
                        "AC-02",
                        "处置任务手动关闭",
                        "Web 预警中心进入详情且当前用户为默认接收人",
                        "提交处置并关闭",
                        "业务预警关闭；预警中心对应条状态变为「已处置」；APP 无处置入口",
                    ],
                    [
                        "AC-03",
                        "自动关闭类引导",
                        "打开预警中心中自动关闭类处置任务详情",
                        "Web 按提示前往实名制/闸机侧处理，次日定时关闭；APP 仅查看",
                        "条件满足后业务关闭且预警中心条变为已处置；APP 不操作",
                    ],
                    [
                        "AC-04",
                        "不双挂待办/通知",
                        "存在人员预警",
                        "查看我的待办 / 通知信息（及 APP 流程中心/消息提醒）",
                        "不再出现同条人员预警；仅预警中心展示",
                    ],
                    [
                        "AC-05",
                        "批量已读与消除",
                        "预警中心存在未读通知与任一条目",
                        "批量已读、批量消除",
                        "未读通知变已读；消除后列表移除且业务预警未关闭",
                    ],
                    [
                        "AC-06",
                        "APP 无业务菜单",
                        "APP 登录",
                        "浏览业务菜单",
                        "无实名制看板/台账/预警清单等",
                    ],
                ]
                while len(table.rows) < 1 + len(data):
                    table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
                for i, row_data in enumerate(data):
                    for j, val in enumerate(row_data):
                        set_cell_text(table.rows[i + 1].cells[j], val)
                for i in range(1 + len(data), len(table.rows)):
                    for j in range(len(head)):
                        set_cell_text(table.rows[i].cells[j], "")
                format_table(table)

        # 字段提示表「流程名称」等 — 更新为预警中心口径
        if head[:2] == ["预警类型", "提示文字"] and len(rows) > 1 and rows[1][0] == "流程名称":
            data = [
                ["列表-所属模块", "固定「人员实名」"],
                ["列表-项目名称", "项目简称（如 T2项目）"],
                ["列表-预警描述", "规则名称：姓名。触发原因摘要（不含人员编号）"],
                ["列表-处理人", "展示姓名（默认接收人或关闭操作人）"],
                ["列表-类型", "处置任务 / 通知"],
                ["列表-状态", "处置任务：待处置/已处置；通知：未读/已读"],
                ["列表-消息时间", "触发时间或关闭时间"],
                ["详情页", "Web 端且当前用户为默认接收人时，处置能力与 F-LABOR-06 对齐；APP/非默认接收人只读"],
                [
                    "操作-处理预警",
                    "按手动/自动区分。手动：仅 Web 默认接收人可处置并关闭；自动：Web 弹窗引导「去处理」；APP 无此操作",
                ],
            ]
            while len(table.rows) < 1 + len(data):
                table._tbl.append(deepcopy(table._tbl.tr_lst[-1]))
            for i, row_data in enumerate(data):
                for j, val in enumerate(row_data):
                    set_cell_text(table.rows[i + 1].cells[j], val)
            for i in range(1 + len(data), len(table.rows)):
                for j in range(len(head)):
                    set_cell_text(table.rows[i].cells[j], "")
            format_table(table)


def insert_shots_after_6104(doc):
    anchor = find_para(doc, exact="6.10.4 原型页面")
    if anchor is None:
        raise RuntimeError("找不到 6.10.4 原型页面")
    # 找到下一标题前的最后一个段落作为插入点；直接插在标题后
    cur = anchor
    # 若已有图注则跳过重复插入
    for p in doc.paragraphs:
        if "图 6.10-1" in p.text:
            return

    for fname, caption in reversed(SHOTS):
        # 因为 insert_after 是紧跟 anchor，倒序插入可保持正序
        pass

    cur = anchor
    for fname, caption in SHOTS:
        img_p = insert_paragraph_after(cur, "")
        run = img_p.add_run()
        path = SHOT_DIR / fname
        if not path.exists():
            raise FileNotFoundError(path)
        # APP 窄屏图略窄
        width = Inches(3.2) if "app" in fname else Inches(6.0)
        run.add_picture(str(path), width=width)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = insert_paragraph_after(img_p, caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        cur = cap


def main():
    doc = Document(str(DOC_PATH))

    update_meta_and_revision(doc)

    pairs = [
        (
            "└── 预警消息与详情                               （F-LABOR-10）",
            "└── 预警中心与详情                               （F-LABOR-10）",
        ),
        (
            "4）任务类预警仅默认接收人可在 Web（个人中心或预警清单）处置；自动关闭类次日定时关闭；APP 仅查看；",
            "4）处置任务类预警仅默认接收人可在 Web（个人中心·预警中心或预警清单）处置；自动关闭类次日定时关闭；APP 预警中心仅查看；",
        ),
        (
            "指挥部 Web / APP 经个人中心触达预警（F-LABOR-10）：Web 端任务类仅默认接收人可处置，APP 仅查看；项目侧亦可在预警清单处置。",
            "指挥部 Web / APP 经个人中心「预警中心」触达预警（F-LABOR-10）：Web 端处置任务类仅默认接收人可处置，APP 仅查看；项目侧亦可在预警清单处置。人员预警不进入「我的待办/通知信息」。",
        ),
        (
            "触发/升级/关闭/通知均可推送个人中心消息。",
            "触发/升级/关闭/通知均可投递个人中心「预警中心」。",
        ),
        (
            "仍推送个人中心消息；",
            "仍投递个人中心预警中心；",
        ),
        (
            "APP 无预警清单菜单，仅消息中心",
            "APP 无预警清单菜单，仅个人中心·预警中心",
        ),
        (
            "APP：仅个人中心预警消息/详情；无本模块业务菜单",
            "APP：仅个人中心·预警中心与详情；无本模块业务菜单",
        ),
        (
            "个人中心 | 人员预警消息通知与详情；Web 任务类仅默认接收人可处置，APP 仅查看 | 预警触发/升级/关闭后推送或短时可见",
            "个人中心 | 预警中心（处置任务/通知）与详情；Web 处置任务类仅默认接收人可处置，APP 仅查看；不进入待办/通知 | 预警触发/升级/关闭后投递预警中心",
        ),
        (
            "个人中心消息 | 个人中心模块 | 支持预警消息与详情；Web 端任务类仅默认接收人可处置，APP 仅查看",
            "个人中心·预警中心 | 个人中心模块 | 支持预警中心列表与详情；Web 端处置任务类仅默认接收人可处置，APP 仅查看；待办/通知不承载人员预警",
        ),
        (
            "F-LABOR-10 | 个人中心 | 个人中心预警消息与详情 | 消息通知与详情；Web 任务类仅默认接收人可处置，APP 仅查看 | Web+APP | — | Web 个人中心（仅默认接收人）与项目 Web 可处置；APP 仅查看",
            "F-LABOR-10 | 个人中心 | 个人中心·预警中心与详情 | 预警中心汇总处置任务/通知；筛选分页；批量已读/消除；Web 处置任务类仅默认接收人可处置，APP 仅查看 | Web+APP | — | 不进入待办/通知；消除≠业务关闭",
        ),
        (
            "三端个人中心可查看预警消息与详情；任务类仅默认接收人可在 Web 个人中心或项目 Web 预警清单处置；APP 仅查看不做处置",
            "三端经个人中心「预警中心」查看处置任务/通知与详情；处置任务类仅默认接收人可在 Web 预警中心或项目 Web 预警清单处置；APP 仅查看不做处置",
        ),
        (
            "项目看板、台账、考勤、预警处置、轨迹外跳；个人中心预警消息/详情",
            "项目看板、台账、考勤、预警处置、轨迹外跳；个人中心·预警中心与详情",
        ),
        (
            "人员预警消息 | 个人中心 | 本模块 → 对端 | 触发/升级/关闭 | Web 个人中心支持默认接收人处置；APP 仅查看",
            "人员预警 → 预警中心 | 个人中心 | 本模块 → 对端 | 触发/升级/关闭 | 投递预警中心；Web 默认接收人可处置；APP 仅查看；不进待办/通知",
        ),
        (
            "推送个人中心消息，需做好工人健康情况排查，不强制要求退场。",
            "投递个人中心预警中心，需做好工人健康情况排查，不强制要求退场。",
        ),
        (
            "关闭代办事项并更新预警状态",
            "预警中心对应条变为已处置并更新业务预警状态",
        ),
        (
            "关闭代办并更新预警状态",
            "预警中心对应条变为已处置并更新业务预警状态",
        ),
    ]
    n = apply_pairs(doc, pairs)
    print("pair replacements:", n)

    rewrite_610_section(doc)
    update_610_tables(doc)
    insert_shots_after_6104(doc)
    set_heading_space_after(doc)

    # 权限表措辞
    apply_pairs(
        doc,
        [
            (
                "查看消息/详情；Web 处置（仅默认接收人）",
                "查看预警中心/详情；Web 处置（仅默认接收人）",
            ),
        ],
    )

    doc.save(str(DOC_PATH))
    print("saved:", DOC_PATH)


if __name__ == "__main__":
    main()
