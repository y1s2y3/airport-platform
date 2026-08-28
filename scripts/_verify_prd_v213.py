# -*- coding: utf-8 -*-
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
doc = Document(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx")
out = []
for row in doc.tables[0].rows:
    out.append("META: " + row.cells[0].text.strip() + " = " + row.cells[1].text.strip()[:80])
for row in doc.tables[1].rows:
    out.append("REV: " + " | ".join(c.text.strip()[:90] for c in row.cells))

keys = ["双路径", "路径A", "路径B", "标准上报", "标准接口", "图 5.2-1", "图 7.1-1", "闸机", "5.2", "6.4.1", "7.1"]
for p in doc.paragraphs:
    t = p.text.strip()
    if any(k in t for k in keys):
        out.append(f"P|{p.style.name}|{t[:180]}")

draw = sum(1 for p in doc.paragraphs if p._element.xpath(".//*[local-name()='drawing']"))
out.append(f"DRAWINGS={draw}")
bak = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_备份_V2.1.2_20260815.docx")
out.append(f"BACKUP_EXISTS={bak.exists()} size={bak.stat().st_size if bak.exists() else 0}")

# confirm old core flow still present
core = next((p.text for p in doc.paragraphs if p.text.strip().startswith("核心流程：")), "")
out.append("CORE_HAS_ROMA=" + str("一期经 ROMA" in core))
out.append("CORE_FIRST_LINE=" + core.splitlines()[1][:80] if len(core.splitlines()) > 1 else "missing")

Path(r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_prd_v213_verify.txt").write_text(
    "\n".join(out), encoding="utf-8"
)
print("\n".join(out[:40]))
print("...")
print(out[-5:])
