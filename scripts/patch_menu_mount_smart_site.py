# -*- coding: utf-8 -*-
"""将人员/车辆 PRD·架构中指挥部菜单挂载从「施工现场管理」改为「智慧工地监管」。路由不变。"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document

BASE = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
TODAY = "2026-08-13"
REPORT = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_menu_mount_patch_report.txt"
)

REPLACEMENTS = [
    ("挂施工现场管理 / 智慧工地监管", "挂智慧工地监管"),
    ("挂施工现场管理/智慧工地监管", "挂智慧工地监管"),
    ("施工现场管理 → 人员实名制管理", "智慧工地监管 → 人员实名制管理"),
    ("施工现场管理 → 车辆管理", "智慧工地监管 → 车辆管理"),
    ("施工现场管理 · 人员实名制管理", "智慧工地监管 · 人员实名制管理"),
    ("施工现场管理 · 车辆管理", "智慧工地监管 · 车辆管理"),
    ("挂施工现场管理", "挂智慧工地监管"),
    ("位于施工现场管理，不在安全看板", "位于智慧工地监管，不在安全看板"),
    ("位于施工现场管理，", "位于智慧工地监管，"),
    ("指挥部挂施工现场管理", "指挥部挂智慧工地监管"),
    ("指挥部「施工现场管理 → 人员实名制管理」", "指挥部「智慧工地监管 → 人员实名制管理」"),
    ("指挥部「施工现场管理 → 车辆管理」", "指挥部「智慧工地监管 → 车辆管理」"),
    ("指挥部菜单**挂「施工现场管理」**", "指挥部菜单**挂「智慧工地监管」**"),
    ("仅指挥部「施工现场管理」可见", "仅指挥部「智慧工地监管」可见"),
    ("施工现场管理（指挥部挂载）；智慧工地监管（项目包装）", "智慧工地监管（指挥部与项目侧挂载）"),
]


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def apply_replacements(text: str) -> tuple[str, int]:
    n = 0
    new = text
    for old, repl in REPLACEMENTS:
        if old in new:
            c = new.count(old)
            new = new.replace(old, repl)
            n += c
    # 兜底：指挥部语境下残留的「施工现场管理」→「智慧工地监管」
    if "施工现场管理" in new:
        for old, repl in [
            ("「施工现场管理」", "「智慧工地监管」"),
            ("施工现场管理下", "智慧工地监管下"),
            ("施工现场管理内", "智慧工地监管内"),
            ("施工现场管理中", "智慧工地监管中"),
        ]:
            if old in new:
                c = new.count(old)
                new = new.replace(old, repl)
                n += c
    return new, n


def replace_in_paragraph(para) -> int:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    new, n = apply_replacements(full)
    if new != full:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = new
    return n


def bump_meta_and_revision(doc: Document, ver: str, note: str, log: list[str]) -> None:
    if not doc.tables:
        return
    meta = doc.tables[0]
    for row in meta.rows:
        key = row.cells[0].text.strip()
        if key == "版本号" and len(row.cells) > 1:
            set_cell_text(row.cells[1], ver)
        if key == "更新日期" and len(row.cells) > 1:
            set_cell_text(row.cells[1], TODAY)

    if len(doc.tables) < 2:
        return
    rev = doc.tables[1]
    headers = [c.text.strip() for c in rev.rows[0].cells]
    if not headers or ("版本" not in headers[0] and "修订" not in "".join(headers)):
        # 也可能第一列就是版本
        pass
    if len(rev.rows) < 2:
        return
    first = rev.rows[1].cells[0].text.strip()
    if first == ver:
        return
    new_tr = deepcopy(rev.rows[1]._tr)
    rev._tbl.insert(1, new_tr)
    row = rev.rows[1]
    vals = [ver, "—", TODAY, note]
    for i, v in enumerate(vals):
        if i < len(row.cells):
            set_cell_text(row.cells[i], v)
    log.append(f"  + revision {ver}")


def process_docx(path: Path, ver: str, note: str, log: list[str]) -> None:
    bak = path.with_name(
        f"{path.stem}_菜单挂载修订前_副本_{datetime.now().strftime('%Y%m%d')}{path.suffix}"
    )
    if not bak.exists():
        shutil.copy2(path, bak)
        log.append(f"backup {bak.name}")

    doc = Document(str(path))
    total = 0
    for para in doc.paragraphs:
        total += replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += replace_in_paragraph(para)

    bump_meta_and_revision(doc, ver, note, log)
    doc.save(str(path))
    log.append(f"saved {path.name}, replacements≈{total}")

    doc2 = Document(str(path))
    left = []
    for p in doc2.paragraphs:
        if "施工现场管理" in p.text:
            left.append("P: " + p.text[:160])
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                if "施工现场管理" in cell.text:
                    left.append("T: " + cell.text.replace("\n", " | ")[:160])
    log.append(f"  leftover 施工现场管理: {len(left)}")
    for x in left[:20]:
        log.append("   - " + x)
    # confirm 智慧工地监管 HQ tree
    hits = 0
    for p in doc2.paragraphs:
        if "智慧工地监管" in p.text and ("人员实名制" in p.text or "车辆管理" in p.text):
            hits += 1
    log.append(f"  智慧工地监管+模块 hits in paras: {hits}")


def patch_architecture_md(path: Path, ver_bump: str, log: list[str]) -> None:
    text = path.read_text(encoding="utf-8")
    orig = text
    n = 0
    for old, repl in REPLACEMENTS:
        if old in text:
            c = text.count(old)
            text = text.replace(old, repl)
            n += c

    # version bump
    text = re.sub(
        r"(\| 版本号 \| )V[\d.]+( \|)",
        rf"\g<1>{ver_bump}\2",
        text,
        count=1,
    )
    text = re.sub(
        r"(\| 更新日期 \| )\d{4}-\d{2}-\d{2}( \|)",
        rf"\g<1>{TODAY}\2",
        text,
        count=1,
    )
    rev_line = (
        f"| {ver_bump} 修订 | 指挥部人员/车辆菜单挂载由「施工现场管理」调整为「智慧工地监管」"
        f"（页面路由 `/labor/*`、`/vehicle/*` 不变）。 |"
    )
    if f"| {ver_bump} 修订 |" not in text:
        # insert after 模块编码 / 侧栏菜单名 row, before first V*.* 修订 or after 模块编码
        lines = text.splitlines(keepends=True)
        out = []
        inserted = False
        for i, line in enumerate(lines):
            out.append(line)
            if inserted:
                continue
            # after last meta identity row before existing revisions
            if re.match(r"\| V[\d.]+ 修订 \|", line):
                out.insert(-1, rev_line + "\n")
                inserted = True
        text = "".join(out) if inserted else text
        if not inserted:
            # fallback: after 模块编码 line
            text2 = []
            for line in text.splitlines(keepends=True):
                text2.append(line)
                if "| 模块编码 |" in line and not inserted:
                    text2.append(rev_line + "\n")
                    inserted = True
            text = "".join(text2)

    if text != orig:
        path.write_text(text, encoding="utf-8")
        log.append(f"patched md {path.name}, replacements≈{n}")
    else:
        log.append(f"no change md {path.name}")

    left = [ln.strip() for ln in text.splitlines() if "施工现场管理" in ln]
    log.append(f"  leftover lines: {len(left)}")
    for ln in left:
        log.append("   - " + ln[:140])


def patch_legacy_prd_md(path: Path, log: list[str]) -> None:
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    orig = text
    for old, repl in REPLACEMENTS:
        text = text.replace(old, repl)
    # labor specific compare table
    text = text.replace(
        "| 指挥部挂载 | 安全看板 | 施工现场管理 |",
        "| 指挥部挂载 | 安全看板 | 智慧工地监管 |",
    )
    if text != orig:
        path.write_text(text, encoding="utf-8")
        log.append(f"patched legacy {path.name}")
    else:
        log.append(f"no change legacy {path.name}")


def main() -> None:
    log: list[str] = []
    note = "指挥部菜单挂载由「施工现场管理」调整为「智慧工地监管」；页面路由不变。"
    process_docx(BASE / "人员实名制_PRD_待评审.docx", "V2.1.3", note, log)
    process_docx(BASE / "车辆管理_PRD.docx", "V2.2.2", note, log)
    patch_architecture_md(BASE / "product-architecture-人员实名制管理-v1.md", "V1.45", log)
    patch_architecture_md(BASE / "product-architecture-车辆进出场管理-v1.md", "V1.34", log)
    patch_legacy_prd_md(BASE / "prd-labor-v1.md", log)
    patch_legacy_prd_md(BASE / "prd-vehicle-v1.md", log)
    REPORT.write_text("\n".join(log), encoding="utf-8")
    print("OK wrote", REPORT)


if __name__ == "__main__":
    main()
