# -*- coding: utf-8 -*-
"""从人员实名制 PRD 删除已废止的「API 业务契约」小节标题。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)


def main() -> None:
    bak = SRC.with_name(
        f"人员实名制_PRD_待评审_去API契约节前_副本_{datetime.now().strftime('%Y%m%d')}.docx"
    )
    if not bak.exists():
        shutil.copy2(SRC, bak)
        print("backup", bak.name)
    else:
        print("backup exists", bak.name)

    d = Document(str(SRC))

    to_delete: list[int] = []
    for i, para in enumerate(d.paragraphs):
        t = para.text.strip()
        if "API 业务契约" not in t and not (t.startswith("6.") and "业务契约" in t):
            continue
        to_delete.append(i)
        j = i + 1
        while j < len(d.paragraphs):
            nt = d.paragraphs[j].text.strip()
            ns = d.paragraphs[j].style.name if d.paragraphs[j].style else ""
            if ns.startswith("Heading"):
                break
            if nt == "" or "本期不写" in nt or "规格验收以本章" in nt or "不作为验收依据" in nt:
                to_delete.append(j)
                j += 1
                continue
            break

    print("delete indices", sorted(set(to_delete)))
    for i in sorted(set(to_delete), reverse=True):
        el = d.paragraphs[i]._element
        el.getparent().remove(el)

    # 6.2 节内顺延：原 6.2.10/11/12 → 6.2.9/10/11
    renames = [
        ("6.2.12 ", "6.2.11 "),
        ("6.2.11 ", "6.2.10 "),
        ("6.2.10 ", "6.2.9 "),
    ]
    for para in d.paragraphs:
        full = "".join(r.text for r in para.runs) if para.runs else para.text
        for old, new in renames:
            if full.startswith(old):
                new_full = new + full[len(old) :]
                if para.runs:
                    para.runs[0].text = new_full
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.text = new_full
                print("renumber:", new_full[:80])
                break

    d.save(str(SRC))
    print("saved", SRC.name)

    d2 = Document(str(SRC))
    hits = [p.text for p in d2.paragraphs if "业务契约" in p.text]
    print("remaining paragraph hits:", hits)
    for table in d2.tables:
        for row in table.rows:
            for cell in row.cells:
                if "业务契约" in cell.text:
                    print("remaining table cell:", cell.text[:80].replace("\n", " "))
    print("6.2 Heading 4:")
    for p in d2.paragraphs:
        if p.style and p.style.name == "Heading 4" and p.text.startswith("6.2."):
            print(" ", p.text)


if __name__ == "__main__":
    main()
