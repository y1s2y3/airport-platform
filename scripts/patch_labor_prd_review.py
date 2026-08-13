# -*- coding: utf-8 -*-
"""Patch 人员实名制_PRD_待评审.docx per review decisions."""
from docx import Document

PATH = r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def set_cell_text(cell, value):
    if not cell.paragraphs:
        cell.text = value
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def main():
    doc = Document(PATH)

    pairs = [
        (
            "状态仅「待处理→已关闭」；9 条规则可配（默认全部关闭）；不做强制拦截",
            "任务类状态「待处理→已关闭」，通知类终态「已通知」；9 条规则可配（默认开/关以 6.7.8「预警配置规则」为准）；不做强制拦截",
        ),
        (
            "三端个人中心可查看预警消息与详情；处置仅项目 Web",
            "三端个人中心可查看预警消息与详情，任务类可由默认接收人处置；亦可在项目 Web 预警清单处置",
        ),
        (
            "跨项目合计：在岗且为特种作业人员的人数（工种/职务以“特种-”开头）。",
            "跨项目合计：在岗且 isSpecial（是否特种作业）为是的人数。",
        ),
        ("高龄提醒（男60/女50）", "高龄提醒（男65/女60）"),
        ("关，22", "关，20"),
        (
            "在岗人员按年龄分段：30岁以下（年龄≤30）、31-40岁、41-50岁、51-60岁、60岁以上。",
            "在岗人员按年龄分段：30岁以下（年龄≤30）、31-40岁、41-50岁、51-60岁（含60）、60岁以上（年龄>60，不含60）。",
        ),
        (
            "进入看板后按当前组织范围聚合人员主档、考勤与预警：指挥部为全项目汇总；项目端为当前项目。指标口径与指挥部「实名制统计」一致，具体计算见本节 6.3.8「通用口径」及各字段「说明」列。结构环图、出勤趋势、预警趋势均在同一统计日与近 30 日窗口内计算。",
            "进入看板后按当前项目聚合人员主档、考勤与预警（本功能仅项目 Web；指挥部跨项目汇总见 F-LABOR-01 实名制统计）。指标计算口径与指挥部「实名制统计」对齐，具体见本节 6.3.8「通用口径」及各字段「说明」列。结构环图、出勤趋势、预警趋势均在同一统计日与近 30 日窗口内计算。",
        ),
        ("指挥部/项目 Web 同布局", "仅项目 Web；指标口径与指挥部实名制统计对齐"),
        (
            "预警处置在「预警清单/详情」完成。",
            "任务类预警可在个人中心（Web/APP）或项目 Web「预警清单/详情」处置；通知类仅查阅。口径与 F-LABOR-06 一致。",
        ),
        (
            "只读展示人员预警及其处置记录；消息投递遵循个人中心公共能力。",
            "展示人员预警及其处置记录；任务类支持默认接收人在个人中心处置并关闭；消息投递遵循个人中心公共能力。",
        ),
        ("消息通知与详情只读", "消息通知与详情；任务类可处置"),
        ("处置在项目 Web", "个人中心与项目 Web 均可处置（任务类）"),
        (
            "指挥部维护全平台共用（不按项目拆分）劳务黑名单；名单中人员发生进场时，生成「黑名单人员进场预警」；二期不做闸机强制拦截，仅预警并由责任人处置。",
            "指挥部维护全平台共用（不按项目拆分）劳务黑名单；每日定时扫描各项目在岗人员，命中黑名单且规则开启时生成「黑名单人员进场预警」；二期不做闸机强制拦截，仅预警并由责任人处置。",
        ),
        (
            "维护劳务黑名单台账；命中进场事件时触发人员预警；移出后不再对新进场触发（已生成预警按处置闭环）。",
            "维护劳务黑名单台账；每日定时扫描在岗人员命中黑名单时触发人员预警；移出黑名单或人员离场后，已生成预警按自动关闭/处置闭环，且不再对新命中触发。",
        ),
        (
            "进场命中生成「黑名单人员进场预警」；不做强制拦截",
            "每日扫描在岗命中生成「黑名单人员进场预警」；不做强制拦截",
        ),
        (
            "劳务黑名单中的人员发生进场（刷卡或登记入场等）且该规则已开启时触发。",
            "每日定时任务扫描：人员在岗状态为「在岗」且证件号命中全平台劳务黑名单、且该规则已开启时触发。",
        ),
        ("黑名单人员在岗同步且规则启用", "黑名单人员处于在岗且规则启用（日扫命中）"),
        (
            "6.2.9 API 业务契约",
            "6.2.9 API 业务契约（本期不写，规格验收以本章字段与业务规则为准）",
        ),
        # specialCount project明细 if same wording exists
        (
            "本项目在岗且为特种作业人员的人数。",
            "本项目在岗且 isSpecial（是否特种作业）为是的人数。",
        ),
    ]

    count = 0
    for p in doc.paragraphs:
        count += replace_in_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    count += replace_in_paragraph(p, pairs)
    print("replacement hits:", count)

    for ti, table in enumerate(doc.tables[:6]):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            if cells[0] == "版本号" and cells[1] == "V1.0":
                set_cell_text(row.cells[1], "V1.1")
                print("meta version -> V1.1 @ table", ti)
            if cells[0] == "更新日期" and "2026-08-11" in cells[1]:
                set_cell_text(row.cells[1], "2026-08-12")
                print("meta date -> 2026-08-12 @ table", ti)
            if cells[0] == "状态" and cells[1] == "草稿":
                set_cell_text(row.cells[1], "待评审")
                print("meta status -> 待评审 @ table", ti)

    rev_done = False
    for ti, table in enumerate(doc.tables[:10]):
        headers = [c.text.strip() for c in table.rows[0].cells]
        if not headers or headers[0] != "版本":
            continue
        if "说明" not in headers and "日期" not in headers:
            continue
        # avoid duplicate V1.1
        existing = [r.cells[0].text.strip() for r in table.rows]
        if "V1.1" in existing:
            print("revision V1.1 already exists @ table", ti)
            rev_done = True
            break
        new_row = table.add_row()
        mapping = {
            "版本": "V1.1",
            "修订人": "",
            "日期": "2026-08-12",
            "说明": (
                "评审口径修订：个人中心可处置；明确已通知状态；规则默认以6.7.8为准；"
                "管理人员考勤默认20；高龄默认65/60；特种人数用isSpecial；看板仅项目、统计归指挥部；"
                "黑名单日扫在岗；年龄段60岁以上不含60；API契约本期不写"
            ),
        }
        for i, h in enumerate(headers):
            if i >= len(new_row.cells):
                break
            set_cell_text(new_row.cells[i], mapping.get(h, ""))
        print("revision row added @ table", ti, headers)
        rev_done = True
        break
    if not rev_done:
        print("WARN: revision table not found")

    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        if table.rows[0].cells[0].text.strip() != "接口用途":
            continue
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue
            for ci, cell in enumerate(row.cells):
                set_cell_text(
                    cell,
                    "—（本期不写 API 业务契约）" if ci == 0 else "—",
                )
        print("API table neutralized @", ti)
        break

    # Also update blacklist trigger parenthetical if still old in same cell
    extra_pairs = [
        (
            "(每日定时任务触发，一个人最多有一条该类型待处理预警)",
            "（一个人最多有一条该类型待处理预警）",
        ),
    ]
    # Only apply the above inside cells that already mention 黑名单 + 在岗 scan
    extra_hits = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full = "\n".join(p.text for p in cell.paragraphs)
                if "黑名单" in full and "在岗" in full:
                    for p in cell.paragraphs:
                        extra_hits += replace_in_paragraph(p, extra_pairs)
    print("extra blacklist paren hits:", extra_hits)

    out_v11 = PATH.replace("_待评审.docx", "_待评审_V1.1.docx")
    try:
        doc.save(PATH)
        print("saved original:", PATH)
        original_ok = True
    except PermissionError:
        original_ok = False
        print("WARN: original locked, skip overwrite")
    doc.save(out_v11)
    print("saved V1.1:", out_v11)
    if not original_ok:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
