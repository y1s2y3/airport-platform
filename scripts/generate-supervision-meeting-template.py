"""生成监理例会纪要 Word 模版，输出到 public/templates/。"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public" / "templates"
OUT_FILE = OUT_DIR / "监理例会纪要模板.docx"


def set_heading_space_after(paragraph, pt=6):
    paragraph.paragraph_format.space_after = Pt(pt)


def format_table_rows(table):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    set_heading_space_after(paragraph)
    return paragraph


def build_document():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("建设工程指挥部监理例会纪要")
    run.bold = True
    run.font.size = Pt(16)
    set_heading_space_after(title, 12)

    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    rows = [
        ("项目部", "（如：三跑道项目部）"),
        ("项目名称", "（如：三跑道扩建工程）"),
        ("召开日期", "YYYY-MM-DD"),
        ("项目经理/负责人参会", "姓名/职务，多人用顿号分隔"),
        ("项目部长/副部长参会", "姓名/职务，多人用顿号分隔"),
    ]
    for idx, (label, placeholder) in enumerate(rows):
        info.rows[idx].cells[0].text = label
        info.rows[idx].cells[1].text = placeholder
    format_table_rows(info)

    doc.add_paragraph()

    add_heading(doc, "一、上次会议决议落实情况", level=2)
    doc.add_paragraph("（填写上次例会提出问题的闭环情况）")

    add_heading(doc, "二、本周施工及安全质量情况", level=2)
    doc.add_paragraph("（填写本周主要施工内容、危大工程/危险作业开展情况、安全质量管理措施等）")

    add_heading(doc, "三、存在问题及整改要求", level=2)
    issue_table = doc.add_table(rows=2, cols=5)
    issue_table.style = "Table Grid"
    headers = ["序号", "类型（安全/质量）", "问题描述", "整改责任人", "整改期限"]
    for col, header in enumerate(headers):
        issue_table.rows[0].cells[col].text = header
    issue_table.rows[1].cells[0].text = "1"
    issue_table.rows[1].cells[1].text = "安全"
    issue_table.rows[1].cells[2].text = "（示例）临边防护不到位"
    issue_table.rows[1].cells[3].text = "张三"
    issue_table.rows[1].cells[4].text = "YYYY-MM-DD"
    format_table_rows(issue_table)
    doc.add_paragraph("（可按实际情况增删行；系统支持从纪要中识别安全/质量问题并归档隐患）")

    add_heading(doc, "四、下步工作安排", level=2)
    doc.add_paragraph("（填写下周重点施工计划、协调事项及管控要求）")

    add_heading(doc, "五、其他事项", level=2)
    doc.add_paragraph("（选填）")

    note = doc.add_paragraph()
    note.add_run("填报说明：").bold = True
    note.add_run(
        "1. 填写完成后导出 Word 与 PDF 各一份上传系统，文字须清晰可辨；"
        "2. 同步上传签到表照片、会议现场照片（能看清主要参会人员，含水印）；"
        "3. 未召开会议须在系统备注中注明原因。"
    )

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_FILE)
    print(f"Wrote {OUT_FILE}")


if __name__ == "__main__":
    main()
