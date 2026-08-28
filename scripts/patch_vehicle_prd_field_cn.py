# -*- coding: utf-8 -*-
"""为车辆管理_PRD.docx 字段表补充「中文名」列（仅此改动）。"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx"
)

# 字段名 → 中文名（优先；说明仅作兜底）
FIELD_CN = {
    "todayInCount": "今日进场车次",
    "todayOutCount": "今日出场车次",
    "onSiteCount": "在场车辆",
    "projectName": "项目名称",
    "projectId": "项目ID",
    "projectCode": "项目编码",
    "keyword": "关键词",
    "url": "系统地址",
    "systemName": "系统名称",
    "enabled": "启用",
    "updatedAt": "更新时间",
    "direction": "方向",
    "id": "记录ID",
    "plateNo": "车牌号",
    "vehicleType": "车辆类型",
    "gateName": "道闸名称",
    "recordTime": "记录时间",
    "externalId": "外部唯一号",
    "操作": "操作",
}


def set_cell_text(cell, text):
    # 清空并写入，尽量保留第一个段落
    if not cell.paragraphs:
        cell.text = text
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def infer_cn(field_name: str, remark: str) -> str:
    if field_name in FIELD_CN:
        return FIELD_CN[field_name]
    # 操作行
    if field_name.strip() == "操作":
        return "操作"
    # 说明较短且不像长句时可用
    remark = (remark or "").strip()
    if not remark:
        return field_name
    # 取说明中分号/句号前的短标题
    for sep in ("；", ";", "。", "（", "("):
        if sep in remark:
            head = remark.split(sep, 1)[0].strip()
            if 1 <= len(head) <= 20 and not head.startswith("`"):
                return head
    if len(remark) <= 16 and "\n" not in remark:
        return remark
    return field_name


def insert_col_after(table, after_idx: int):
    """在每行 after_idx 列之后插入一列（复制单元格结构）。"""
    for row in table.rows:
        src_tc = row.cells[after_idx]._tc
        new_tc = deepcopy(src_tc)
        # 清空新单元格文本
        for t in new_tc.iterchildren(qn("w:p")):
            for r in t.findall(".//" + qn("w:t")):
                r.text = ""
        src_tc.addnext(new_tc)
    # 刷新：python-docx cells 缓存可能过期，重新取 table
    return table


def process(doc: Document) -> int:
    changed = 0
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        if "字段名" not in headers:
            continue
        if "中文名" in headers:
            continue
        field_idx = headers.index("字段名")
        remark_idx = headers.index("说明") if "说明" in headers else -1

        insert_col_after(table, field_idx)
        # 插入后表头：… 字段名 | (新) | 类型 …
        # 重新读 cells
        set_cell_text(table.rows[0].cells[field_idx + 1], "中文名")

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            field_name = row.cells[field_idx].text.strip()
            remark = ""
            if remark_idx >= 0:
                # 说明列因插入右移一列
                remark = row.cells[remark_idx + 1].text.strip()
            cn = infer_cn(field_name, remark)
            if field_name == "updatedAt":
                cn = "更新时间"
            if field_name == "enabled":
                cn = "启用"
            if field_name == "url":
                cn = "系统地址"
            if field_name == "projectId" and "每项目" in remark:
                cn = "配置项目"
            set_cell_text(row.cells[field_idx + 1], cn)
            changed += 1
        print("patched table headers ->", [c.text.strip() for c in table.rows[0].cells])
    return changed


def main():
    doc = Document(str(SRC))
    n = process(doc)
    doc.save(str(SRC))
    print("saved", SRC.name, "rows filled:", n)


if __name__ == "__main__":
    main()
