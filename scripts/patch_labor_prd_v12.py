# -*- coding: utf-8 -*-
"""Round-2 PRD patch: from V1.1 base -> save to 副本_20260812.docx as V1.2."""
from docx import Document
from pathlib import Path

SRC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_V1.1.docx"
)
DST = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_副本_20260812.docx"
)
ALSO = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_V1.2.docx"
)


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def set_cell_text(cell, value):
    if not cell.paragraphs:
        cell.text = value
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def apply_all(doc, pairs):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, pairs)
    return n


def main():
    doc = Document(str(SRC))

    pairs = [
        # APP only view; Web dispose = default recipient only
        (
            "任务类预警可在个人中心（Web/APP）或项目 Web「预警清单/详情」处置；通知类仅查阅。口径与 F-LABOR-06 一致。",
            "任务类预警：Web 个人中心（仅默认接收人）或项目 Web「预警清单/详情」可处置；APP 个人中心仅查看、不做处置；通知类仅查阅。口径与 F-LABOR-06 一致。",
        ),
        (
            "展示人员预警及其处置记录；任务类支持默认接收人在个人中心处置并关闭；消息投递遵循个人中心公共能力。",
            "展示人员预警及其处置记录；Web 端任务类仅默认接收人可在个人中心处置并关闭；APP 端仅查看；消息投递遵循个人中心公共能力。",
        ),
        (
            "三端个人中心可查看预警消息与详情，任务类可由默认接收人处置；亦可在项目 Web 预警清单处置",
            "三端个人中心可查看预警消息与详情；任务类仅默认接收人可在 Web 个人中心或项目 Web 预警清单处置；APP 仅查看不做处置",
        ),
        (
            "消息通知与详情；任务类可处置",
            "消息通知与详情；Web 任务类仅默认接收人可处置，APP 仅查看",
        ),
        (
            "个人中心与项目 Web 均可处置（任务类）",
            "Web 个人中心（仅默认接收人）与项目 Web 可处置；APP 仅查看",
        ),
        (
            "人员预警消息通知与详情；任务类可处置",
            "人员预警消息通知与详情；Web 任务类仅默认接收人可处置，APP 仅查看",
        ),
        (
            "支持预警消息与详情处置",
            "支持预警消息与详情；Web 端任务类仅默认接收人可处置，APP 仅查看",
        ),
        (
            "支持在个人中心处置预警",
            "Web 个人中心支持默认接收人处置；APP 仅查看",
        ),
        (
            "APP/指挥部 Web 仅经个人中心查看消息与详情（F-LABOR-10）。",
            "指挥部 Web / APP 经个人中心触达预警（F-LABOR-10）：Web 端任务类仅默认接收人可处置，APP 仅查看；项目侧亦可在预警清单处置。",
        ),
        (
            "自动处理:按下方列表弹窗提示对应内容，取消按钮关闭弹窗，去处理按钮跳转人员实名制，并自动筛选对应人员",
            "自动处理：Web 端按下方列表弹窗提示对应内容，可「去处理」跳转人员实名制并自动筛选对应人员；APP 端仅展示提示文案，无「去处理」跳转、不做处置",
        ),
        (
            "手动处理:弹窗展示处理说明*，附件字段，上传后自动关闭预警",
            "手动处理：仅 Web 端默认接收人可操作——弹窗展示处理说明*、附件，提交后关闭预警；APP 端无此操作",
        ),
        (
            "仅配置的默认接收人可处置，其他分级管控用户仅接收通知",
            "处置权限：仅配置的默认接收人可处置；其他分级管控用户仅接收通知、只读查看；APP 端一律只读",
        ),
        (
            "查看消息/详情/处置",
            "查看消息/详情；Web 处置（仅默认接收人）",
        ),
        (
            "均可看到对应预警消息并可打开详情",
            "均可看到对应预警消息并可打开详情；APP 仅查看，不可处置",
        ),
        (
            "打开个人中心任务详情",
            "Web 个人中心任务详情且当前用户为默认接收人",
        ),
        (
            "提交处置信息后关闭代办事项并更新预警状态",
            "提交处置信息后关闭代办事项并更新预警状态；APP 无处置入口",
        ),
        (
            "根据提示处理人员问题，如不足16周岁人员退场（改表，并触发定时任务）",
            "Web 端根据提示前往人员实名制/闸机侧处理（如不足16周岁退场），次日定时任务关闭；APP 仅查看提示、不跳转不处置",
        ),
        (
            "只读展示编号、类型、状态、原因、人员、项目、时间线",
            "展示编号、类型、状态、原因、人员、项目、时间线；Web 默认接收人对任务类可处置，APP/非默认接收人只读",
        ),
        (
            "项目侧可处置角色",
            "项目 Web 可处置角色，或 Web 个人中心默认接收人；APP 不可处置",
        ),
        # HQ: no drill-down
        (
            "展示全项目人员汇总指标与按项目明细，支持关键词筛选、预警未处置下钻、查看项目详情（切项目进台账/看板上下文）。",
            "展示全项目人员汇总指标与按项目明细，支持关键词筛选、查看项目详情（切项目进台账/看板上下文）；「预警未处置」仅展示数量，不做单独下钻。",
        ),
        (
            "全项目汇总指标、按项目明细、预警未处置下钻",
            "全项目汇总指标、按项目明细；预警未处置仅展示不下钻",
        ),
        ("预警未处置下钻", "预警未处置（仅展示）"),
        (
            "点击「预警未处置数量」等指标",
            "查看「预警未处置数量」等指标",
        ),
        (
            "进入预警相关列表（按待处理过滤）",
            "不做下钻跳转；指标仅展示",
        ),
        # reportDays unify
        (
            "分级链最多 8 级，上报天数逐级递增（配置见 F-LABOR-07 实名制配置）",
            "分级链最多 8 级，上报天数逐级递增（后级 ≥ 前级，允许相等；配置见 F-LABOR-07）",
        ),
        (
            "校验 ≤8 级、天数逐级递增",
            "校验 ≤8 级、天数逐级递增（后级 ≥ 前级，允许相等）",
        ),
        ("各级天数递增提示", "不满足后级≥前级时提示"),
        ("上报天数不递增", "上报天数不满足后级≥前级"),
        # elderly: remove 仅施工方端
        (
            "通知类；仅施工方端提示，需做好工人健康情况排查，不强制要求退场。",
            "通知类；推送个人中心消息，需做好工人健康情况排查，不强制要求退场。",
        ),
        (
            "向施工方端完成提示推送并留档后。",
            "完成个人中心提示推送并留档后。",
        ),
        # blacklist wording (avoid bare「进场预警」substring replace)
        (
            "全平台共用新增/移出；进场生成预警",
            "全平台共用新增/移出；日扫在岗命中生成预警",
        ),
        (
            "黑名单进场生成 “黑名单进场预警”，由责任人处置",
            "黑名单日扫在岗命中生成「黑名单人员进场预警」（系统自动关闭类）；默认接收人可在 Web 查看/引导处理",
        ),
        (
            "将人员移出黑名单后，不再对新进场触发；已生成预警自动关闭。",
            "将人员移出黑名单或人员离场后，日扫不再对新命中触发；已生成预警按自动关闭规则关闭。",
        ),
        ("进场同步完成", "日扫校验完成"),
        # dashboard
        (
            "按组织范围加载指标、环图、趋势与预警清单",
            "按当前项目加载指标、环图、趋势与预警清单",
        ),
        ("项目端列表展示；指挥部侧可省略。", "本项目列表展示。"),
        ("项目端可展示时分；指挥部侧可省略。", "可展示时分。"),
        (
            "在岗且为特种作业人员为是的人数。",
            "在岗且 isSpecial（是否特种作业）为是的人数。",
        ),
        ("已选组织范围且有人员", "已选当前项目且有人员"),
    ]

    print("hits:", apply_all(doc, pairs))

    # Fix AC-01 name cell if exact match "进场预警"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t == "进场预警":
                    set_cell_text(cell, "日扫命中预警")
                    print("AC title cell fixed")
                if t == "后级 ≥ 前级":
                    set_cell_text(cell, "后级 ≥ 前级（即逐级递增口径，允许相等）")
                    print("reportDays cell fixed")
                if t == "手动且待处理":
                    # leave; companion cell already patched via 项目侧可处置角色
                    pass

    for table in doc.tables[:6]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            if cells[0] == "版本号":
                set_cell_text(row.cells[1], "V1.2")
                print("version V1.2")
            if cells[0] == "更新日期":
                set_cell_text(row.cells[1], "2026-08-12")
            if cells[0] == "状态":
                set_cell_text(row.cells[1], "待评审")

    for table in doc.tables[:10]:
        headers = [c.text.strip() for c in table.rows[0].cells]
        if not headers or headers[0] != "版本":
            continue
        if "说明" not in headers and "日期" not in headers:
            continue
        existing = [r.cells[0].text.strip() for r in table.rows]
        if "V1.2" not in existing:
            new_row = table.add_row()
            mapping = {
                "版本": "V1.2",
                "修订人": "",
                "日期": "2026-08-12",
                "说明": (
                    "评审修订：APP个人中心仅查看不做处置；处置仅默认接收人（Web）；"
                    "指挥部预警未处置取消下钻；上报天数递增=后级≥前级；去掉高龄仅施工方端；"
                    "黑名单/看板残留措辞清理"
                ),
            }
            for i, h in enumerate(headers):
                if i < len(new_row.cells):
                    set_cell_text(new_row.cells[i], mapping.get(h, ""))
            print("revision V1.2 added")
        break

    for path in (DST, ALSO):
        try:
            doc.save(str(path))
            print("saved", path)
        except PermissionError:
            print("LOCKED", path)


if __name__ == "__main__":
    main()
