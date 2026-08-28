# -*- coding: utf-8 -*-
"""从可打开备份恢复车辆 PRD，仅做文本替换与修订行写入（不插表行，避免损坏 OOXML）。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

BASE = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
CUR = BASE / "车辆管理_PRD.docx"
BAK = BASE / "车辆管理_PRD_菜单挂载修订前_副本_20260813.docx"
NOTE = "指挥部菜单挂载由「施工现场管理」调整为「智慧工地监管」；页面路由不变。"
VER = "V2.2.2"
TODAY = "2026-08-13"

REPLACEMENTS = [
    ("挂施工现场管理 / 智慧工地监管", "挂智慧工地监管"),
    ("挂施工现场管理/智慧工地监管", "挂智慧工地监管"),
    ("施工现场管理 → 车辆管理", "智慧工地监管 → 车辆管理"),
    ("施工现场管理 · 车辆管理", "智慧工地监管 · 车辆管理"),
    ("挂施工现场管理", "挂智慧工地监管"),
    ("位于施工现场管理，不在安全看板", "位于智慧工地监管，不在安全看板"),
    ("位于施工现场管理，", "位于智慧工地监管，"),
    ("指挥部挂施工现场管理", "指挥部挂智慧工地监管"),
    ("指挥部「施工现场管理 → 车辆管理」", "指挥部「智慧工地监管 → 车辆管理」"),
    ("指挥部菜单**挂「施工现场管理」**", "指挥部菜单**挂「智慧工地监管」**"),
    ("「施工现场管理」", "「智慧工地监管」"),
]


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def apply_replacements(text: str) -> str:
    new = text
    for old, repl in REPLACEMENTS:
        new = new.replace(old, repl)
    return new


def replace_in_paragraph(para) -> int:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    new = apply_replacements(full)
    if new == full:
        return 0
    if para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new
    return 1


def main() -> None:
    if not BAK.exists():
        raise SystemExit(f"missing backup: {BAK}")

    broken = BASE / f"车辆管理_PRD_损坏不可打开_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    if CUR.exists():
        shutil.copy2(CUR, broken)
        print("quarantined broken ->", broken.name)

    shutil.copy2(BAK, CUR)
    print("restored from backup")

    doc = Document(str(CUR))
    n = 0
    for para in doc.paragraphs:
        n += replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    n += replace_in_paragraph(para)

    # meta
    meta = doc.tables[0]
    for row in meta.rows:
        key = row.cells[0].text.strip()
        if key == "版本号" and len(row.cells) > 1:
            set_cell_text(row.cells[1], VER)
        if key == "更新日期" and len(row.cells) > 1:
            set_cell_text(row.cells[1], TODAY)

    # revision: 版本|日期|说明 — shift without structural insert
    rev = doc.tables[1]
    header = [c.text.strip() for c in rev.rows[0].cells]
    assert header[0] in ("版本", "Version"), header

    # Collect existing data rows (non-empty version)
    data = []
    for row in rev.rows[1:]:
        ver = row.cells[0].text.strip()
        if not ver or ver == "版本":
            continue
        vals = [c.text for c in row.cells]
        data.append(vals)

    # Prepend new revision into existing row slots only (reuse empty rows)
    new_row = [VER, TODAY, NOTE]
    while len(new_row) < len(rev.rows[1].cells):
        new_row.append("")
    all_data = [new_row] + data

    for i, vals in enumerate(all_data):
        row_i = i + 1
        if row_i >= len(rev.rows):
            # no more physical rows — stop (do not insert XML rows)
            print("WARN: not enough revision rows to write all history; wrote", i, "of", len(all_data))
            break
        fill = list(vals)
        while len(fill) < len(rev.rows[row_i].cells):
            fill.append("")
        for j, v in enumerate(fill[: len(rev.rows[row_i].cells)]):
            set_cell_text(rev.rows[row_i].cells[j], v.strip() if isinstance(v, str) else v)

    # clear leftover rows beyond written
    written = min(len(all_data), len(rev.rows) - 1)
    for row_i in range(written + 1, len(rev.rows)):
        for c in rev.rows[row_i].cells:
            set_cell_text(c, "")

    doc.save(str(CUR))
    print("saved replacements", n)

    # leftover check
    doc2 = Document(str(CUR))
    left = []
    for p in doc2.paragraphs:
        if "施工现场管理" in p.text and "调整为" not in p.text:
            left.append(p.text[:100])
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                if "施工现场管理" in cell.text and "调整为" not in cell.text:
                    left.append(cell.text[:100])
    print("leftover mount refs", len(left))
    for x in left:
        print(" -", x)
    meta2 = {r.cells[0].text.strip(): r.cells[1].text.strip() for r in doc2.tables[0].rows if len(r.cells) > 1}
    print("meta", meta2.get("版本号"), meta2.get("更新日期"))
    print("rev1", [c.text.strip()[:50] for c in doc2.tables[1].rows[1].cells])


if __name__ == "__main__":
    main()
