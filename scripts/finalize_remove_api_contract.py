# -*- coding: utf-8 -*-
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

SRC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)


def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = text
        for extra in cell.paragraphs[1:]:
            for r in extra.runs:
                r.text = ""
    else:
        cell.text = text


def main() -> None:
    d = Document(str(SRC))

    # Ensure no API contract leftovers
    for table in list(d.tables):
        if not table.rows:
            continue
        header = "|".join(c.text.strip() for c in table.rows[0].cells)
        if "接口用途" in header and "业务入参" in header:
            table._element.getparent().remove(table._element)
            print("removed leftover API table")

    hits = []
    for p in d.paragraphs:
        if "业务契约" in p.text or "API 业务" in p.text:
            hits.append(("p", p.text[:80]))
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if "业务契约" in cell.text or ("API" in cell.text and "契约" in cell.text):
                    hits.append(("c", cell.text[:80].replace("\n", " ")))
    print("hits before rev insert", hits)

    # Meta version/date
    meta = d.tables[0]
    for row in meta.rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], "V2.1.2")
        if key == "更新日期":
            set_cell_text(row.cells[1], "2026-08-13")

    # Revision table
    rev = d.tables[1]
    headers = [c.text.strip() for c in rev.rows[0].cells]
    assert "版本" in headers[0]
    first = rev.rows[1].cells[0].text.strip() if len(rev.rows) > 1 else ""
    if first != "V2.1.2":
        new_tr = deepcopy(rev.rows[1]._tr)
        rev._tbl.insert(1, new_tr)
        # re-bind
        row = rev.rows[1]
        # columns: 版本, 修订人?, 日期, 说明 — from earlier: ['版本','修订人','日期','说明']
        set_cell_text(row.cells[0], "V2.1.2")
        if len(row.cells) > 1:
            set_cell_text(row.cells[1], row.cells[1].text.strip() or "—")
        if len(row.cells) > 2:
            set_cell_text(row.cells[2], "2026-08-13")
        if len(row.cells) > 3:
            set_cell_text(
                row.cells[3],
                "删除 6.2.9「API 业务契约」空壳标题及接口用途表；6.2 节后续小节编号顺延（业务规则→6.2.9）。",
            )
        print("inserted revision V2.1.2")
    else:
        print("revision V2.1.2 already present")

    d.save(str(SRC))

    d2 = Document(str(SRC))
    print("version", d2.tables[0].rows[1].cells[0].text, d2.tables[0].rows[1].cells[1].text)
    for row in d2.tables[0].rows:
        if row.cells[0].text.strip() == "版本号":
            print("meta version", row.cells[1].text.strip())
    print("rev row1", [c.text.strip()[:40] for c in d2.tables[1].rows[1].cells])
    print("6.2 H4:")
    for p in d2.paragraphs:
        if p.style and p.style.name == "Heading 4" and p.text.startswith("6.2."):
            print(" ", p.text)
    left = []
    for p in d2.paragraphs:
        if "业务契约" in p.text:
            left.append(p.text)
    for t in d2.tables:
        for row in t.rows:
            for cell in row.cells:
                if "业务契约" in cell.text:
                    left.append(cell.text[:60])
    print("remaining 业务契约", left)


if __name__ == "__main__":
    main()
