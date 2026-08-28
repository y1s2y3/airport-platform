# -*- coding: utf-8 -*-
"""车辆 PRD：批量上报失败=整批拒绝；注明后端敲定合理性。"""
from __future__ import annotations

import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\车辆管理_PRD.docx")
LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_vehicle_batch_log.txt"
)
VER = "V2.2.7"
TODAY = "2026-08-13"
NOTE = (
    "批量上报失败策略定为整批拒绝（任一条校验失败则整批不落库）；"
    "是否合理由后端结合接口实现与对接方约束最终敲定。"
)
BATCH_RULE = (
    "批量失败策略：任一条校验失败（含缺必填、项目 ID 无效/越权、方向非法、缺 external_id 等）"
    "则整批拒绝，本批全部不落库，并返回失败明细；不采用部分成功。"
    "（产品暂定；是否合理由后端结合接口实现与对接方约束最终敲定。）"
)


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def bump_meta(doc: Document) -> None:
    for row in doc.tables[0].rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], VER)
        if key == "更新日期":
            set_cell_text(row.cells[1], TODAY)


def add_bullet_under_669(doc: Document, log: list[str]) -> None:
    if any("整批拒绝" in (p.text or "") for p in doc.paragraphs):
        log.append("batch rule already in paras")
        return
    target = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("6.6.9") and "业务规则" in para.text:
            target = i
            break
    if target is None:
        log.append("WARN no 6.6.9")
        return
    last = None
    for j in range(target + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            break
        if p.text.strip():
            last = p
    if last is None:
        log.append("WARN no bullets")
        return
    new_el = deepcopy(last._element)
    for node in new_el.iter():
        if node.tag.endswith("}t"):
            node.text = ""
    last._element.addnext(new_el)
    first = True
    for node in new_el.iter():
        if node.tag.endswith("}t"):
            if first:
                node.text = BATCH_RULE
                first = False
            else:
                node.text = ""
    log.append("added 6.6.9 batch rule")


def patch_interaction_table(doc: Document, log: list[str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            op = row.cells[0].text.strip()
            if op == "单条/批量上报" or (op.startswith("单条") and "批量" in op):
                # 异常处理列
                if len(row.cells) >= 4:
                    set_cell_text(
                        row.cells[3],
                        "缺必填/项目无法识别或越权/方向非法/缺 external_id → 拒绝；"
                        "批量时任一条失败则整批拒绝（全部不落库），返回失败明细；记对接日志。"
                        "整批策略是否合理由后端最终敲定。",
                    )
                    log.append("updated 单条/批量上报 异常处理")
            if "重复上报" in op:
                pass


def ensure_ac(doc: Document, log: list[str]) -> None:
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        h = " | ".join(c.text for c in table.rows[0].cells)
        if "验收项" not in h:
            continue
        blob = "\n".join("|".join(c.text for c in r.cells) for r in table.rows)
        if "越权项目拒收" not in blob and "无 external_id 拒收" not in blob:
            continue
        # this is 6.6 AC
        for row in table.rows[1:]:
            if "整批" in row.cells[1].text or row.cells[0].text.strip() == "AC-07":
                log.append("AC batch exists")
                return
        new_tr = deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        vals = [
            "AC-07",
            "批量整批拒绝",
            "一批含合法条与非法条（如缺车牌）",
            "调用批量上报",
            "整批不落库；合法条也不写入；返回失败明细",
        ]
        for i, v in enumerate(vals):
            if i < len(row.cells):
                set_cell_text(row.cells[i], v)
        log.append(f"appended AC-07 on T{ti}")
        return
    log.append("WARN 6.6 AC not found")


def main() -> None:
    log: list[str] = []
    bak = SRC.with_name(
        f"车辆管理_PRD_批量策略前_副本_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    shutil.copy2(SRC, bak)
    log.append("backup " + bak.name)

    doc = Document(str(SRC))
    bump_meta(doc)
    add_bullet_under_669(doc, log)
    patch_interaction_table(doc, log)
    ensure_ac(doc, log)

    # 7.1 失败策略 briefly
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 5:
                continue
            name = row.cells[0].text.strip()
            if "标准上报接口" in name and "F-VEHICLE-06" in name:
                fail = row.cells[4].text.strip()
                if "整批" not in fail:
                    set_cell_text(
                        row.cells[4],
                        fail.rstrip("；。")
                        + "；批量任一条失败则整批拒绝（是否合理由后端最终敲定）",
                    )
                    log.append("7.1 失败策略补整批")

    doc.save(str(SRC))
    log.append("python saved")

    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        wdoc = word.Documents.Open(str(SRC), ReadOnly=False)
        tbl = wdoc.Tables(2)
        first = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first != VER:
            tbl.Rows(2).Select()
            word.Selection.InsertRowsAbove(1)
            row = tbl.Rows(2)
            ncols = row.Cells.Count
            vals = [VER, TODAY, NOTE] if ncols == 3 else [VER, "—", TODAY, NOTE]
            for i, v in enumerate(vals, start=1):
                if i <= ncols:
                    row.Cells(i).Range.Text = v
            log.append("rev " + VER)

        tmp = SRC.with_name("_vehicle_batch_tmp.docx")
        if tmp.exists():
            tmp.unlink()
        wdoc.SaveAs2(str(tmp), FileFormat=12)
        wdoc.Close(False)
        shutil.copy2(tmp, SRC)
        tmp.unlink()

        wdoc = word.Documents.Open(str(SRC), ReadOnly=True)
        log.append(f"WORD_OK {wdoc.Paragraphs.Count}")
        wdoc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    d2 = Document(str(SRC))
    blob = "\n".join(p.text for p in d2.paragraphs)
    for t in d2.tables:
        for row in t.rows:
            blob += "\n" + "|".join(c.text for c in row.cells)
    for n in ["整批拒绝", "后端", "AC-07", "批量整批拒绝", "最终敲定"]:
        log.append(f"has[{n}]={n in blob}")
    LOG.write_text("\n".join(log), encoding="utf-8")
    print("done")


if __name__ == "__main__":
    main()
