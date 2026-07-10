# -*- coding: utf-8 -*-
"""更新《智慧工地建设管控一体化平台需求规格说明书》.docx"""
from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
DOC_FILE = ROOT / "《智慧工地建设管控一体化平台需求规格说明书》.docx"
TODAY = date(2026, 7, 10).strftime("%Y年%m月%d日")
PROJECT_NAME = "深圳机场智慧工地建设管控一体化平台"

# 标题自动编号与标题文字统一：黑体 + 与各级标题相同字号
HEADING_FONT_SPECS = {
    1: {"name": "黑体", "size_pt": 14, "bold": False},
    2: {"name": "黑体", "size_pt": 12, "bold": False},
    3: {"name": "黑体", "size_pt": 12, "bold": False},
    4: {"name": "黑体", "size_pt": 12, "bold": False},
}
HEADING_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)[、.\s]+\s*")


def strip_manual_heading_number(text: str) -> str:
    t = text.strip()
    while True:
        m = HEADING_NUM_RE.match(t)
        if not m:
            break
        t = t[m.end() :].lstrip()
    return t


def is_heading_paragraph(paragraph) -> bool:
    name = paragraph.style.name if paragraph.style else ""
    return name.startswith("Heading") or name.startswith("标题")


def fix_all_heading_numbers(doc) -> int:
    """Remove manual numbering from heading paragraphs to avoid conflict with Word auto-numbering."""
    count = 0
    for p in doc.paragraphs:
        if not is_heading_paragraph(p):
            continue
        raw = p.text.strip()
        if not raw:
            continue
        cleaned = strip_manual_heading_number(raw)
        if cleaned == raw:
            continue
        if p.runs:
            p.runs[0].text = cleaned
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = cleaned
        count += 1
    return count


def _heading_level_from_paragraph(paragraph) -> int | None:
    name = paragraph.style.name if paragraph.style else ""
    if name.startswith("Heading "):
        try:
            return int(name.split(" ", 1)[1])
        except ValueError:
            return None
    if name.startswith("标题 "):
        try:
            return int(name.split(" ", 1)[1])
        except ValueError:
            return None
    return None


def _half_points(size_pt: float) -> str:
    return str(int(round(size_pt * 2)))


def _build_rpr_element(font_name: str, size_pt: float, bold: bool = False):
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rpr.append(rfonts)
    sz_val = _half_points(size_pt)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), sz_val)
    rpr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), sz_val)
    rpr.append(szcs)
    if bold:
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))
    return rpr


def _replace_child_rpr(parent, rpr):
    existing = parent.find(qn("w:rPr"))
    if existing is not None:
        parent.remove(existing)
    parent.append(rpr)


def apply_heading_run_font(run, level: int):
    spec = HEADING_FONT_SPECS.get(level, HEADING_FONT_SPECS[4])
    add_run_font(run, spec["name"], spec["size_pt"], spec["bold"])


def sync_heading_style_fonts(doc) -> int:
    count = 0
    for level, spec in HEADING_FONT_SPECS.items():
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = spec["name"]
        style.font.size = Pt(spec["size_pt"])
        style.font.bold = spec["bold"]
        _replace_child_rpr(style.element, _build_rpr_element(spec["name"], spec["size_pt"], spec["bold"]))
        count += 1
    return count


def _resolve_heading_abstract_num_id(doc) -> str | None:
    try:
        h1 = doc.styles["Heading 1"].element
    except KeyError:
        return None
    ppr = h1.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    num_id_el = numpr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    num_id = num_id_el.get(qn("w:val"))
    numbering_root = doc.part.numbering_part.element
    for num in numbering_root.findall(qn("w:num")):
        if num.get(qn("w:numId")) != num_id:
            continue
        abs_el = num.find(qn("w:abstractNumId"))
        if abs_el is not None:
            return abs_el.get(qn("w:val"))
    return None


def sync_heading_numbering_fonts(doc) -> int:
    """Make auto-number prefix use the same font/size as each heading level."""
    abstract_id = _resolve_heading_abstract_num_id(doc)
    if not abstract_id:
        return 0
    numbering_root = doc.part.numbering_part.element
    count = 0
    for abstract in numbering_root.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in abstract.findall(qn("w:lvl")):
            ilvl = int(lvl.get(qn("w:ilvl"), "99"))
            if ilvl >= 4:
                continue
            spec = HEADING_FONT_SPECS.get(ilvl + 1, HEADING_FONT_SPECS[4])
            _replace_child_rpr(lvl, _build_rpr_element(spec["name"], spec["size_pt"], spec["bold"]))
            count += 1
    return count


def sync_heading_paragraph_fonts(doc) -> int:
    count = 0
    for para in doc.paragraphs:
        level = _heading_level_from_paragraph(para)
        if level is None:
            continue
        for run in para.runs:
            apply_heading_run_font(run, level)
        count += 1
    return count


def sync_all_heading_fonts(doc) -> dict[str, int]:
    return {
        "styles": sync_heading_style_fonts(doc),
        "numbering_levels": sync_heading_numbering_fonts(doc),
        "paragraphs": sync_heading_paragraph_fonts(doc),
    }



def add_run_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    sz_val = _half_points(size)
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), sz_val)
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), sz_val)


def set_heading_space_after(paragraph, pt=6):
    paragraph.paragraph_format.space_after = Pt(pt)


def add_heading(doc, text, level):
    p = doc.add_heading(strip_manual_heading_number(text), level=level)
    set_heading_space_after(p)
    for run in p.runs:
        apply_heading_run_font(run, level)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    add_run_font(r)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("• " + text)
    add_run_font(r)
    p.paragraph_format.left_indent = Cm(0.74)
    return p


def format_table_rows(table, min_height_cm=0.6):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(min_height_cm)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val)
    format_table_rows(table)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
    doc.add_paragraph()
    return table


def build_coc_screen_structure():
    return [
        "COC调度大屏 → 通用：标题栏（工程指挥部项目树、状态筛选、关闭返回）/ 会议签到浮窗 / 会议录屏",
        "COC调度大屏 → 工程指挥部视图：视频监控（监控列表+3×2宫格+巡检对讲设备）/ 项目进度 / 红黑榜 / 劳务分析 / 隐患分析",
        "COC调度大屏 → 领导讲话：3×3视频墙、设备邀请、全员静音、截屏生成任务单/提示函/处罚单",
        "COC调度大屏 → 项目视图：项目维度监控、劳务统计、隐患清单、项目调度入口",
        "COC调度大屏 → 项目调度：对讲主画面+关联监控+危险作业清单+管理人员清单+人员风险核验+任务单/提示函/处罚单+隐患清单+质量验评风险项",
        "COC调度大屏 → 进度详情：控制性计划甘特图",
        "COC后台管理 → 每日施工作业 / 任务单 / 提示函 / 处罚单 / 黑红榜单 / 巡检仪管理 / 智能安全帽管理 / 监理会议管理",
    ]


def build_coc_function_rows():
    return [
        ("COC大屏", "通用", "项目树", "工程指挥部+项目切换", "树形选择组织/项目，支持在建/前期/历史状态筛选与名称搜索", "调研"),
        ("COC大屏", "通用", "会议签到", "签到浮窗", "展示关键岗位人员姓名、单位、岗位及签到/未签到状态", "调研"),
        ("COC大屏", "通用", "会议管控", "页面录屏", "开始/结束调度会议，浏览器录屏并本地保存，AI语音转写生成纪要（试点）", "调研"),
        ("COC大屏", "指挥部", "视频监控", "监控列表", "左侧展示项目摄像头列表，支持在线/离线统计、重点视频筛选、拖拽排序", "调研/原型"),
        ("COC大屏", "指挥部", "视频监控", "视频宫格", "3×2分页展示实时视频，支持全屏、云台、截图标记问题", "调研/原型"),
        ("COC大屏", "指挥部", "视频监控", "巡检对讲", "展示巡检仪/Web对讲设备，点击进入项目调度", "调研/原型"),
        ("COC大屏", "指挥部", "领导讲话", "视频墙", "3×3展示入会设备，支持邀请、静音、截屏一键生成单据", "调研"),
        ("COC大屏", "指挥部", "红黑榜", "榜单展示", "红榜/黑榜双列展示，含项目、现场图、描述，与后台联动刷新", "调研"),
        ("COC大屏", "指挥部", "项目进度", "进度总览", "展示项目总数、平均完成率、滞后占比及柱状对比，可下钻甘特图", "调研"),
        ("COC大屏", "指挥部", "劳务分析", "多项目统计", "汇总各项目在场人数、特种作业、管理人员等", "调研"),
        ("COC大屏", "指挥部", "隐患分析", "等级分布", "环形图展示待处理隐患等级分布，支持跳转项目", "调研"),
        ("COC大屏", "项目", "项目调度", "远程对讲", "对讲主画面+3×2关联监控，支持设备切换", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "危险作业清单", "展示当日危险作业类型、地点、状态，支持更多弹窗与详情", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "管理人员清单", "展示项目关键管理人员及在场状态", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "人员风险核验", "针对安全教育、特种作业、实名制等维度展示人员风险项", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "任务单/提示函/处罚单", "调度页快速起草、提交，同步至COC后台台账", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "隐患清单", "展示项目安全/质量隐患，支持筛选与更多弹窗", "调研/原型"),
        ("COC大屏", "项目", "项目调度", "质量验评风险项", "展示质量验评待关注项，支持更多弹窗", "调研/原型"),
        ("COC后台", "业务", "每日施工作业", "危险作业填报", "项目每日16:30前填报次日危险作业计划，生成统计", "调研"),
        ("COC后台", "业务", "任务单", "全流程管理", "创建、下发、签收、整改反馈、复核闭环，支持截图/会议一键生成", "调研"),
        ("COC后台", "业务", "提示函", "全流程管理", "提示函创建、下发、签收与闭环", "调研"),
        ("COC后台", "业务", "处罚单", "全流程管理", "开具、申诉、复核、缴纳凭证上传与归档，可联动红黑榜", "调研"),
        ("COC后台", "业务", "黑红榜单", "按期维护", "维护红榜/黑榜期次，保存后大屏自动刷新", "调研"),
        ("COC后台", "设备", "巡检仪管理", "设备注册", "巡检仪编号、名称、绑定项目/人员、在线状态", "调研/原型"),
        ("COC后台", "设备", "智能安全帽", "设备台账", "安全帽设备绑定、在线/定位状态", "调研/原型"),
        ("COC后台", "业务", "监理会议管理", "纪要上传", "上传监理例会纪要（Word/PDF），AI识别安全质量问题并归档隐患", "调研"),
    ]


def build_master_function_list():
    rows = [
        ("组织管理", "组织架构", "组织树维护", "维护企业组织树，管理组织成员、岗位与组织信息，支持角色授权与数据权限（指挥部/项目层级）", "PC"),
        ("组织管理", "用户管理", "用户台账", "管理系统用户账号、所属组织与角色绑定，支持启用/停用、密码重置", "PC"),
        ("组织管理", "角色管理", "角色配置", "定义角色并配置菜单、数据权限，实现 RBAC", "PC"),
        ("组织管理", "岗位管理", "岗位台账", "维护岗位编码、名称、级别、职责及关联角色", "PC"),
        ("组织管理", "菜单管理", "菜单树", "维护系统菜单树、路由与按钮级权限标识", "PC"),
        ("组织管理", "日志管理", "登录/操作日志", "记录登录与关键业务操作，支持检索导出", "PC"),
    ]
    for r in build_coc_function_rows():
        rows.append(("COC调度中心", r[0], r[2], r[4], "PC"))
    rows.extend([
        ("安全管理", "人员实名制管理", "（待展开）", "劳务看板、人员实名制、考勤、黑名单、预警等", "PC"),
        ("安全管理", "车辆管理", "（待展开）", "车辆看板、进出场、轨迹监管、车牌管理等", "PC"),
        ("安全管理", "视频监控", "（待展开）", "监控列表、NVR设备管理", "PC"),
        ("质量管理", "质量验评", "（待展开）", "质量报验、验评流程、省统表等", "PC"),
    ])
    return rows


def clear_body_from(doc, start_index):
    body = doc.element.body
    children = list(body)
    para_idx = -1
    remove_from = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para_idx += 1
            if para_idx == start_index:
                remove_from = i
                break
    if remove_from is not None:
        for child in children[remove_from:]:
            body.remove(child)


def remove_trailing_junk(doc):
    """Remove leftover header/footer duplicate blocks at document end."""
    junk_markers = (
        "基础管理-需求规格说明书",
        "需求设计输出物",
        "评审形式",
        "参与评审人员",
        "文档编制单位",
        "文档审核单位",
        "编制部门",
        "编制人员",
        "审核人员",
        "版本：V1.0",
        "审核部门",
        "审核人员",
        "日期",
    )
    while doc.paragraphs:
        last = doc.paragraphs[-1]
        text = last.text.strip()
        if not text:
            p = last._element
            p.getparent().remove(p)
            continue
        if any(text.startswith(m) for m in junk_markers):
            p = last._element
            p.getparent().remove(p)
            continue
        break


def append_intro_and_analysis(doc):
    add_heading(doc, "引言", level=1)
    add_heading(doc, "目的", level=2)
    add_para(
        doc,
        f"本文档为{PROJECT_NAME}需求规格说明文档，作为系统设计、开发实现与测试验收的依据。"
        "本文档涵盖组织管理、COC调度中心（含调度大屏与COC后台管理）、安全管理、质量管理等业务板块，"
        "其中组织管理复用现有平台能力，COC调度中心结合调研资料与当前原型系统逻辑详细描述，"
        "安全管理与质量管理本期仅建立章节标题，详细需求将在后续版本补充。",
    )
    add_heading(doc, "参考文件", level=2)
    refs = [
        "《COC系统升级_功能清单.xlsx》（COC调度中心需求方案）",
        "《COC系统升级_PRD.docx》（COC系统升级产品需求）",
        "《COC大屏需求确认.docx》（COC大屏功能范围确认）",
        "《COC调研_需求调研报告_20260609.docx》《COC需求沟通会_20260609.md》（调研过程资料）",
        "《建设工程指挥部施工作业统计表.xlsx》《危险作业填报导入表.xlsx》等线下业务资料",
        "Airport_Expansion_Project_Informatization 原型系统（组织管理、COC大屏、COC后台管理）",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    add_heading(doc, "术语与缩写解释", level=2)
    add_table(
        doc,
        ["名词", "说明"],
        [
            ("COC", "Construction Operation Center，工程调度指挥中心"),
            ("RBAC", "基于角色的访问控制"),
            ("NVR", "网络硬盘录像机，视频监控存储设备"),
            ("危险作业", "动火、高处、有限空间等需专项管控的施工作业"),
            ("任务单/提示函/处罚单", "远程调度产生的三类督办文书，在COC后台闭环管理"),
        ],
        [Cm(3), Cm(13)],
    )

    add_heading(doc, "项目综述", level=1)
    add_heading(doc, "项目目标", level=2)
    add_para(
        doc,
        "深圳机场扩建工程参建项目多、作业面分散，指挥部需实时掌握各项目现场视频、人员、危险作业及安全质量动态。"
        "建设智慧工地建设管控一体化平台，实现组织与权限统一、COC远程可视调度、安全质量数据汇聚与闭环处置，"
        "支撑指挥部综合研判与风险前置管控。",
    )
    add_heading(doc, "项目范围", level=2)
    scopes = [
        "组织管理：组织架构、用户管理、角色管理、岗位管理、菜单管理、日志管理（复用现有系统，本期不展开详细需求）。",
        "COC调度中心：COC调度大屏（工程指挥部视图、项目视图、领导讲话、项目调度、会议签到与录屏）及COC后台管理（每日施工作业、任务单、提示函、处罚单、黑红榜、巡检仪、智能安全帽、监理会议）。",
        "安全管理：人员实名制管理、车辆管理、视频监控等（本期仅建立章节标题）。",
        "质量管理：质量验评等功能（本期仅建立章节标题）。",
    ]
    for s in scopes:
        add_bullet(doc, s)
    add_heading(doc, "客户组织架构", level=2)
    add_para(
        doc,
        "系统面向深圳机场扩建工程指挥部及下属各参建项目。组织树包括：深圳机场集团/工程指挥部、"
        "各在建/前期/历史项目、外部单位、施工项目、其他组织等节点；支持指挥部层级与项目层级数据权限配置。",
    )
    add_heading(doc, "用户分析", level=2)
    add_table(
        doc,
        ["角色名称", "主要职责及场景", "主要使用功能"],
        [
            ("指挥长/指挥部领导", "远程调度、综合研判", "COC大屏、领导讲话、项目进度、红黑榜"),
            ("COC调度室", "日常监控、会议组织、设备管理", "COC大屏、巡检仪/安全帽管理、会议录屏"),
            ("安质部主管", "安全质量监管、单据下发", "COC后台任务单/提示函/处罚单、黑红榜"),
            ("项目经理", "项目日常管理、整改闭环", "项目调度、每日施工作业填报、单据签收"),
            ("系统管理员", "平台配置", "组织管理、菜单权限"),
        ],
        [Cm(3.5), Cm(6), Cm(6.5)],
    )

    add_heading(doc, "产品业务分析", level=1)
    add_heading(doc, "用户原始需求", level=2)
    add_table(
        doc,
        ["优先级", "板块", "原始需求描述"],
        [
            ("P0", "COC调度中心", "平台内集成音视频调度，替代外部QQ等工具；支持截图/语音一键生成任务单、提示函、处罚单"),
            ("P0", "COC调度中心", "危险作业平台化填报，指挥部总览今日危险作业数量及项目明细"),
            ("P0", "COC调度中心", "督办任务增加复核流程与成果附件必填；支持撤回与申诉"),
            ("P0", "COC调度中心", "监理例会纪要上传至系统，AI识别安全质量问题（试点）"),
            ("P1", "COC调度中心", "视频监控宫格布局优化、项目状态筛选、重点摄像头标记、在线/离线统计"),
            ("P1", "组织管理", "RBAC权限、组织树、数据权限（指挥部/项目层级）"),
            ("P2", "安全管理", "人员实名制、车辆管理等（后续版本展开）"),
            ("P2", "质量管理", "质量验评等（后续版本展开）"),
        ],
        [Cm(1.5), Cm(3), Cm(11.5)],
    )
    add_heading(doc, "业务流程分析", level=2)
    flows = [
        "组织管理：维护组织树与用户 → 配置角色菜单与数据权限 → 用户登录后按权限访问业务模块。",
        "COC调度：选择项目/设备 → 远程对讲+视频监控 → 截图/语音标记问题 → 生成任务单/提示函/处罚单 → 项目签收整改 → 复核闭环 → 纳入红黑榜。",
        "危险作业：项目每日16:30前填报次日计划 → 指挥部/项目调度页展示当日清单 → 高风险项远程核验。",
        "监理会议：上传例会纪要 → AI摘取安全质量问题 → 人工核对 → 归档隐患/任务跟踪。",
    ]
    for f in flows:
        add_bullet(doc, f)
    add_heading(doc, "产品架构图", level=2)
    add_para(
        doc,
        "平台采用 Web 一体化架构：工作台为统一入口。组织管理为底层支撑；COC调度大屏为指挥展示层，"
        "COC后台为调度业务管理后台；安全管理、质量管理为并列业务板块。"
        "各模块通过统一组织树、项目树及标准数据接口实现数据贯通。",
    )
    add_heading(doc, "功能清单", level=2)
    add_table(
        doc,
        ["应用板块", "一级功能", "二级功能", "功能描述", "交互端"],
        build_master_function_list(),
        [Cm(2.2), Cm(2.8), Cm(2.8), Cm(7.5), Cm(1.5)],
    )


def append_detail_design(doc):
    add_heading(doc, "详细设计方案", level=1)

    add_heading(doc, "组织管理", level=2)
    add_para(
        doc,
        "组织管理模块复用现有智慧工地建设管控一体化平台能力，不在本文档重复描述详细需求。"
        "功能范围包括：组织架构（组织树、成员、岗位、角色授权、数据权限配置）、用户管理、角色管理、"
        "岗位管理、菜单管理、日志管理（系统日志、登录日志、操作日志）。"
        "详细字段、用例及权限规则参见平台一期组织管理模块规格说明及当前系统原型（/settings/*）。",
    )

    add_heading(doc, "COC调度中心", level=2)
    add_heading(doc, "功能概述", level=3)
    add_para(
        doc,
        "COC调度中心是工程指挥部远程调度与综合管控的核心模块，包含 COC调度大屏 与 COC后台管理 两部分。"
        "COC调度大屏提供工程指挥部一屏统览（视频监控、项目进度、红黑榜、劳务/隐患分析）及项目级远程调度能力；"
        "COC后台管理负责任务单、提示函、处罚单、黑红榜、危险作业填报、设备台账及监理会议等业务数据的维护与闭环。",
    )
    add_heading(doc, "业务流程", level=3)
    for item in [
        "调度会议：COC调度室发起会议 → 开启录屏/语音转写 → 邀请项目巡检设备入会 → 领导讲话或问题标记 → 生成任务单/提示函/处罚单 → 项目签收整改 → 复核闭环 → 归档。",
        "视频巡检：选择项目 → 浏览监控宫格 → 发现隐患截图标记 → 选择单据类型 → 自动分派责任人跟进。",
        "危险作业：项目每日16:30前在「每日施工作业」填报次日计划 → 调度大屏/项目调度页展示当日清单 → 指挥部远程核验。",
        "红黑榜：后台维护榜单期次 → 保存触发大屏刷新；处罚单/隐患记录可勾选入榜。",
        "监理会议：上传监理例会纪要 → AI识别安全/质量段落 → 人工核对 → 形成隐患清单或督办任务。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "功能结构图", level=3)
    for item in build_coc_screen_structure():
        add_bullet(doc, item)
    add_heading(doc, "原型页面", level=3)
    add_para(doc, "原型系统路径：#/coc（COC调度大屏）；COC后台管理为平台菜单「COC后台管理」下各子菜单。")
    pages = [
        ("工程指挥部首页", "默认进入指挥部视角：左侧监控列表与巡检对讲，中部3×2视频监控，左下项目进度，右侧红黑榜、劳务分析、隐患分析。"),
        ("领导讲话", "3×3视频墙展示参会设备（含指挥部调度席首格），支持设备邀请、全员静音、截屏生成任务单/提示函/处罚单。"),
        ("项目视图", "选择具体项目后展示项目维度监控、劳务统计、隐患清单，右侧提供「项目调度」入口。"),
        ("项目调度页", "对讲主画面+3×2关联监控；下方危险作业清单、管理人员清单、人员风险核验；右侧任务单/提示函/处罚单、隐患清单、质量验评风险项；各清单支持半透明「更多」弹窗。"),
        ("会议签到", "大屏右侧浮窗，展示关键岗位人员签到状态，支持按项目筛选。"),
        ("COC后台-任务单/提示函/处罚单", "台账列表、创建、下发、签收、整改、复核、申诉全流程。"),
        ("COC后台-每日施工作业", "危险作业计划填报与统计导出。"),
        ("COC后台-黑红榜单", "按期维护红榜/黑榜，与大屏联动。"),
        ("COC后台-巡检仪/智能安全帽", "设备注册、绑定项目/人员、在线状态维护。"),
        ("COC后台-监理会议管理", "监理例会信息登记、纪要上传、隐患识别与清单管理。"),
    ]
    for title, desc in pages:
        add_heading(doc, title, level=4)
        add_para(doc, desc)
    add_heading(doc, "字段描述", level=3)
    add_table(
        doc,
        ["用户端", "功能模块", "子功能", "主要字段", "说明"],
        [
            (r[0], r[1], r[2], "见业务模板", r[4]) for r in build_coc_function_rows()
        ],
        [Cm(1.8), Cm(2.2), Cm(2.5), Cm(2.5), Cm(7)],
    )
    add_table(
        doc,
        ["字段", "控件", "是否必填", "说明"],
        [
            ("项目名称", "文本", "是", "关联项目树节点"),
            ("作业类型", "下拉", "是", "动火/高处/有限空间等"),
            ("作业地点", "文本", "是", "作业位置描述"),
            ("计划时间", "日期时间", "是", "危险作业计划时段"),
            ("单据类型", "单选", "是", "任务单/提示函/处罚单"),
            ("问题描述", "多行文本", "是", "截图/调度识别的问题描述"),
            ("责任单位", "下拉", "是", "项目/分包单位"),
            ("整改期限", "日期", "是", "要求完成日期"),
            ("成果附件", "附件", "是", "整改闭环必传（复核前）"),
            ("榜单类型", "单选", "是", "红榜/黑榜"),
            ("期次", "文本", "是", "如2026年第6期"),
        ],
        [Cm(3), Cm(2.5), Cm(2), Cm(8.5)],
    )
    add_heading(doc, "功能用例", level=3)
    add_table(
        doc,
        ["功能菜单", "操作名称", "前置条件", "预期结果"],
        [
            ("COC大屏", "切换项目", "具有大屏访问权限", "切换指挥部/项目视图并刷新关联数据"),
            ("COC大屏", "截图生成单据", "调度中选中视频画面", "打开截图标记弹窗，选择单据类型并提交至后台"),
            ("COC大屏", "领导讲话", "COC调度室权限", "打开3×3视频墙，控制设备入会/静音"),
            ("项目调度", "远程对讲", "选中巡检/对讲设备", "建立对讲连接并展示关联监控"),
            ("项目调度", "查看危险作业", "项目已填报作业计划", "列表展示当日作业，更多弹窗支持筛选详情"),
            ("COC后台-任务单", "创建下发", "安质部权限", "生成任务单并推送至项目责任人"),
            ("COC后台-任务单", "整改提交", "项目责任人", "上传成果附件，状态变为待复核"),
            ("COC后台-任务单", "复核通过", "安质部复核人", "任务闭环归档"),
            ("COC后台-处罚单", "申诉", "被处罚方", "发起申诉进入复核流程"),
            ("COC后台-每日施工作业", "填报", "项目施工权限", "保存次日危险作业计划"),
            ("COC后台-黑红榜", "发布", "安质部权限", "保存后大屏红黑榜自动刷新"),
        ],
        [Cm(3.5), Cm(3), Cm(3.5), Cm(6)],
    )
    add_heading(doc, "关联模块", level=3)
    for item in [
        "组织管理：COC大屏、COC后台菜单与按钮权限由角色管理统一配置；数据权限控制项目可见范围。",
        "项目基础信息：项目树、项目简称、状态（在建/前期/历史）来源于基础数据模块。",
        "视频监控：NVR设备与通道数据供大屏宫格播放及项目调度关联监控调用。",
        "人员实名制：劳务分析、人员风险核验、会议签到人员数据来源于实名制模块（后续版本深度集成）。",
        "质量管理：质量验评风险项与质量隐患数据联动（后续版本展开）。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "权限要求", level=3)
    add_table(
        doc,
        ["页面功能", "授权对象", "操作权限"],
        [
            ("COC调度大屏", "指挥部领导、COC调度室", "查看、调度、截图、会议控制"),
            ("领导讲话", "指挥长、COC调度室", "设备邀请、静音、截屏生成单据"),
            ("项目调度", "COC调度室、安质部、项目经理", "查看；调度室可发起对讲与截图"),
            ("COC后台-任务单/提示函/处罚单", "安质部", "创建、下发、复核"),
            ("COC后台-任务单/提示函/处罚单", "项目经理、施工", "签收、整改反馈、申诉"),
            ("黑红榜单", "安质部", "新增、编辑、删除、发布"),
            ("每日施工作业", "项目经理、施工", "填报、查看"),
            ("巡检仪/智能安全帽", "COC调度室", "注册、绑定、维护"),
            ("监理会议管理", "COC调度室、安质部", "上传纪要、管理隐患清单"),
        ],
        [Cm(4.5), Cm(4), Cm(5.5)],
    )

    add_heading(doc, "安全管理", level=2)
    add_para(doc, "安全管理板块详细需求将在后续版本补充，本期建立以下子模块标题：")
    for title in [
        "人员实名制管理（劳务看板、人员实名制、人员轨迹、考勤明细、考勤统计、劳务黑名单、预警清单、实名制配置、设备管理等）",
        "车辆管理（车辆管理看板、进出场记录、车辆轨迹监管、车牌管理、设备管理、预警清单等）",
        "视频监控（监控列表、NVR设备管理等）",
    ]:
        add_heading(doc, title, level=3)

    add_heading(doc, "质量管理", level=2)
    add_para(doc, "质量管理板块详细需求将在后续版本补充，本期建立以下子模块标题：")
    for title in [
        "质量验评（质量报验、验评流程、省统表、四方验评等业务）",
        "质量检查与隐患（质量隐患清单、抽检、闭环处置等）",
    ]:
        add_heading(doc, title, level=3)


def append_third_party(doc):
    add_heading(doc, "第三方接口说明", level=1)
    add_table(
        doc,
        ["接口名称", "对接系统", "数据方向", "主要数据内容", "说明"],
        [
            ("视频流媒体", "海康/宇视等NVR平台", "接入", "实时视频流、回放、云台控制", "通过NVR注册统一管理"),
            ("巡检对讲", "执法记录仪/Web对讲", "双向", "实时音视频流", "支持远程呼叫与项目调度"),
            ("劳务/闸机标准接口", "各项目劳务系统", "接入", "人员信息、考勤记录", "供人员风险核验、会议签到引用"),
            ("OA用户组织同步", "OA办公系统", "接入", "组织机构、用户账号", "组织管理模块使用"),
            ("AI语音/文档识别", "大模型服务", "接入", "语音转写、纪要摘取", "会议录屏、监理纪要识别（试点）"),
        ],
        [Cm(3), Cm(3.5), Cm(1.5), Cm(4.5), Cm(3.5)],
    )


def update_revision_table(doc):
    for table in doc.tables:
        if len(table.rows) >= 2 and table.rows[0].cells[0].text.strip() == "版本号":
            row = table.rows[1]
            row.cells[0].text = "V1.1"
            row.cells[1].text = TODAY
            row.cells[2].text = "补充组织管理、COC调度中心详细需求；预留安全管理、质量管理章节"
            return


def save_document(doc: Document, path: Path = DOC_FILE) -> None:
    """Save in place; do not write sibling copies."""
    try:
        doc.save(str(path))
    except PermissionError:
        print(f"无法保存：{path}")
        print("请关闭 Word 中打开的该文档后重新运行。")
        sys.exit(1)


def fix_heading_numbers_file(path: Path) -> int:
    doc = Document(str(path))
    count = fix_all_heading_numbers(doc)
    sync_all_heading_fonts(doc)
    save_document(doc, path)
    return count


def sync_heading_fonts_file(path: Path) -> dict[str, int]:
    doc = Document(str(path))
    stats = sync_all_heading_fonts(doc)
    save_document(doc, path)
    return stats


def main():
    if len(sys.argv) > 1 and sys.argv[1] == "--fix-only":
        target = Path(sys.argv[2]) if len(sys.argv) > 2 else DOC_FILE
        n = fix_heading_numbers_file(target)
        print(f"已修正 {n} 个标题编号: {target}")
        return

    if len(sys.argv) > 1 and sys.argv[1] == "--sync-fonts":
        target = Path(sys.argv[2]) if len(sys.argv) > 2 else DOC_FILE
        stats = sync_heading_fonts_file(target)
        print(f"已同步标题编号字体: {target} -> {stats}")
        return

    doc = Document(str(DOC_FILE))
    update_revision_table(doc)
    clear_body_from(doc, 72)  # from "引言"
    append_intro_and_analysis(doc)
    append_detail_design(doc)
    append_third_party(doc)
    remove_trailing_junk(doc)
    fix_all_heading_numbers(doc)
    sync_all_heading_fonts(doc)
    save_document(doc, DOC_FILE)
    print(f"已更新: {DOC_FILE}")


if __name__ == "__main__":
    main()
