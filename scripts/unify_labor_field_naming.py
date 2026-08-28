# -*- coding: utf-8 -*-
"""将人员实名制 PRD 字段名统一为 snake_case（与架构一致）；保留 rule_key 驼峰枚举。"""
import re
import shutil
from pathlib import Path

from docx import Document

DIR = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
PRD = DIR / "人员实名制_PRD_待评审.docx"
BAK = DIR / "人员实名制_PRD_待评审_字段命名统一前_副本_20260813.docx"
ARCH = DIR / "product-architecture-人员实名制管理-v1.md"
ARCH_VEH = DIR / "product-architecture-车辆进出场管理-v1.md"

# 规则编码：架构约定为 camelCase，禁止改成 snake_case
RULE_KEYS = {
    "noLevel3Education",
    "specialCertMissing",
    "workOver12h",
    "ageLimit",
    "elderlyReminder",
    "idCardExpired",
    "absentDays",
    "managerAttendance",
    "blacklistEntry",
}


def camel_to_snake(name: str) -> str:
    name = re.sub(r"[（(].*?[）)]", "", name).strip()
    if not name or name in RULE_KEYS:
        return name
    if "_" in name and not re.search(r"[a-z][A-Z]", name):
        return name
    s1 = re.sub(r"(.)([A-Z][a-z]+)", r"\1_\2", name)
    return re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", s1).lower()


def build_renames(doc: Document):
    fields = set()
    for table in doc.tables:
        if not table.rows:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if "字段名" not in h:
            continue
        i = h.index("字段名")
        for row in table.rows[1:]:
            n = row.cells[i].text.strip()
            if n and n not in ("操作", "—", "-"):
                fields.add(n)
    pairs = []
    for f in fields:
        if f in RULE_KEYS:
            continue
        if "defaultRecipient" in f or f.startswith("defaultRecipientId"):
            pairs.append((f, "default_recipient_id"))
            continue
        if re.search(r"[a-z][A-Z]", f) or (any(c.isupper() for c in f[1:]) and f[0].islower()):
            sn = camel_to_snake(f)
            if sn != f:
                pairs.append((f, sn))
    # 长到短
    pairs.sort(key=lambda x: len(x[0]), reverse=True)
    return pairs


def apply_pairs(text: str, pairs) -> str:
    out = text
    for old, new in pairs:
        out = out.replace(old, new)
    return out


def patch_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = apply_pairs(text, pairs)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def patch_prd():
    shutil.copy2(PRD, BAK)
    doc = Document(str(PRD))
    pairs = build_renames(doc)
    print("rename pairs:", len(pairs))
    for a, b in pairs[:15]:
        print(" ", a, "->", b)
    n = 0
    for p in doc.paragraphs:
        n += patch_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += patch_paragraph(p, pairs)
    # version note if meta exists
    for table in doc.tables[:8]:
        headers = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        if headers and headers[0] == "版本":
            existing = [r.cells[0].text.strip() for r in table.rows]
            if "字段命名" not in "".join(c.text for r in table.rows for c in r.cells):
                nr = table.add_row()
                mapping = {
                    "版本": "字段命名修订",
                    "修订人": "",
                    "日期": "2026-08-13",
                    "说明": "字段名与产品架构统一为 snake_case（如 project_id、personnel_no）；规则编码 rule_key 仍为 camelCase（如 managerAttendance）",
                }
                for i, h in enumerate(headers):
                    if i < len(nr.cells):
                        cell = nr.cells[i]
                        val = mapping.get(h, "")
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = val
                            for r in cell.paragraphs[0].runs[1:]:
                                r.text = ""
                        else:
                            cell.paragraphs[0].text = val
            break
    doc.save(str(PRD))
    print("prd hits", n)

    # verify
    doc2 = Document(str(PRD))
    left = []
    for table in doc2.tables:
        if not table.rows:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if "字段名" not in h:
            continue
        i = h.index("字段名")
        for row in table.rows[1:]:
            name = row.cells[i].text.strip()
            if name and re.search(r"[a-z][A-Z]", name) and name not in RULE_KEYS:
                left.append(name)
    print("remaining camel fields", sorted(set(left)))


def ensure_naming_note(path: Path, module_label: str):
    text = path.read_text(encoding="utf-8")
    note = (
        "\n> **字段命名约定**：PRD / 架构 / 库表与接口字段名统一使用 **snake_case**"
        "（例：`project_id`、`plate_no`、`today_in_count`）。"
        "预警等**规则编码**（`rule_key`）可保留既有 **camelCase** 枚举值"
        "（例：`managerAttendance`），与字段名区分。\n"
    )
    if "字段命名约定" in text:
        return 0
    # insert after ## 2. 字段定义 heading
    marker = "## 2. 字段定义"
    if marker not in text:
        marker = "## 2."
    idx = text.find(marker)
    if idx < 0:
        return 0
    end = text.find("\n", idx)
    text = text[: end + 1] + note + text[end + 1 :]

    # bump version lightly if vehicle/labor meta present
    if module_label == "labor":
        if "| 版本号 | V1.42 |" in text:
            text = text.replace("| 版本号 | V1.42 |", "| 版本号 | V1.43 |", 1)
            text = text.replace(
                "| V1.42 修订 |",
                "| V1.43 修订 | 明确字段 **snake_case** 命名约定（与 PRD 对齐）；`rule_key` 仍为 camelCase。 |\n| V1.42 修订 |",
                1,
            )
        text = text.replace("| 更新日期 | 2026-08-12 |", "| 更新日期 | 2026-08-13 |", 1)
    path.write_text(text, encoding="utf-8")
    return 1


def main():
    patch_prd()
    print("labor arch note", ensure_naming_note(ARCH, "labor"))
    # vehicle arch: add note if missing
    vt = ARCH_VEH.read_text(encoding="utf-8")
    if "字段命名约定" not in vt:
        marker = "## 2. 字段定义"
        idx = vt.find(marker)
        if idx >= 0:
            end = vt.find("\n", idx)
            note = (
                "\n> **字段命名约定**：PRD / 架构 / 库表与接口字段名统一使用 **snake_case**"
                "（例：`project_id`、`plate_no`、`today_in_count`）。与人员实名制模块一致。\n"
            )
            vt = vt[: end + 1] + note + vt[end + 1 :]
            if "| 版本号 | V1.31 |" in vt:
                vt = vt.replace("| 版本号 | V1.31 |", "| 版本号 | V1.32 |", 1)
                vt = vt.replace(
                    "| V1.31 修订 |",
                    "| V1.32 修订 | 文首补充字段 **snake_case** 命名约定说明。 |\n| V1.31 修订 |",
                    1,
                )
            ARCH_VEH.write_text(vt, encoding="utf-8")
            print("vehicle arch note added")
    else:
        print("vehicle arch note already present")


if __name__ == "__main__":
    main()
