# -*- coding: utf-8 -*-
"""重绘中文字体流程图并替换 PRD 中已插入图片。"""
from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

sys.stdout.reconfigure(encoding="utf-8")

FIG_DIR = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\docs\prd-shots\labor-data-ingest"
)
DOC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审.docx"
)


def setup_font():
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        p = Path(path)
        if p.exists():
            font_manager.fontManager.addfont(str(p))
            name = font_manager.FontProperties(fname=str(p)).get_name()
            plt.rcParams["font.sans-serif"] = [name]
            plt.rcParams["axes.unicode_minus"] = False
            print("font:", name, path)
            return name
    raise RuntimeError("未找到可用中文字体")


def box(ax, x, y, w, h, text, fc="#F5F7FA", ec="#409EFF"):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9)


def arrow(ax, x1, y1, x2, y2):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.1,
            color="#606266",
        )
    )


def make_charts():
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(10.5, 5.4), dpi=160)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.2)
    ax.axis("off")
    ax.set_title("人员实名制数据接入双路径（总体）", fontsize=12, pad=8)
    box(
        ax,
        0.4,
        4.6,
        4.2,
        1.3,
        "路径A（现有，逻辑不变）\n已对接一期实名制的项目\n一期系统 → ROMA",
        fc="#F0F9EB",
        ec="#67C23A",
    )
    box(
        ax,
        7.4,
        4.6,
        4.2,
        1.3,
        "路径B（补充）\n一期未对接的项目\n闸机/实名制子系统\n→ 标准上报接口",
        fc="#FDF6EC",
        ec="#E6A23C",
    )
    box(ax, 3.5, 2.6, 5.0, 1.1, "智慧工程建设管控一体化平台\n人员实名制数据接入", fc="#ECF5FF")
    box(
        ax,
        3.0,
        0.4,
        6.0,
        1.3,
        "落库人员主档（统一口径）\n人员基础信息 / 单位信息 / 特种作业信息 / 安全教育信息",
        fc="#FEF0F0",
        ec="#F56C6C",
    )
    arrow(ax, 2.5, 4.6, 5.0, 3.7)
    arrow(ax, 9.5, 4.6, 7.0, 3.7)
    arrow(ax, 6.0, 2.6, 6.0, 1.7)
    p1 = FIG_DIR / "fig-data-ingest-overview.png"
    fig.tight_layout()
    fig.savefig(p1, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(10.5, 3.8), dpi=160)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    ax.set_title("路径B：标准接口上报流程（一期未对接项目）", fontsize=12, pad=8)
    steps = [
        (0.3, "项目现场闸机\n或实名制子系统"),
        (3.0, "调用本平台\n标准上报接口"),
        (5.7, "校验鉴权\n与字段完整性"),
        (8.4, "写入人员主档\n及关联信息"),
    ]
    for i, (x, text) in enumerate(steps):
        box(ax, x, 1.4, 2.4, 1.3, text, fc="#FDF6EC", ec="#E6A23C")
        if i < len(steps) - 1:
            arrow(ax, x + 2.4, 2.05, steps[i + 1][0], 2.05)
    box(
        ax,
        3.5,
        0.2,
        5.0,
        0.8,
        "上报字段：人员基础信息、单位信息、特种作业信息、安全教育信息",
        fc="#F5F7FA",
        ec="#909399",
    )
    p2 = FIG_DIR / "fig-data-ingest-api.png"
    fig.tight_layout()
    fig.savefig(p2, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print("charts:", p1, p2)
    return p1, p2


def replace_images(p1: Path, p2: Path):
    doc = Document(str(DOC))
    caps = {
        "图 5.2-1 人员实名制数据接入双路径（总体）": p1,
        "图 7.1-1 路径B：标准接口上报流程（一期未对接项目）": p2,
    }
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t not in caps:
            continue
        prev = paras[i - 1] if i > 0 else None
        if prev is None:
            continue
        for child in list(prev._p):
            prev._p.remove(child)
        run = prev.add_run()
        run.add_picture(str(caps[t]), width=Inches(6.2))
        prev.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("replaced:", t)
    doc.save(str(DOC))
    print("saved:", DOC)


def main():
    setup_font()
    p1, p2 = make_charts()
    replace_images(p1, p2)


if __name__ == "__main__":
    main()
