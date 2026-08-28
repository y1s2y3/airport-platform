# -*- coding: utf-8 -*-
"""V2.1.8：在 §6.5.8 增加「进出场记录」独立字段说明表。"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
DOC = BASE / "人员实名制_PRD_评审通过.docx"
BACKUP = BASE / "人员实名制_PRD_评审通过_backup_20260821_v218.docx"

ACCESS_ROWS = [
    ["实体/对象", "字段名", "中文名", "类型", "必填", "说明"],
    ["进出场记录", "id", "记录 ID", "string", "是", "流水主键"],
    ["", "project_id", "项目", "ref", "是", "所属项目"],
    ["", "personnel_id", "人员", "ref", "条件", "与主档关联；也可按证件号逻辑关联"],
    ["", "name", "姓名", "string", "是", "上报姓名"],
    ["", "id_card", "身份证（脱敏）", "string", "否", "列表/弹窗脱敏展示；查看明文记日志（若展示）"],
    ["", "direction", "进出场方向", "enum", "是", "进场（in）/ 出场（out）；单条事件方向固定"],
    ["", "record_time", "时间", "datetime", "是", "刷卡/识别发生时间；派生日汇总时按此排序"],
    ["", "gate_name", "闸机", "string", "否", "上报闸机名称文案；无设备台账外键"],
    ["", "photo", "考勤照片", "string/url", "否", "抓拍/附件；无图可空"],
    ["", "external_id", "外部流水号", "string", "否", "对接方幂等键；有则按约定去重"],
]

ACCESS_LOGIC = (
    "进出场记录为入站真源：一期系统上报，或施工现场实名制系统按标准接口上报；"
    "本平台只读展示（弹窗「查看进出场记录」），不补录、不改流水、不直连闸机私有协议。"
    "单条记录表示一次进场或出场事件；考勤明细日汇总中的进/出场时间、闸机、进出场、工时、在场状态均由本表按 §6.5.10 规则计算。"
)


def set_para_text(p, value: str) -> None:
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    first = cell.paragraphs[0]
    set_para_text(first, value)
    for p in cell.paragraphs[1:]:
        set_para_text(p, "")


def format_table(table: Table) -> None:
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_heading_space_after(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            p.paragraph_format.space_after = Pt(6)


def insert_paragraph_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = paragraph._p.makeelement(qn("w:p"), {})
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def insert_table_before_element(doc: Document, anchor_elm, rows: list[list[str]]) -> Table:
    """Insert a table immediately before anchor_elm (e.g. an existing tbl)."""
    # Create table at end then move
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i].cells[j], val)
    format_table(table)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor_elm.addprevious(tbl)
    return table


def update_meta(doc: Document) -> None:
    t0 = doc.tables[0]
    for row in t0.rows:
        k = row.cells[0].text.strip()
        if k == "版本号":
            set_cell_text(row.cells[1], "V2.1.8")
        if k == "更新日期":
            set_cell_text(row.cells[1], "2026-08-21")


def insert_revision(doc: Document) -> None:
    t1 = doc.tables[1]
    new_row = t1.add_row()
    vals = (
        "V2.1.8",
        "—",
        "2026-08-21",
        "考勤明细 §6.5.8：单独增加「进出场记录」字段说明表（方向/时间/闸机/照片等）及逻辑说明；"
        "与日汇总派生字段表分列。",
    )
    for i, v in enumerate(vals):
        if i < len(new_row.cells):
            set_cell_text(new_row.cells[i], v)
    tbl = t1._tbl
    tr = new_row._tr
    tbl.remove(tr)
    t1.rows[0]._tr.addnext(tr)
    format_table(t1)


def find_day_field_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:3] != ["字段名", "中文名", "类型"]:
            continue
        names = {r.cells[0].text.strip() for r in table.rows}
        if "clock_in" in names:
            return table
    return None


def patch_field_section(doc: Document) -> None:
    # Rename combined title paragraph
    title_para = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if "进出场记录（入站真源，只读）" in t and "日汇总" in t:
            title_para = p
            break
        if t == "进出场记录（入站真源，只读）；考勤明细日汇总（派生，只读）":
            title_para = p
            break

    day_table = find_day_field_table(doc)
    if day_table is None:
        raise SystemExit("day field table not found")

    # Insert before day summary table: logic + access title + access table + day title
    anchor = day_table._tbl

    # day title paragraph before existing table
    day_title = doc.add_paragraph("考勤明细日汇总（ATTENDANCE_DAY，派生，只读）")
    day_title_p = day_title._p
    day_title_p.getparent().remove(day_title_p)
    anchor.addprevious(day_title_p)

    # access table
    insert_table_before_element(doc, day_title_p, ACCESS_ROWS)

    # access title + logic (in reverse order before table which is now before day_title)
    # After insert_table, access tbl is before day_title. Insert titles before access tbl.
    # Find the access table we just inserted = previous sibling of day_title
    access_tbl = day_title_p.getprevious()
    while access_tbl is not None and access_tbl.tag != qn("w:tbl"):
        access_tbl = access_tbl.getprevious()
    if access_tbl is None:
        raise SystemExit("access table not found after insert")

    access_title = doc.add_paragraph("进出场记录（PERSON_ACCESS_RECORD，入站真源，只读）")
    access_title_p = access_title._p
    access_title_p.getparent().remove(access_title_p)
    access_tbl.addprevious(access_title_p)

    logic = doc.add_paragraph(ACCESS_LOGIC)
    logic_p = logic._p
    logic_p.getparent().remove(logic_p)
    access_title_p.addprevious(logic_p)

    if title_para is not None:
        # remove old combined title to avoid duplicate
        set_para_text(title_para, "")


def main() -> None:
    if not DOC.exists():
        raise SystemExit(f"missing {DOC}")
    shutil.copy2(DOC, BACKUP)
    print("backup ->", BACKUP)

    doc = Document(str(DOC))
    # idempotency: skip if already has access field table header with direction in 6-col entity table
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:3] == ["实体/对象", "字段名", "中文名"]:
            body = " ".join(c.text for r in table.rows for c in r.cells)
            if "进出场记录" in body and "direction" in body and "record_time" in body:
                print("already patched, only bump if needed")
                ver = doc.tables[0].rows[1].cells[1].text.strip()
                if ver == "V2.1.8":
                    return
                break

    update_meta(doc)
    insert_revision(doc)
    patch_field_section(doc)
    set_heading_space_after(doc)
    doc.save(str(DOC))

    # verify
    d2 = Document(str(DOC))
    ver = d2.tables[0].rows[1].cells[1].text.strip()
    found = False
    for table in d2.tables:
        h = [c.text.strip() for c in table.rows[0].cells]
        if h[:2] == ["实体/对象", "字段名"]:
            joined = " ".join(c.text for r in table.rows for c in r.cells)
            if "direction" in joined and "进出场记录" in joined:
                found = True
                print("access field table rows", len(table.rows))
    print("version", ver, "access_table", found)
    assert ver == "V2.1.8" and found


if __name__ == "__main__":
    main()
