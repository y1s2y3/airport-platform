# -*- coding: utf-8 -*-
from docx import Document
from pathlib import Path

p = Path(r"调研记录/4、安全管理/人员、车辆管理/人员实名制_PRD_评审通过.docx")
d = Document(str(p))
out = Path(r"Airport_Expansion_Project_Informatization/scripts/_prd_attend_extract.txt")
lines = []
keys = ("考勤", "进出场", "ATTENDANCE", "修订记录", "版本号", "V2.1", "clock_", "gate_", "工时", "在场状态")
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if not t:
        continue
    if any(k in t for k in keys):
        style = para.style.name if para.style else ""
        lines.append(f"P{i}|{style}|{t}")
for ti, table in enumerate(d.tables):
    hit = False
    rows_txt = []
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
        joined = " | ".join(cells)
        if any(k in joined for k in keys + ("版本", "修订")):
            hit = True
        if ri < 10 or any(k in joined for k in ("考勤", "进出场", "clock", "gate", "工时", "V2.1")):
            rows_txt.append(f"  R{ri}: {joined[:240]}")
    if hit:
        lines.append(f"TABLE{ti}")
        lines.extend(rows_txt[:15])
        lines.append("---")
out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "lines", len(lines))
