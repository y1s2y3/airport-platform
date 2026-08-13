# -*- coding: utf-8 -*-
"""V1.4: ROMA/通知重生/黑名单条数/导出/必填空值/模块编码/5.2流程."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

PATH = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_V1.3.docx"
)
OUTS = [
    PATH,
    Path(
        r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_副本_20260812.docx"
    ),
    Path(
        r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_待评审_V1.4.docx"
    ),
]


def replace_in_paragraph(p, pairs):
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)
    return 1


def set_cell_text(cell, value):
    if not cell.paragraphs:
        cell.text = value
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def apply_all(doc, pairs):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, pairs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_paragraph(p, pairs)
    return n


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def main():
    doc = Document(str(PATH))

    pairs = [
        # ROMA does not overwrite local L3 education
        (
            "读：ROMA 主档；\n可编辑范围仅限三级安全教育记录；查看手机号、证件号等敏感信息时写入操作日志（业务归类为「人员实名制管理」）。",
            "读：ROMA 主档（基本身份/单位工种/特种证等）；可编辑范围仅限三级安全教育记录。本地已保存的三级安全教育记录不被 ROMA 回写覆盖。查看手机号、证件号等敏感信息时写入操作日志（业务归类为「人员实名制管理」）。",
        ),
        (
            "读：ROMA 主档；",
            "读：ROMA 主档（基本身份/单位工种/特种证等）；本地三级安全教育不被 ROMA 回写覆盖；",
        ),
        (
            "二期不支持新建人员主档；人员仅经一期 → ROMA 同步进入；界面无新建入口",
            "二期不支持新建人员主档；人员仅经一期 → ROMA 同步进入；界面无新建入口；本地三级安全教育记录不被 ROMA 回写覆盖",
        ),
        # notification no regenerate
        (
            "通知类（高龄提醒、身份证过期提醒）：状态为已通知，无需关闭，不参与分级上报与超期升级；仍推送个人中心消息；详情只读、无处置并关闭；时间线仅保留触发记录。待处理/待人工处置指标不含通知类；处置方式筛选含「通知」。",
            "通知类（高龄提醒、身份证过期提醒）：状态为已通知，无需关闭，不参与分级上报与超期升级；仍推送个人中心消息；详情只读、无处置并关闭；时间线仅保留触发记录。同一人同一规则在同一项目仅生成一条，已通知后即使条件仍成立也不重生。待处理/待人工处置指标不含通知类；处置方式筛选含「通知」。",
        ),
        (
            "(每日定时任务触发，一个人最多有一条该类型待处理预警)",
            "(每日定时任务触发；同一人同一规则同一项目最多一条；通知类已生成后不重生)",
        ),
        (
            "(每日定时任务触发，一个人最多有一条该类型待处理预警)。",
            "(每日定时任务触发；同一人同一规则同一项目最多一条；通知类已生成后不重生)。",
        ),
        # blacklist 1 per project
        (
            "每日定时任务扫描：人员在岗状态为「在岗」且证件号命中全平台劳务黑名单、且该规则已开启时触发。\n（一个人最多有一条该类型待处理预警）",
            "每日定时任务扫描：人员在岗状态为「在岗」且证件号命中全平台劳务黑名单、且该规则已开启时触发。同一人多项目同时在岗时，每个项目各生成 1 条待处理预警（按项目隔离，不合并为全平台 1 条）。",
        ),
        (
            "每日定时任务扫描：人员在岗状态为「在岗」且证件号命中全平台劳务黑名单、且该规则已开启时触发。",
            "每日定时任务扫描：人员在岗状态为「在岗」且证件号命中全平台劳务黑名单、且该规则已开启时触发；同一人多项目同时在岗时每项目各 1 条。",
        ),
        (
            "（一个人最多有一条该类型待处理预警）",
            "（同一人同一规则同一项目最多一条；黑名单跨项目按每项目 1 条）",
        ),
        (
            "维护劳务黑名单台账；每日定时扫描在岗人员命中黑名单时触发人员预警；移出黑名单或人员离场后，已生成预警按自动关闭/处置闭环，且不再对新命中触发。",
            "维护劳务黑名单台账；每日定时扫描在岗人员命中黑名单时按项目触发预警（同一人多项目在岗则每项目 1 条）；移出黑名单或该项目人员离场后，已生成预警按次日自动关闭，且该项目不再因同一命中重复生成。",
        ),
        # export
        (
            "导出\n导出按钮\n导出当前筛选结果\n无数据提示",
            "导出\n导出按钮\n按当前筛选结果，导出列表全部字段的全量数据（字段=列表列；本期不描述异步导出细节）\n无数据提示",
        ),
        (
            "导出\n导出按钮\n导出当前筛选\n无数据提示",
            "导出\n导出按钮\n按当前筛选结果，导出列表全部字段的全量数据（字段=列表列；本期不描述异步导出细节）\n无数据提示",
        ),
        (
            "台账/考勤/预警列表分页查询；导出大数据量异步。",
            "台账/考勤/预警列表分页查询；导出按列表字段全量导出（字段同列表列），本期不描述异步导出实现细节。",
        ),
        # photo / clockOut empty
        (
            "photo\n照片\nstring\n是\n只读",
            "photo\n照片\nstring\n否\n只读；ROMA 无照片时显示默认占位图，不阻断详情查看",
        ),
        (
            "clockOut\n出场时间\ndatetime\n是\n—",
            "clockOut\n出场时间\ndatetime\n否\n未出场时可为空，界面显示「—」；不因空值报错",
        ),
        (
            "clockIn\n进场时间\ndatetime\n是\n—",
            "clockIn\n进场时间\ndatetime\n条件\n有进场记录时必有值；仅出场异常数据允许空并显示「—」",
        ),
    ]

    print("hits:", apply_all(doc, pairs))

    # More precise cell-level fixes for photo/clockOut if table cells split
    for table in doc.tables:
        rows = table.rows
        for ri, row in enumerate(rows):
            c0 = row.cells[0].text.strip() if row.cells else ""
            if c0 == "photo" and len(row.cells) >= 5:
                # check 必填 col
                if row.cells[3].text.strip() == "是":
                    set_cell_text(row.cells[3], "否")
                if "占位" not in row.cells[4].text:
                    set_cell_text(
                        row.cells[4],
                        "只读；ROMA 无照片时显示默认占位图，不阻断详情查看",
                    )
                    print("photo cell fixed", ri)
            if c0 == "clockOut" and len(row.cells) >= 5:
                set_cell_text(row.cells[3], "否")
                set_cell_text(
                    row.cells[4],
                    "未出场时可为空，界面显示「—」；不因空值报错",
                )
                print("clockOut cell fixed")
            if c0 == "clockIn" and len(row.cells) >= 5 and row.cells[3].text.strip() == "是":
                set_cell_text(row.cells[3], "条件")
                set_cell_text(
                    row.cells[4],
                    "有进场记录时必有值；异常空值显示「—」",
                )
                print("clockIn cell fixed")

    # Export interaction rows: find cells with exact 导出当前筛选结果
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t in ("导出当前筛选结果", "导出当前筛选"):
                    set_cell_text(
                        cell,
                        "按当前筛选，导出列表全部字段全量数据（字段=列表列；本期不描述异步细节）",
                    )
                    print("export cell fixed")

    # Add module encoding to meta table
    meta = doc.tables[0]
    has_code = any(r.cells[0].text.strip() == "模块编码" for r in meta.rows)
    if not has_code:
        # insert after 模块名称 row
        insert_at = None
        for i, r in enumerate(meta.rows):
            if r.cells[0].text.strip() == "模块名称":
                insert_at = i
                break
        # add_row appends at end; then we set values. Order less critical for meta.
        nr = meta.add_row()
        set_cell_text(nr.cells[0], "模块编码")
        set_cell_text(nr.cells[1], "LABOR（功能ID 前缀 F-LABOR）")
        print("module code added")

    # Fill 5.2 after heading
    flow_text = (
        "核心流程（只读监管 + 局部补录 + 预警闭环）：\n"
        "1）一期经 ROMA 同步人员主档/特种证/考勤 → 项目 Web 台账、考勤、看板与指挥部统计只读展示；\n"
        "2）项目侧仅可补录三级安全教育（本地保存，ROMA 回写不覆盖）；\n"
        "3）按项目实名制配置（轨迹外链、默认接收人、分级管控、9 条规则）定时生成/关闭预警；通知类一次性生成不重生；\n"
        "4）任务类预警仅默认接收人可在 Web（个人中心或预警清单）处置；自动关闭类次日定时关闭；APP 仅查看；\n"
        "5）黑名单全平台维护，日扫各项目在岗命中，每项目各生成 1 条预警，引导处理到位后次日自动关闭；\n"
        "6）人员轨迹仅外链跳转：指挥部有地址可跳（含停用），项目侧须启用且地址有效。"
    )
    for p in doc.paragraphs:
        if p.text.strip() == "5.2 本模块核心业务流程":
            # if next already filled, skip
            nxt = p._p.getnext()
            already = False
            if nxt is not None:
                # peek text
                from docx.oxml.ns import qn

                texts = nxt.findall(".//" + qn("w:t"))
                joined = "".join(t.text or "" for t in texts)
                if "核心流程" in joined:
                    already = True
            if not already:
                insert_paragraph_after(p, flow_text)
                print("5.2 flow inserted")
            break

    # Add business rule bullets into 6.4.9 / 6.6 / 6.8 / 6.5 if missing via pairs already
    # Extra explicit rules paragraphs near known rule headers
    extra_rules = [
        (
            "枚举库表存英文编码，界面展示中文",
            "枚举库表存英文编码，界面展示中文\n"
            "本地三级安全教育不被 ROMA 回写覆盖\n"
            "照片：非强必填；无图显示占位，不阻断\n"
            "导出：按列表字段全量导出当前筛选结果，本期不描述异步细节",
        ),
        (
            "APP 无预警清单菜单，仅消息中心",
            "APP 无预警清单菜单，仅消息中心\n"
            "通知类：同一人同一规则同一项目仅一条，已通知后不重生\n"
            "自动关闭类：统一次日定时任务关闭",
        ),
        (
            "规则开关在 F-LABOR-07「实名制配置」中的黑名单进场规则",
            "规则开关在 F-LABOR-07「实名制配置」中的黑名单进场规则\n"
            "同一人多项目同时在岗：每个项目各生成 1 条黑名单预警",
        ),
        (
            "只读；不做手工补录、不做考勤机设备管理",
            "只读；不做手工补录、不做考勤机设备管理\n"
            "出场时间（clockOut）允许为空（未出场显示「—」）\n"
            "导出：按列表字段全量导出当前筛选结果，本期不描述异步细节",
        ),
    ]
    print("extra rule hits:", apply_all(doc, extra_rules))

    # Meta version bump V1.4
    for table in doc.tables[:6]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            if cells[0] == "版本号":
                set_cell_text(row.cells[1], "V1.4")
            if cells[0] == "更新日期":
                set_cell_text(row.cells[1], "2026-08-12")
            if cells[0] == "状态":
                set_cell_text(row.cells[1], "待评审")

    for table in doc.tables[:10]:
        headers = [c.text.strip() for c in table.rows[0].cells]
        if not headers or headers[0] != "版本":
            continue
        existing = [r.cells[0].text.strip() for r in table.rows]
        if "V1.4" not in existing:
            nr = table.add_row()
            mapping = {
                "版本": "V1.4",
                "修订人": "",
                "日期": "2026-08-12",
                "说明": (
                    "开发边界补齐：ROMA不覆盖本地三级教育；通知类不重生；黑名单按项目每条1条；"
                    "导出=列表字段全量且不描述异步；照片/出场时间空值规则；模块编码LABOR；补写5.2核心流程"
                ),
            }
            for i, h in enumerate(headers):
                if i < len(nr.cells):
                    set_cell_text(nr.cells[i], mapping.get(h, ""))
            print("revision V1.4 added")
        break

    # Out of scope already has 短信 - OK

    for out in OUTS:
        try:
            doc.save(str(out))
            print("saved", out.name)
        except PermissionError:
            print("LOCKED", out.name)


if __name__ == "__main__":
    main()
