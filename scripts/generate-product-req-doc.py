# -*- coding: utf-8 -*-
"""生成《深圳机场扩建工程智慧工程建设管控一体化平台 需求规格说明书》"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_FILE = ROOT / "产品需求规格说明书.docx"
COC_SHOTS = ROOT / "docs" / "coc-screen-shots"
SAFETY_SHOTS = ROOT / "docs" / "safety-req-shots"

DOC_NO = "SZAIR-REQ-20260705"
VERSION = "V1.0"
TODAY = date(2026, 7, 5).strftime("%Y年%m月%d日")
PROJECT_NAME = "深圳机场扩建工程智慧工程建设管控一体化平台"


def set_heading_space_after(paragraph, pt=6):
    paragraph.paragraph_format.space_after = Pt(pt)


def format_table_rows(table, min_height_cm=0.6):
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(min_height_cm)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_run_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    set_heading_space_after(p)
    for run in p.runs:
        add_run_font(run, "黑体", 14 if level == 1 else 12, False)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        add_run_font(r1, bold=True)
        r2 = p.add_run(text)
        add_run_font(r2)
    else:
        r = p.add_run(text)
        add_run_font(r)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    add_run_font(r)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val)
    format_table_rows(table)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
    doc.add_paragraph()
    return table


def add_screenshot(doc, title, desc, image_name, shot_dir):
    add_heading(doc, title, level=4)
    add_para(doc, desc)
    img_path = shot_dir / image_name if image_name else None
    if img_path and img_path.exists():
        doc.add_paragraph()
        doc.add_picture(str(img_path), width=Inches(6.2))
    doc.add_paragraph()


def add_module_section(doc, module_title, overview, flows, structure, screenshots, fields, use_cases, relations, permissions):
    add_heading(doc, module_title, level=2)
    add_heading(doc, "功能概述", level=3)
    add_para(doc, overview)
    add_heading(doc, "业务流程", level=3)
    for item in flows:
        add_bullet(doc, item)
    add_heading(doc, "功能结构", level=3)
    for item in structure:
        add_bullet(doc, item)
    if screenshots:
        add_heading(doc, "原型页面", level=3)
        add_para(doc, "以下截图为当前原型核心页面，供研发核对布局与功能范围。")
        for title, desc, fname, shot_dir in screenshots:
            add_screenshot(doc, title, desc, fname, shot_dir)
    if fields:
        add_heading(doc, "字段描述", level=3)
        add_table(doc, fields[0], fields[1], fields[2] if len(fields) > 2 else None)
    if use_cases:
        add_heading(doc, "功能用例", level=3)
        add_table(doc, use_cases[0], use_cases[1], use_cases[2] if len(use_cases) > 2 else None)
    if relations:
        add_heading(doc, "关联模块", level=3)
        for item in relations:
            add_bullet(doc, item)
    if permissions:
        add_heading(doc, "权限要求", level=3)
        add_table(doc, permissions[0], permissions[1], permissions[2] if len(permissions) > 2 else None)


def build_master_function_list():
    rows = []
    # 基础设置
    settings = [
        ("基础设置", "系统设置", "组织结构", "统一组织节点树，OA/外部单位/其他用户分类展示，支持 OA 同步"),
        ("基础设置", "系统设置", "用户管理", "管理系统用户账号、组织归属、角色绑定，支持启用/停用、密码重置"),
        ("基础设置", "系统设置", "部门管理", "维护部门层级与人员归属，支持增删改及负责人设置"),
        ("基础设置", "系统设置", "角色管理", "定义角色并配置菜单、数据权限，实现 RBAC"),
        ("基础设置", "系统设置", "权限管理", "细粒度配置功能权限与数据范围，支持按角色批量授权"),
        ("基础设置", "系统设置", "菜单管理", "维护菜单树、路由与按钮级权限标识"),
        ("基础设置", "系统对接", "第三方用户管理", "对接 OA/一期平台用户同步与映射"),
        ("基础设置", "系统对接", "第三方角色管理", "维护第三方角色与平台角色映射"),
        ("基础设置", "系统对接", "第三方组织管理", "同步第三方组织并与平台组织结构对照"),
        ("基础设置", "系统对接", "第三方菜单管理", "配置第三方菜单与平台菜单映射，支持 SSO 跳转"),
        ("基础设置", "日志管理", "登录日志", "记录登录时间、IP、终端、结果，支持检索导出"),
        ("基础设置", "日志管理", "操作日志", "记录关键业务操作，含操作人、时间、对象与结果"),
    ]
    # 基础输出（基础数据管理）
    basic = [
        ("基础输出", "项目基础信息", "项目台账", "指挥部维护全部项目基础信息，含参建单位、概况、安全画像等"),
        ("基础输出", "项目基础信息", "新增/编辑", "支持新增、编辑项目，自动生成项目简称"),
        ("基础输出", "分包单位管理", "分包台账", "按项目维护分包单位安全生产许可、负责人及资质证书"),
        ("基础输出", "分包单位管理", "详情查看", "查看分包单位完整登记信息与证书附件"),
        ("基础输出", "工程划分库", "树结构维护", "统一维护单位/子单位/分部/子分部/分项工程层级"),
        ("基础输出", "工程划分库", "版本与停用", "支持新增、编辑、停用、版本管理及 BIM 关联"),
    ]
    # 人员实名制
    labor = [
        ("安全管理", "人员实名制", "劳务看板", "展示在场人数、进场/出场、预警等核心指标与图表分析"),
        ("安全管理", "人员实名制", "人员实名制", "维护劳务人员台账，支持检索、脱敏、详情查看"),
        ("安全管理", "人员实名制", "考勤明细", "同步门禁刷卡进出记录，支持多维检索"),
        ("安全管理", "人员实名制", "考勤统计", "按项目/单位/工种统计出勤人数与出勤率"),
        ("安全管理", "人员实名制", "劳务黑名单", "维护指挥部统一黑名单，进场自动预警"),
        ("安全管理", "人员实名制", "预警清单", "汇总实名制预警，支持处置与分级上报"),
        ("安全管理", "人员实名制", "实名制配置", "配置预警规则、分级管控及现场对接开关"),
        ("安全管理", "人员实名制", "数据对接", "提供劳务/闸机系统标准接口"),
    ]
    # 车辆
    vehicle = [
        ("安全管理", "车辆管理", "车辆管理看板", "展示进出场量、在场/在途数量及预警统计"),
        ("安全管理", "车辆管理", "进出场记录", "记录车辆进出场，对接工地自建车辆系统"),
        ("安全管理", "车辆管理", "车辆轨迹监管", "GIS 展示轨迹与电子围栏，越界触发报警"),
        ("安全管理", "车辆管理", "车牌管理", "维护车辆台账，支持授权道闸与导入导出"),
        ("安全管理", "车辆管理", "设备管理", "管理轨迹监测设备、道闸识别设备"),
        ("安全管理", "车辆管理", "预警清单", "汇总车辆监管与轨迹监管预警"),
        ("安全管理", "车辆管理", "数据对接", "提供车辆系统标准接口"),
    ]
    # 视频监控
    video = [
        ("安全管理", "视频监控", "监控列表", "按项目查看摄像头通道，支持编辑、排序"),
        ("安全管理", "视频监控", "设备管理", "管理 NVR 录像机注册、绑定项目及通道配置"),
    ]
    # COC
    coc_screen = [
        ("COC调度中心", "COC大屏", "工程指挥部首页", "视频监控、进度总览、红黑榜、劳务/隐患分析一屏统览"),
        ("COC调度中心", "COC大屏", "领导讲话", "九宫格视频墙、设备入会、静音控制、截屏生成单据"),
        ("COC调度中心", "COC大屏", "项目首页", "项目维度监控、劳务统计、隐患清单"),
        ("COC调度中心", "COC大屏", "项目调度", "对讲主画面、关联监控、调度记录、任务/提示/处罚单"),
        ("COC调度中心", "COC大屏", "会议签到", "指挥部/项目参会人员签到状态浮窗"),
        ("COC调度中心", "COC大屏", "会议管控", "开始/结束调度会议、页面录屏"),
    ]
    coc_admin = [
        ("COC调度中心", "COC后台", "每日施工作业", "危险作业计划填报与总览"),
        ("COC调度中心", "COC后台", "任务单", "远程调度任务单全流程管理"),
        ("COC调度中心", "COC后台", "提示函", "提示函创建、下发与闭环"),
        ("COC调度中心", "COC后台", "处罚单", "处罚单开具、申诉、复核与归档"),
        ("COC调度中心", "COC后台", "黑红榜单", "按期维护红榜/黑榜，大屏联动刷新"),
        ("COC调度中心", "COC后台", "巡检仪管理", "巡检仪注册、绑定项目/人员"),
        ("COC调度中心", "COC后台", "智能安全帽管理", "安全帽设备台账与绑定"),
        ("COC调度中心", "COC后台", "监理会议管理", "监理例会纪要上传与隐患识别"),
    ]
    rows.extend(settings + basic + labor + vehicle + video + coc_screen + coc_admin)
    return rows


def build_labor_function_rows():
    client = "一体化平台"
    src = "需求确认"
    return [
        (client, "人员实名制", "劳务看板", "汇总指标", "展示在场人数、今日进场/出场、登记人数、待处置预警等核心指标。", src),
        (client, "人员实名制", "劳务看板", "年龄/工种结构", "环形图展示各年龄段、各工种人员占比。", src),
        (client, "人员实名制", "劳务看板", "出勤趋势", "折线图展示近阶段出勤人数与出勤率变化。", src),
        (client, "人员实名制", "劳务看板", "各项目统计", "指挥部视角对比各项目在场、登记、预警情况。", src),
        (client, "人员实名制", "劳务看板", "预警动态", "展示近期实名制预警摘要，支持跳转预警清单。", src),
        (client, "人员实名制", "人员实名制", "人员列表", "按项目维护劳务人员台账，展示编号、姓名、工种、单位、入场状态等。", src),
        (client, "人员实名制", "人员实名制", "检索筛选", "支持姓名/编号/身份证/手机号/单位及工种、入场状态筛选。", src),
        (client, "人员实名制", "人员实名制", "手机号脱敏", "列表默认脱敏展示手机号，查看需记录操作日志。", src),
        (client, "人员实名制", "人员实名制", "新增/编辑", "未对接现场实名制系统的项目支持手工维护；已对接项目禁止手工新增。", src),
        (client, "人员实名制", "人员实名制", "人员详情", "查看基本信息、资质证书、安全教育、进退场记录、考勤摘要。", src),
        (client, "人员实名制", "考勤明细", "进出记录", "同步门禁刷卡进场/出场明细，含闸机、方向、时间。", src),
        (client, "人员实名制", "考勤统计", "统计汇总", "按项目、单位、工种统计出勤人数、出勤率。", src),
        (client, "人员实名制", "劳务黑名单", "黑名单台账", "维护指挥部统一劳务黑名单，进场或打卡自动生成预警。", src),
        (client, "人员实名制", "预警清单", "预警列表", "汇总各类实名制规则触发预警，支持状态筛选与处置。", src),
        (client, "人员实名制", "实名制配置", "预警规则", "指挥部配置预警规则开关及阈值。", src),
        (client, "人员实名制", "实名制配置", "分级管控", "配置四级上报天数及接收人。", src),
        (client, "人员实名制", "实名制配置", "现场对接", "按项目配置是否对接现场实名制系统。", src),
        (client, "人员实名制", "数据对接", "标准接口", "支持不同厂家劳务/闸机系统推送人员及考勤数据。", src),
    ]


def build_vehicle_function_rows():
    client = "一体化平台"
    src = "需求确认"
    return [
        (client, "车辆管理", "车辆管理看板", "汇总指标", "展示今日进出场、在场/在途、登记车辆及异常预警。", src),
        (client, "车辆管理", "进出场记录", "进出记录", "记录车辆进出场信息，对接工地自建车辆系统，不做准入限制。", src),
        (client, "车辆管理", "车辆轨迹监管", "GIS轨迹", "GIS 地图展示车辆位置与历史轨迹。", src),
        (client, "车辆管理", "车辆轨迹监管", "电子围栏", "支持创建禁行/允许区域围栏。", src),
        (client, "车辆管理", "车辆轨迹监管", "轨迹报警", "越界行为触发报警并纳入预警清单。", src),
        (client, "车辆管理", "车牌管理", "车辆台账", "维护车牌、类型、所属单位、司机、准入证明等。", src),
        (client, "车辆管理", "车牌管理", "授权道闸", "为车辆配置可通行道闸，支持批量授权。", src),
        (client, "车辆管理", "设备管理", "设备台账", "管理轨迹监测设备、道闸识别设备。", src),
        (client, "车辆管理", "预警清单", "预警列表", "汇总车辆监管与轨迹监管预警。", src),
        (client, "车辆管理", "数据对接", "标准接口", "同步车辆资质、进出场及轨迹数据。", src),
    ]


def build_coc_function_rows():
    client = "COC调度大屏"
    src = "需求确认"
    rows = [
        (client, "通用", "标题栏", "平台标题与关闭", "展示 COC 调度指挥中心名称、logo；支持关闭返回工作台。", "招投标"),
        (client, "通用", "会议签到", "签到浮窗", "展开/收起签到面板，展示人员姓名、单位、岗位及签到状态。", "调研结果"),
        (client, "通用", "会议管控", "会议录屏", "开始/结束调度会议，浏览器页面录屏并可本地保存。", "调研结果"),
        (client, "通用", "项目列表", "项目树", "工程指挥部根节点 + 各项目子节点，支持状态筛选与名称搜索。", "调研结果"),
        (client, "通用", "视频监控", "宫格展示", "3×2 分页展示摄像头，显示在线/离线状态。", "招投标"),
        (client, "通用", "视频监控", "重点视频", "支持标记重点摄像头并筛选。", "调研结果"),
        (client, "通用", "视频监控", "云台操作", "支持放大缩小、旋转、声音控制、主/子码流切换。", "招投标"),
        (client, "通用", "视频监控", "问题截图", "截图标记问题，可生成任务单/提示函/处罚单/隐患。", "调研结果"),
        (client, "通用", "手持对讲", "巡检对讲", "展示巡检终端与 Web 对讲席，点击进入项目调度。", "招投标"),
        (client, "指挥部", "领导讲话", "视频墙", "3×3 分页展示入会设备，支持邀请、静音、截屏生成单据。", "调研结果"),
        (client, "指挥部", "红黑榜单", "榜单展示", "双列展示红榜/黑榜，含项目简称、现场图、描述。", "调研结果"),
        (client, "指挥部", "项目进度", "进度总览", "展示项目总数、平均完成率、滞后节点占比及柱状对比。", "调研结果"),
        (client, "指挥部", "劳务分析", "多项目统计", "汇总各项目在场人数、特种作业、管理人员等。", "调研结果"),
        (client, "指挥部", "安质隐患", "等级分布", "环形图展示待处理隐患等级分布。", "调研结果"),
        (client, "项目层级", "项目调度", "对讲+监控", "对讲主画面 + 3×2 关联监控，下方调度记录与文书面板。", "调研结果"),
        (client, "项目层级", "人员风险核验", "风险分析", "针对安全教育、特种作业等分析人员风险。", "调研结果"),
        (client, "项目层级", "危险作业", "作业清单", "展示当日危险作业计划，支持全量查看。", "调研结果"),
    ]
    coc_admin = [
        ("COC后台", "通用", "巡检仪管理", "设备注册", "新增巡检仪编号、名称、绑定项目/人员。", "招投标"),
        ("COC后台", "通用", "智能安全帽", "设备台账", "管理安全帽设备绑定与在线状态。", "调研结果"),
        ("COC后台", "指挥部", "任务单", "台账管理", "远程调度任务单创建、下发、签收与闭环。", "调研结果"),
        ("COC后台", "指挥部", "提示函", "台账管理", "提示函全流程管理。", "调研结果"),
        ("COC后台", "指挥部", "处罚单", "台账管理", "处罚单开具、申诉、复核与红黑榜联动。", "调研结果"),
        ("COC后台", "指挥部", "黑红榜单", "按期管理", "维护红榜/黑榜，保存后大屏自动刷新。", "调研结果"),
        ("COC后台", "通用", "监理会议", "纪要上传", "上传监理例会纪要，系统识别归档隐患。", "调研结果"),
        ("COC后台", "项目", "每日施工作业", "危险作业", "每日填报危险作业计划并生成统计。", "调研结果"),
    ]
    return rows + coc_admin


def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    lines = [
        ("深圳机场扩建工程", 16, True),
        (PROJECT_NAME, 18, True),
        ("需求规格说明书", 22, True),
    ]
    for text, size, bold in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        add_run_font(r, "黑体", size, bold)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("中建三局集团有限公司")
    add_run_font(r, "宋体", 14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("二〇二六年")
    add_run_font(r2, "宋体", 14)
    doc.add_page_break()


def add_revision_table(doc):
    add_heading(doc, "需求版本修订记录", level=1)
    add_table(
        doc,
        ["版本号", "维护日期", "维护内容", "维护人", "审核人"],
        [
            (VERSION, TODAY, "新建文档，覆盖基础设置、基础输出、安全管理、COC调度中心四大板块", "", ""),
        ],
        [Cm(2), Cm(3), Cm(8), Cm(2.5), Cm(2.5)],
    )


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    add_cover_page(doc)

    # 文档信息表
    add_table(
        doc,
        ["文档属性", "内容"],
        [
            ("文档编号", DOC_NO),
            ("版本", VERSION),
            ("编制日期", TODAY),
            ("适用范围", "深圳机场扩建工程智慧工程建设管控一体化平台（一期已完成功能）"),
            ("交付对象", "研发团队"),
        ],
        [Cm(4), Cm(12)],
    )

    add_revision_table(doc)

    # 1 引言
    add_heading(doc, "1、引言", level=1)
    add_heading(doc, "1.1 目的", level=2)
    add_para(
        doc,
        f"本文档为{PROJECT_NAME}需求规格说明文档，以此作为系统设计、代码编写、测试阶段的开发设计依据。"
        "本文档包含系统体系结构、业务逻辑、功能设计、权限要求及第三方接口说明，覆盖基础设置、基础数据、"
        "安全管理（人员实名制、车辆管理、视频监控）及 COC 调度指挥中心四大业务板块。"
        "其中「基础输出」对应平台基础数据管理模块。",
    )
    add_heading(doc, "1.2 参考文件", level=2)
    refs = [
        "《产品规则书示例》（版式与章节结构参考）",
        "《COC大屏需求确认.docx》（COC 功能范围框定）",
        "《安全管理（人员、车辆）需求确认.docx》（安全管理功能范围框定）",
        "《COC调研_需求调研报告_20260609.docx》（COC 调研过程辅助资料）",
        "《安全管理功能清单.xlsx》（安全管理调研功能清单）",
        "深圳机场扩建工程智慧工程建设管控一体化平台原型代码（Airport_Expansion_Project_Informatization）",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    add_heading(doc, "1.3 术语与缩写解释", level=2)
    add_table(
        doc,
        ["名词", "说明"],
        [
            ("COC", "Construction Operation Center，工程调度指挥中心"),
            ("RBAC", "基于角色的访问控制（Role-Based Access Control）"),
            ("NVR", "网络硬盘录像机（Network Video Recorder），视频监控存储设备"),
            ("GIS", "地理信息系统，用于车辆轨迹与电子围栏展示"),
            ("实名制", "劳务人员实名登记与考勤管理体系"),
            ("OA", "办公自动化系统，用于组织与用户同步"),
            ("WBS", "工作分解结构，项目进度计划层级"),
        ],
        [Cm(3), Cm(13)],
    )

    # 2 项目综述
    add_heading(doc, "2、项目综述", level=1)
    add_heading(doc, "2.1 项目目标", level=2)
    add_para(
        doc,
        "深圳机场扩建工程体量大、参建单位多、作业面分散，各项目现场普遍建设智慧工地劳务系统、车辆管理系统及视频监控系统，"
        "但数据分散、标准不一，指挥部难以及时掌握人员准入、车辆监管及现场动态。建设智慧工程建设管控一体化平台，"
        "实现基础数据统一、安全管理数据汇聚与规则预警、COC 远程可视调度，支撑指挥部综合研判与风险前置管控。",
    )
    add_heading(doc, "2.2 项目范围", level=2)
    add_para(doc, "本期需求规格说明书覆盖以下四大板块：")
    scopes = [
        "基础设置：系统设置（组织、用户、部门、角色、权限、菜单）、系统对接（第三方用户/角色/组织/菜单）、日志管理（登录日志、操作日志）。",
        "基础输出（基础数据管理）：项目基础信息、分包单位管理、工程划分库，为各业务模块提供统一主数据。",
        "安全管理：人员实名制管理（劳务看板、人员档案、考勤、黑名单、预警、配置）、车辆管理（看板、进出场、轨迹监管、台账、预警）、视频监控（监控列表、NVR 设备管理）。",
        "COC 调度指挥中心：COC 调度大屏（指挥部/项目视图、领导讲话、项目调度、会议签到与录屏）及 COC 后台管理（任务单、提示函、处罚单、黑红榜、巡检仪、安全帽、监理会议、危险作业）。",
    ]
    for s in scopes:
        add_bullet(doc, s)
    add_heading(doc, "2.3 客户组织架构", level=2)
    add_para(
        doc,
        "系统面向深圳机场扩建工程指挥部及下属各参建项目。组织架构包括：工程指挥部（决策与调度层）、"
        "各在建/前期/历史项目（执行层）、监理单位、施工单位及分包单位。平台通过统一组织树实现指挥部与项目两级数据视角切换。",
    )
    add_heading(doc, "2.4 用户分析", level=2)
    add_table(
        doc,
        ["角色名称", "主要职责及场景", "主要使用功能"],
        [
            ("指挥长/指挥部领导", "远程调度、综合研判、应急指挥", "COC 大屏、领导讲话、项目进度、红黑榜"),
            ("COC 调度室", "日常监控、会议组织、设备管理", "COC 大屏、巡检仪/安全帽管理、会议录屏"),
            ("安质部主管", "安全质量监管、单据下发、榜单维护", "COC 后台任务单/提示函/处罚单、黑红榜、预警清单"),
            ("项目经理", "项目日常管理、隐患整改、单据签收", "项目调度、预警处置、每日施工作业填报"),
            ("安全/劳务管理员", "人员车辆日常管理", "人员实名制、车辆管理、考勤统计"),
            ("系统管理员", "平台配置与权限管理", "系统设置、系统对接、菜单权限"),
        ],
        [Cm(3.5), Cm(6), Cm(6.5)],
    )

    # 3 产品业务分析
    add_heading(doc, "3、产品业务分析", level=1)
    add_heading(doc, "3.1 用户原始需求", level=2)
    add_table(
        doc,
        ["优先级", "板块", "原始需求描述"],
        [
            ("P0", "COC调度中心", "提升视频监控操作便捷性，支持项目状态筛选、重点摄像头标记、平台内音视频调度"),
            ("P0", "COC调度中心", "危险作业平台化填报，领导总览今日危险作业数量及明细"),
            ("P0", "COC调度中心", "督办任务复核与附件必填，罚单/提示单模板化生成"),
            ("P0", "安全管理", "人员实名制数据汇聚、准入核验与预警规则、分级处置"),
            ("P0", "安全管理", "车辆进出场记录与轨迹监管、电子围栏越界报警"),
            ("P1", "基础输出", "统一维护项目基础信息、分包单位、工程划分库"),
            ("P1", "基础设置", "RBAC 权限体系、OA/一期平台用户组织同步"),
            ("P2", "COC调度中心", "AI 识别监理例会纪要安全质量问题（试点验证）"),
        ],
        [Cm(1.5), Cm(3), Cm(11.5)],
    )
    add_heading(doc, "3.2 业务流程分析", level=2)
    add_para(doc, "各板块核心业务流程如下：")
    flows = [
        "基础设置：第三方系统（OA/一期）→ 用户/组织同步 → 角色映射 → 菜单权限分配 → 用户登录鉴权。",
        "基础输出：指挥部维护项目主数据 → 各业务模块引用项目树 → 项目级数据隔离与指挥部汇总视图。",
        "人员实名制：现场闸机/劳务系统推送数据 → 准入规则校验 → 触发预警 → 项目自处置或分级上报 → 闭环。",
        "车辆管理：工地车辆系统推送进出场/轨迹 → GIS 围栏比对 → 越界报警 → 预警清单处置。",
        "视频监控：NVR 注册绑定项目 → 通道同步至监控列表 → COC 大屏/项目调度调用播放。",
        "COC 调度：选择项目/设备 → 远程对讲+视频监控 → 截图/语音标记问题 → 生成任务单/提示函/处罚单 → 项目整改闭环 → 纳入红黑榜。",
    ]
    for f in flows:
        add_bullet(doc, f)
    add_heading(doc, "3.3 产品架构", level=2)
    add_para(
        doc,
        "平台采用 Web 端一体化架构：工作台为统一入口，左侧菜单按业务板块组织。"
        "基础设置与基础输出（基础数据）为底层支撑；安全管理、视频监控为业务数据层；COC 调度大屏为指挥展示层，"
        "COC 后台为调度业务管理后台。各模块通过统一项目树、组织树及标准数据接口实现数据贯通。",
    )
    add_heading(doc, "3.4 功能清单", level=2)
    master_rows = build_master_function_list()
    add_table(
        doc,
        ["应用板块", "一级功能", "二级功能", "功能描述"],
        master_rows,
        [Cm(2.5), Cm(3), Cm(3), Cm(7.5)],
    )

    # 4 详细设计方案
    add_heading(doc, "4、详细设计方案", level=1)

    # 4.1 基础设置
    add_module_section(
        doc,
        "4.1 基础设置",
        "基础设置模块为平台提供统一的组织、用户、权限及第三方系统对接能力，采用 RBAC 模型实现功能权限与数据范围控制，"
        "并记录登录与操作日志以满足等保审计要求。",
        [
            "系统管理员配置组织结构树，同步或维护 OA/外部单位用户。",
            "创建角色并分配菜单、按钮及数据权限。",
            "用户登录时校验角色权限，按数据范围过滤业务数据。",
            "关键操作写入操作日志，登录行为写入登录日志。",
        ],
        [
            "系统设置 → 组织结构 / 用户管理 / 部门管理 / 角色管理 / 权限管理 / 菜单管理",
            "系统对接 → 第三方用户管理 / 第三方角色管理 / 第三方组织管理 / 第三方菜单管理",
            "日志管理 → 登录日志 / 操作日志",
        ],
        None,
        (
            ["字段", "控件", "是否必填", "说明"],
            [
                ("用户名", "文本框", "是", "登录账号，唯一"),
                ("姓名", "文本框", "是", "用户真实姓名"),
                ("所属组织", "树选择", "是", "关联组织结构节点"),
                ("角色", "多选", "是", "绑定一个或多个角色"),
                ("状态", "开关", "是", "启用/停用"),
                ("角色名称", "文本框", "是", "角色标识"),
                ("菜单权限", "树勾选", "是", "可访问的菜单与按钮"),
                ("数据范围", "下拉", "是", "全部/本部门/本项目"),
            ],
            [Cm(3), Cm(2.5), Cm(2), Cm(8.5)],
        ),
        (
            ["功能菜单", "操作名称", "前置条件", "预期结果"],
            [
                ("用户管理", "新增用户", "具有用户管理权限", "创建用户并分配角色"),
                ("用户管理", "重置密码", "选中用户", "重置为默认密码并记录日志"),
                ("角色管理", "配置权限", "选中角色", "保存菜单与数据权限"),
                ("第三方用户", "增量同步", "配置对接源", "同步新增/变更用户并映射组织"),
                ("登录日志", "检索导出", "具有日志查看权限", "按时间/用户筛选并导出"),
            ],
            [Cm(3), Cm(3), Cm(4), Cm(6)],
        ),
        [
            "基础数据管理：用户所属项目范围影响数据可见性。",
            "各业务模块：菜单与按钮权限由本模块统一配置。",
            "系统对接：OA/一期平台用户组织数据作为用户管理的数据源。",
        ],
        (
            ["页面功能", "授权对象", "操作权限"],
            [
                ("组织结构", "系统管理员", "查看、新增、编辑、删除"),
                ("用户管理", "系统管理员", "查看、新增、编辑、停用、重置密码"),
                ("角色/权限管理", "系统管理员", "查看、新增、编辑、授权"),
                ("日志管理", "系统管理员、审计员", "查看、检索、导出"),
            ],
            [Cm(3.5), Cm(5), Cm(5.5)],
        ),
    )

    # 4.2 基础输出
    add_module_section(
        doc,
        "4.2 基础输出（基础数据管理）",
        "基础输出模块（即基础数据管理）维护平台核心主数据，包括项目基础信息、分包单位登记及工程划分库，为安全管理、COC 调度、质量管理等业务提供统一的项目树与工程结构引用。",
        [
            "指挥部在「项目基础信息」维护全部在建/前期项目台账及安全画像。",
            "各项目在「分包单位管理」登记参建分包单位及资质证书。",
            "质量/安全模块引用「工程划分库」树结构进行部位定位与验评关联。",
        ],
        [
            "项目基础信息 → 项目列表 / 新增编辑 / 安全画像",
            "分包单位管理 → 分包台账 / 详情查看",
            "工程划分库 → 单位工程 / 子单位 / 分部 / 子分部 / 分项（树结构）",
        ],
        None,
        (
            ["字段", "控件", "是否必填", "说明"],
            [
                ("项目名称", "文本框", "是", "项目全称，自动生成简称"),
                ("施工单位", "文本框", "是", "总承包单位"),
                ("监理单位", "文本框", "是", "监理单位名称"),
                ("工程概况", "多行文本", "否", "项目建设概况描述"),
                ("施工许可状态", "下拉", "否", "已办理/未办理/办理中"),
                ("分包单位名称", "文本框", "是", "分包单位全称"),
                ("安全生产许可证", "附件", "是", "许可证扫描件"),
                ("工程节点名称", "文本框", "是", "划分库树节点名称"),
                ("节点类型", "下拉", "是", "单位工程/分部工程等"),
            ],
            [Cm(3), Cm(2.5), Cm(2), Cm(8.5)],
        ),
        (
            ["功能菜单", "操作名称", "前置条件", "预期结果"],
            [
                ("项目基础信息", "新增项目", "指挥部权限", "创建项目并可在各模块项目树中显示"),
                ("项目基础信息", "编辑安全画像", "选中项目", "保存参建各方安全联系人信息"),
                ("分包单位管理", "登记分包", "选中项目", "创建分包单位记录"),
                ("工程划分库", "新增节点", "选中父节点", "在树结构下新增子节点"),
            ],
            [Cm(3.5), Cm(3), Cm(3.5), Cm(6)],
        ),
        [
            "COC 调度中心：项目列表、进度、劳务分析引用项目基础信息。",
            "人员/车辆管理：左侧项目树与数据范围来源于项目主数据。",
            "质量管理（后续）：工程划分库作为质量目录树引用源。",
        ],
        (
            ["页面功能", "授权对象", "操作权限"],
            [
                ("项目基础信息", "指挥部管理员", "查看、新增、编辑"),
                ("项目基础信息", "项目用户", "查看本项目"),
                ("分包单位管理", "项目安全员", "查看、新增、编辑"),
                ("工程划分库", "质量管理员", "查看、新增、编辑、停用"),
            ],
            [Cm(3.5), Cm(5), Cm(5.5)],
        ),
    )

    # 4.3 安全管理
    add_heading(doc, "4.3 安全管理（人员实名制、车辆管理、视频监控）", level=2)
    add_heading(doc, "4.3.1 功能概述", level=3)
    add_para(
        doc,
        "安全管理板块涵盖人员实名制管理、车辆管理及视频监控三大子模块。"
        "人员实名制实现劳务数据汇聚、准入核验与预警闭环；车辆管理实现进出场记录、GIS 轨迹监管与围栏报警；"
        "视频监控提供 NVR 设备与通道管理，为 COC 大屏及业务模块提供视频能力支撑。",
    )
    add_heading(doc, "4.3.2 人员实名制管理", level=3)
    add_table(
        doc,
        ["用户端", "层级", "功能模块", "子功能", "功能描述", "需求源"],
        build_labor_function_rows(),
        [Cm(2), Cm(2), Cm(2.2), Cm(2.2), Cm(6.5), Cm(1.8)],
    )
    labor_shots = [
        ("劳务看板", "指挥部/项目视角展示人员在场、工种与年龄结构、出勤趋势及最新预警动态。", "01-labor-dashboard.png", SAFETY_SHOTS),
        ("人员实名制", "左侧项目树 + 人员列表，支持检索、入场状态筛选；未对接项目可新增人员。", "02-labor-realname.png", SAFETY_SHOTS),
        ("预警清单（人员）", "汇总实名制各类预警，项目树展示待处置数量，支持详情查看与处置。", "03-labor-warning-list.png", SAFETY_SHOTS),
        ("实名制配置", "指挥部配置预警规则与分级管控；项目节点配置现场实名制对接开关。", "04-labor-warning-config.png", SAFETY_SHOTS),
    ]
    add_heading(doc, "原型页面（人员实名制）", level=4)
    for title, desc, fname, shot_dir in labor_shots:
        add_screenshot(doc, title, desc, fname, shot_dir)

    add_heading(doc, "4.3.3 车辆管理", level=3)
    add_table(
        doc,
        ["用户端", "层级", "功能模块", "子功能", "功能描述", "需求源"],
        build_vehicle_function_rows(),
        [Cm(2), Cm(2), Cm(2.2), Cm(2.2), Cm(6.5), Cm(1.8)],
    )
    vehicle_shots = [
        ("车辆管理看板", "展示各项目车辆进出场量、在场/在途数量及异常预警统计。", "05-vehicle-dashboard.png", SAFETY_SHOTS),
        ("进出场记录", "记录车辆进场、出场信息，不做准入限制，支持方向与关键字筛选。", "06-vehicle-access.png", SAFETY_SHOTS),
        ("车辆轨迹监管", "GIS 地图展示车辆轨迹与电子围栏，支持创建禁行/允许区域。", "07-vehicle-track.png", SAFETY_SHOTS),
        ("车牌管理", "维护车辆档案，支持授权道闸、批量授权及导入导出。", "08-vehicle-registry-warning.png", SAFETY_SHOTS),
        ("预警清单（车辆）", "汇总轨迹监管及车辆监管类预警，支持来源与状态筛选。", "09-vehicle-warning-list.png", SAFETY_SHOTS),
    ]
    add_heading(doc, "原型页面（车辆管理）", level=4)
    for title, desc, fname, shot_dir in vehicle_shots:
        add_screenshot(doc, title, desc, fname, shot_dir)

    add_heading(doc, "4.3.4 视频监控", level=3)
    add_para(
        doc,
        "视频监控子模块包含「监控列表」与「设备管理」两个页面。"
        "设备管理维护 NVR 录像机注册信息（名称、绑定项目、IP、通道数、在线状态）；"
        "监控列表按项目树展示摄像头通道，支持编辑名称、位置、类型、在线状态及拖拽排序，"
        "数据供 COC 大屏宫格播放及安全管理模块调用。",
    )
    add_table(
        doc,
        ["字段", "控件", "是否必填", "说明"],
        [
            ("NVR 名称", "文本框", "是", "录像机标识名称"),
            ("绑定项目", "下拉", "是", "关联在建项目"),
            ("IP 地址", "文本框", "是", "NVR 设备 IP"),
            ("通道数", "数字", "否", "视频通道数量"),
            ("摄像头名称", "文本框", "是", "通道显示名称"),
            ("安装位置", "文本框", "否", "摄像头物理位置"),
            ("摄像头类型", "下拉", "否", "球机/枪机/半球等"),
            ("是否重点", "开关", "否", "标记为重点视频"),
        ],
        [Cm(3), Cm(2.5), Cm(2), Cm(8.5)],
    )
    add_heading(doc, "4.3.5 关联模块", level=3)
    for item in [
        "COC 调度大屏：调用监控列表摄像头数据进行宫格播放、截图标记。",
        "人员实名制：准入核验规则在闸机刷卡时触发预警（PER-02/03）。",
        "基础数据：项目树与项目主数据来源于基础数据模块。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.3.6 权限要求", level=3)
    add_table(
        doc,
        ["页面功能", "授权对象", "操作权限"],
        [
            ("劳务看板/人员实名制", "指挥部、项目安全员", "查看；项目级可处置预警"),
            ("实名制配置", "指挥部管理员", "查看、编辑规则与对接配置"),
            ("车辆管理", "指挥部、项目安全员", "查看、编辑台账、处置预警"),
            ("视频监控", "COC调度室、系统管理员", "查看、编辑通道与设备"),
        ],
        [Cm(4), Cm(5), Cm(5)],
    )

    # 4.4 COC
    add_heading(doc, "4.4 COC 调度指挥中心", level=2)
    add_heading(doc, "4.4.1 功能概述", level=3)
    add_para(
        doc,
        "COC 调度指挥中心是工程指挥部远程调度与综合管控的核心入口，包含 COC 调度大屏（全屏展示）与 COC 后台管理。"
        "大屏提供视频监控、项目进度、劳务与隐患分析、红黑榜、领导讲话、项目调度等一屏统览能力；"
        "后台管理任务单、提示函、处罚单、黑红榜、巡检仪、智能安全帽、监理会议及危险作业等业务数据。",
    )
    add_heading(doc, "4.4.2 功能清单", level=3)
    add_table(
        doc,
        ["用户端", "层级", "功能模块", "子功能", "功能描述", "需求源"],
        build_coc_function_rows(),
        [Cm(1.8), Cm(2), Cm(2.2), Cm(2.2), Cm(6.5), Cm(1.8)],
    )
    add_heading(doc, "4.4.3 原型页面", level=3)
    coc_shots = [
        ("工程指挥部首页", "默认进入指挥部视角，左侧监控列表与巡检对讲，中部视频监控与进度总览，右侧红黑榜、劳务分析、隐患分析。", "01-coc-hq-home.png", COC_SHOTS),
        ("领导讲话", "九宫格视频墙展示参会设备，支持全员静音、设备邀请及截屏生成单据。", "04-coc-leader-speech.png", COC_SHOTS),
        ("项目首页", "选择具体项目后展示项目维度监控、劳务统计与隐患清单。", "02-coc-project-view.png", COC_SHOTS),
        ("项目调度页", "对讲主画面 + 关联监控网格，下方调度记录与任务/提示/处罚单面板。", "03-coc-project-dispatch.png", COC_SHOTS),
        ("进度详情", "控制性计划甘特图详情，展示工作项与节点完成状态。", "05-coc-progress-detail.png", COC_SHOTS),
        ("设备调度视图", "点击巡检设备进入全屏调度页面。", "06-coc-dispatch-device-view.png", COC_SHOTS),
    ]
    for title, desc, fname, shot_dir in coc_shots:
        add_screenshot(doc, title, desc, fname, shot_dir)

    add_heading(doc, "4.4.4 业务流程", level=3)
    for item in [
        "调度会议：COC 调度室发起会议 → 开启录屏 → 邀请项目巡检设备入会 → 领导讲话/问题标记 → 生成告知单/处罚单 → 项目签收整改 → 闭环归档。",
        "视频巡检：选择项目 → 浏览监控宫格 → 发现隐患截图标记 → 选择单据类型 → 自动分派责任人跟进。",
        "红黑榜：后台维护榜单 → 保存触发大屏刷新；处罚单记录可导入黑榜。",
        "危险作业：项目每日 16:30 前填报次日作业计划 → 指挥部/项目调度页展示当日清单。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.4.5 权限要求", level=3)
    add_table(
        doc,
        ["页面功能", "授权对象", "操作权限"],
        [
            ("COC 大屏", "指挥部领导、COC调度室", "查看、调度、截图、会议控制"),
            ("领导讲话", "指挥长、COC调度室", "设备邀请、静音、截屏生成单据"),
            ("COC 后台-任务单/提示函/处罚单", "安质部", "创建、下发、跟踪闭环"),
            ("COC 后台-任务单/提示函/处罚单", "项目经理", "签收、整改反馈"),
            ("黑红榜单", "安质部", "新增、编辑、删除"),
            ("巡检仪/安全帽", "COC调度室", "注册、绑定、维护"),
        ],
        [Cm(4.5), Cm(4), Cm(5.5)],
    )

    # 5 第三方接口
    add_heading(doc, "5、第三方接口说明", level=1)
    add_table(
        doc,
        ["接口名称", "对接系统", "数据方向", "主要数据内容", "说明"],
        [
            ("劳务/闸机标准接口", "各项目智慧工地劳务系统", "接入", "人员基本信息、资质证书、考勤记录", "支持多厂家对接；对接后禁止手工新增人员"),
            ("车辆系统标准接口", "工地自建车辆管理系统", "接入", "车辆台账、进出场记录、轨迹定位", "进出场不做准入限制，轨迹用于围栏报警"),
            ("OA 用户组织同步", "OA 办公系统", "接入", "组织机构、用户账号", "增量同步，冲突人工处理"),
            ("一期平台数据对接", "机场扩建一期管控平台", "接入", "项目信息、进度数据、劳务数据", "COC 进度总览、项目列表引用"),
            ("视频流媒体平台", "海康/宇视等 NVR", "接入", "实时视频流、回放、云台控制", "通过 NVR 注册统一管理"),
            ("巡检对讲设备", "执法记录仪/Web对讲", "双向", "实时音视频流", "支持远程呼叫与项目调度"),
        ],
        [Cm(3), Cm(3.5), Cm(1.5), Cm(4.5), Cm(3.5)],
    )

    doc.save(OUT_FILE)
    print(f"已生成: {OUT_FILE}")


if __name__ == "__main__":
    main()
