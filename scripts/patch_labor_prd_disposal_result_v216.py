# -*- coding: utf-8 -*-
"""V2.1.6：处置结果字段 + 详情处置信息模块（仅最终口径，不写变更过程叙事）。"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DOC = Path(
    r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理\人员实名制_PRD_评审通过.docx"
)

# 正文/表格最终口径替换（目标表述，不含「改为/取消」过程词）
EXACT_PAIRS = [
    (
        "「批量处置预警」：勾选后弹窗填报处置说明与附件（与详情处置一致），"
        "仅状态为「待处理」的处置任务关闭为「已关闭」（同步业务预警与预警中心）；"
        "「批量已读」仅对通知类「未读」生效。",
        "「批量处置预警」：勾选后弹窗填报处置结果、处置说明与附件（与详情处置一致；"
        "处置结果为必填单选「已处置/误报」，默认「已处置」），"
        "仅状态为「待处理」的处置任务关闭为「已关闭」（同步业务预警与预警中心）；"
        "「批量已读」仅对通知类「未读」生效。",
    ),
    (
        "批量处置预警：弹窗填报处置说明、附件（与详情处置一致）；"
        "仅「待处理」的处置任务批量关闭为「已关闭」（同步业务预警状态与处置时间线）；"
        "通知类不走批量处置。",
        "批量处置预警：弹窗填报处置结果、处置说明、附件（与详情处置一致；"
        "处置结果必填单选「已处置/误报」，默认「已处置」）；"
        "仅「待处理」的处置任务批量关闭为「已关闭」（同步业务预警状态与处置时间线）；"
        "通知类不走批量处置。",
    ),
]

CELL_EXACT = {
    # T56 交互
    ("打开详情", "展示预警信息、关联人员、时间线"): (
        "打开详情",
        "展示预警信息；若为手动处理且已关闭则展示处置信息（处置结果、处置说明、处置附件）；再展示关联人员与处置时间线",
    ),
    ("处置并关闭", "状态→已关闭"): (
        "处置并关闭",
        "表单最上方必填处置结果（已处置/误报，默认已处置），并填报处置说明、选填附件；提交后状态→已关闭",
    ),
    # T84 批量处置
    (
        "批量处置预警",
        "弹窗填报处置说明、附件（与详情处置一致）；仅「待处理」的处置任务改为「已关闭」（同步业务预警）",
    ): (
        "批量处置预警",
        "弹窗填报处置结果（必填单选已处置/误报，默认已处置）、处置说明、附件（与详情处置一致）；仅「待处理」的处置任务关闭为「已关闭」（同步业务预警）",
    ),
}

# 验收 Then 最终口径（短句）
ACCEPT_THEN = {
    ("AC-01", "手动关闭"): "状态已关闭；关单记录含处置结果",
    ("AC-02", "处置任务手动关闭"): "业务预警与预警中心对应条状态均为「已关闭」，详情可见处置信息；APP 无处置入口",
    ("AC-05", "批量已读与处置"): "未读通知变已读；待处理处置任务变已关闭（业务预警同步关闭）；批量处置须含处置结果",
}


def set_para_text(p, value: str) -> None:
    if p.runs:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(value)


def set_cell_text(cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    set_para_text(cell.paragraphs[0], value)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def format_table(table) -> None:
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def replace_in_paragraph(p, pairs) -> int:
    text = p.text
    if not text:
        return 0
    new = text
    for a, b in pairs:
        if a in new:
            new = new.replace(a, b)
    if new == text:
        return 0
    set_para_text(p, new)
    return 1


def add_row_clone(table, values: list[str]):
    tmpl = table.rows[-1]._tr
    new_tr = deepcopy(tmpl)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell_text(row.cells[i], val)
    return row


def set_meta_version(doc: Document) -> None:
    t0 = doc.tables[0]
    for row in t0.rows:
        key = row.cells[0].text.strip()
        if key == "版本号":
            set_cell_text(row.cells[1], "V2.1.6")
        elif key == "更新日期":
            set_cell_text(row.cells[1], "2026-08-20")


def add_revision(doc: Document) -> None:
    t1 = doc.tables[1]
    # 已存在则跳过
    for row in t1.rows[1:]:
        if row.cells[0].text.strip() == "V2.1.6":
            return
    add_row_clone(
        t1,
        [
            "V2.1.6",
            "—",
            "2026-08-20",
            "详情新增处置与预警中心批量处置表单含必填处置结果（已处置/误报，默认已处置）；"
            "详情页预警信息下增加处置信息模块（处置结果、处置说明、处置附件；仅手动处理且已关闭展示）。",
        ],
    )
    format_table(t1)


def patch_t62(doc: Document) -> None:
    t = doc.tables[62]
    # 已含 disposal_result 则跳过插入
    has_result = any("disposal_result" in r.cells[1].text for r in t.rows)
    if not has_result:
        # 在 content 行前插入 disposal_result
        content_idx = None
        for i, row in enumerate(t.rows):
            if row.cells[1].text.strip() == "content":
                content_idx = i
                break
        if content_idx is None:
            raise RuntimeError("T62 未找到 content 行")
        tmpl = t.rows[content_idx]._tr
        new_tr = deepcopy(tmpl)
        tmpl.addprevious(new_tr)
        # 重新定位新行（content_idx 位置即为新行）
        row = t.rows[content_idx]
        vals = [
            "新增/批量处置表单",
            "disposal_result",
            "处置结果",
            "enum",
            "是",
            "已处置 / 误报；表单最上方单选；默认已处置；关单时写入处置记录",
        ]
        for i, v in enumerate(vals):
            set_cell_text(row.cells[i], v)

    for row in t.rows:
        field = row.cells[1].text.strip()
        if field == "content":
            set_cell_text(row.cells[0], "")
            set_cell_text(row.cells[2], "处置说明")
            set_cell_text(row.cells[4], "是")
            set_cell_text(
                row.cells[5],
                "详情「新增处置」与预警中心「批量处置」必填；提交后写入处置记录",
            )
        elif field == "attachments":
            set_cell_text(row.cells[2], "处置附件")
            set_cell_text(row.cells[5], "选填；提交后写入处置记录")
        elif field == "disposal_result":
            set_cell_text(
                row.cells[5],
                "已处置 / 误报；表单最上方单选；默认已处置；关单时写入处置记录",
            )
    format_table(t)


def ensure_handle_info_section(doc: Document) -> None:
    """在「预警处置（时间线 / 新增表单）」后写入处置信息说明与字段表。"""
    # 定位段落
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("预警处置（时间线"):
            anchor_idx = i
            break
    if anchor_idx is None:
        raise RuntimeError("未找到预警处置字段小节标题")

    # 已存在则只校正标题/表，不重复插入
    already = any(p.text.strip() == "处置信息（只读）" for p in doc.paragraphs)
    if already:
        # 校正已有处置信息表（找含「仅手动处理且已关闭」的表）
        for t in doc.tables:
            blob = "\n".join(c.text for r in t.rows for c in r.cells)
            if "仅手动处理且已关闭" in blob and "disposal_result" in blob:
                _fill_handle_info_table(t)
                format_table(t)
                return
        return

    # 更新小节标题
    set_para_text(
        doc.paragraphs[anchor_idx],
        "预警处置（时间线 / 新增表单 / 批量处置表单）",
    )

    # 在 6.6.10 标题前插入：处置信息标题 + 说明 + 表
    # 找 6.6.10
    h_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("6.6.10"):
            h_idx = i
            break
    if h_idx is None:
        raise RuntimeError("未找到 6.6.10")

    target = doc.paragraphs[h_idx]

    # 插入顺序：先插最靠近标题的段落，再向前插（用 insert_paragraph_before）
    # 最终顺序应为：...处置信息标题、说明、空、表、6.6.10
    # python-docx 无直接 add_table before；用 tbl 插入到段落 XML 之前

    title_p = target.insert_paragraph_before("处置信息（只读）")
    title_p.style = doc.paragraphs[anchor_idx].style
    # 保持 Normal
    title_p.style = "Normal"
    for r in title_p.runs:
        r.bold = True

    desc = (
        "详情页位于「预警信息」下方；仅当处置方式为「手动处理」且状态为「已关闭」时展示。"
        "字段取关单处置记录：处置结果、处置说明、处置附件。"
    )
    desc_p = target.insert_paragraph_before(desc)
    desc_p.style = "Normal"

    # 空行
    target.insert_paragraph_before("")

    # 新建表：克隆 T62 表结构
    src = doc.tables[62]
    new_tbl = deepcopy(src._tbl)
    # 清空至仅表头 1 行
    for tr in list(new_tbl.findall(qn("w:tr")))[1:]:
        new_tbl.remove(tr)
    target._p.addprevious(new_tbl)

    # 找到刚插入的表（含表头且目前只有 1 行）
    handle_table = None
    for t in doc.tables:
        if t._tbl is new_tbl:
            handle_table = t
            break
    if handle_table is None:
        raise RuntimeError("处置信息表插入失败")

    _fill_handle_info_table(handle_table)
    format_table(handle_table)


def _fill_handle_info_table(table) -> None:
    # 保证有表头 + 3 行
    while len(table.rows) < 4:
        add_row_clone(table, ["", "", "", "", "", ""])
    # 多余行删除
    while len(table.rows) > 4:
        table._tbl.remove(table.rows[-1]._tr)

    rows = [
        ["实体/对象", "字段名", "中文名", "类型", "必填", "说明"],
        [
            "处置信息",
            "disposal_result",
            "处置结果",
            "enum",
            "—",
            "已处置 / 误报；只读；来源于关单记录",
        ],
        [
            "",
            "content",
            "处置说明",
            "string",
            "—",
            "只读；来源于关单记录",
        ],
        [
            "",
            "attachments",
            "处置附件",
            "array",
            "—",
            "只读；来源于关单记录；无则展示「—」",
        ],
    ]
    for i, vals in enumerate(rows):
        for j, v in enumerate(vals):
            set_cell_text(table.rows[i].cells[j], v)


def patch_cell_exact(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 3:
                continue
            op = row.cells[0].text.strip()
            resp = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            key = (op, resp)
            if key in CELL_EXACT:
                new_op, new_resp = CELL_EXACT[key]
                set_cell_text(row.cells[0], new_op)
                set_cell_text(row.cells[2], new_resp)
                n += 1
            # 宽松：仅按操作名匹配批量处置
            if op == "批量处置预警" and "处置结果" not in resp and "处置说明" in resp:
                set_cell_text(
                    row.cells[2],
                    "弹窗填报处置结果（必填单选已处置/误报，默认已处置）、处置说明、附件（与详情处置一致）；"
                    "仅「待处理」的处置任务关闭为「已关闭」（同步业务预警）",
                )
                n += 1
            if op == "打开详情" and "处置信息" not in resp and "预警信息" in resp:
                set_cell_text(
                    row.cells[2],
                    "展示预警信息；若为手动处理且已关闭则展示处置信息（处置结果、处置说明、处置附件）；"
                    "再展示关联人员与处置时间线",
                )
                n += 1
            if op == "处置并关闭" and "处置结果" not in resp:
                set_cell_text(
                    row.cells[2],
                    "表单最上方必填处置结果（已处置/误报，默认已处置），并填报处置说明、选填附件；提交后状态→已关闭",
                )
                n += 1
    return n


def patch_acceptance(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr[:2] != ["编号", "验收项"]:
            continue
        for row in table.rows[1:]:
            if len(row.cells) < 5:
                continue
            key = (row.cells[0].text.strip(), row.cells[1].text.strip())
            if key in ACCEPT_THEN:
                set_cell_text(row.cells[4], ACCEPT_THEN[key])
                n += 1
        # AC-05 Given/When 最终口径
        for row in table.rows[1:]:
            if row.cells[0].text.strip() == "AC-05" and "批量" in row.cells[1].text:
                when = row.cells[3].text.strip()
                if "处置结果" not in when:
                    set_cell_text(
                        row.cells[3],
                        "批量已读、批量处置预警（填处置结果+说明+附件）",
                    )
                    n += 1
    return n


def patch_t85_detail(doc: Document) -> None:
    t = doc.tables[85]
    for row in t.rows:
        if row.cells[0].text.strip() == "详情页":
            set_cell_text(
                row.cells[1],
                "Web 端且当前用户为默认接收人时，处置能力与 F-LABOR-06 对齐（含处置结果与处置信息模块）；"
                "APP/非默认接收人只读；手动处理且已关闭时详情展示处置信息",
            )


def main() -> None:
    doc = Document(str(DOC))
    set_meta_version(doc)
    add_revision(doc)

    para_hits = 0
    for p in doc.paragraphs:
        para_hits += replace_in_paragraph(p, EXACT_PAIRS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    para_hits += replace_in_paragraph(p, EXACT_PAIRS)

    patch_t62(doc)
    ensure_handle_info_section(doc)
    cell_hits = patch_cell_exact(doc)
    ac_hits = patch_acceptance(doc)
    patch_t85_detail(doc)

    doc.save(str(DOC))
    print(
        f"saved {DOC.name}; para/cell text hits≈{para_hits}, "
        f"interaction cells={cell_hits}, acceptance={ac_hits}"
    )


if __name__ == "__main__":
    main()
