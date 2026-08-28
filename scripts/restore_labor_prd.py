# -*- coding: utf-8 -*-
"""修复人员实名制 PRD：从 Word 可打开副本恢复，再安全重放内容修订。"""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import etree

BASE = Path(r"C:\Users\13065\Desktop\机场平台2期\调研记录\4、安全管理\人员、车辆管理")
CUR = BASE / "人员实名制_PRD_待评审.docx"
OUT_LOG = Path(
    r"C:\Users\13065\Desktop\机场平台2期\Airport_Expansion_Project_Informatization\scripts\_labor_restore_log.txt"
)

NOTE_MENU = "指挥部菜单挂载由「施工现场管理」调整为「智慧工地监管」；页面路由不变。"
NOTE_API = "删除 6.2.9「API 业务契约」空壳标题及接口用途表；6.2 节后续小节编号顺延（业务规则→6.2.9）。"
VER = "V2.1.3"
TODAY = "2026-08-13"

MENU_REPLACEMENTS = [
    ("挂施工现场管理 / 智慧工地监管", "挂智慧工地监管"),
    ("挂施工现场管理/智慧工地监管", "挂智慧工地监管"),
    ("施工现场管理 → 人员实名制管理", "智慧工地监管 → 人员实名制管理"),
    ("施工现场管理 · 人员实名制管理", "智慧工地监管 · 人员实名制管理"),
    ("挂施工现场管理", "挂智慧工地监管"),
    ("位于施工现场管理，不在安全看板", "位于智慧工地监管，不在安全看板"),
    ("位于施工现场管理，", "位于智慧工地监管，"),
    ("指挥部挂施工现场管理", "指挥部挂智慧工地监管"),
    ("指挥部「施工现场管理 → 人员实名制管理」", "指挥部「智慧工地监管 → 人员实名制管理」"),
    ("仅指挥部「施工现场管理」可见", "仅指挥部「智慧工地监管」可见"),
    ("「施工现场管理」", "「智慧工地监管」"),
]


def log(lines: list[str], msg: str) -> None:
    print(msg)
    lines.append(msg)


def find_good_base() -> Path:
    preferred = [
        "人员实名制_PRD_待评审_去API契约节前_副本_20260813.docx",
        "人员实名制_PRD_待评审_去API契约空壳前_副本_20260813.docx",
        "人员实名制_PRD_待评审_字段命名统一前_副本_20260813.docx",
    ]
    for name in preferred:
        p = BASE / name
        if p.exists():
            return p
    # fallback: any 去API* backup
    for p in sorted(BASE.glob("人员实名制_PRD_待评审_去API*副本*.docx")):
        return p
    raise FileNotFoundError("no good labor PRD backup found")


def set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.text = text
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def replace_in_paragraph(para, pairs: list[tuple[str, str]]) -> int:
    full = "".join(r.text for r in para.runs) if para.runs else (para.text or "")
    new = full
    n = 0
    for old, repl in pairs:
        if old in new:
            c = new.count(old)
            new = new.replace(old, repl)
            n += c
    if new != full:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = new
    return n


def remove_api_section(doc: Document, lines: list[str]) -> None:
    to_delete: list[int] = []
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if "API 业务契约" not in t and not (t.startswith("6.") and "业务契约" in t):
            continue
        to_delete.append(i)
        j = i + 1
        while j < len(doc.paragraphs):
            nt = doc.paragraphs[j].text.strip()
            ns = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if ns and ns.startswith("Heading"):
                break
            if nt == "" or "本期不写" in nt or "规格验收以本章" in nt or "不作为验收依据" in nt:
                to_delete.append(j)
                j += 1
                continue
            break

    # Also remove a following table if it is the API 用途表 (detect by header cells)
    # Safer: only delete paragraphs as original script; then scan tables for 业务契约
    lines.append(f"api delete paragraph indices: {sorted(set(to_delete))}")
    for i in sorted(set(to_delete), reverse=True):
        el = doc.paragraphs[i]._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # 仅删除明确含「业务契约」字样的用途表（避免误删字段说明表）
    tbls_to_remove = []
    for table in doc.tables:
        blob = "\n".join(c.text for row in table.rows for c in row.cells)
        if "业务契约" in blob:
            tbls_to_remove.append(table._tbl)
    for tbl in tbls_to_remove:
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)
            lines.append("removed api-contract table")

    renames = [
        ("6.2.12 ", "6.2.11 "),
        ("6.2.11 ", "6.2.10 "),
        ("6.2.10 ", "6.2.9 "),
    ]
    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs) if para.runs else para.text
        for old, new in renames:
            if full.startswith(old):
                new_full = new + full[len(old) :]
                if para.runs:
                    para.runs[0].text = new_full
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.text = new_full
                lines.append("renumber: " + new_full[:80])
                break


def write_revision_rows_via_word(path: Path, lines: list[str]) -> None:
    """用 Word COM 在修订表顶部插入行（避免 python-docx 破坏 OOXML）。"""
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(path), ReadOnly=False)
        # 假定第 2 个表为修订记录（与 python-docx tables[1] 一致）
        if doc.Tables.Count < 2:
            lines.append("WARN: no revision table in Word")
            doc.Close(False)
            return
        tbl = doc.Tables(2)
        # 在表头下插入两行：V2.1.3、V2.1.2
        entries = [
            (VER, "—", TODAY, NOTE_MENU),
            ("V2.1.2", "—", TODAY, NOTE_API),
        ]
        # 若已有 V2.1.3 则跳过
        first_ver = tbl.Cell(2, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
        if first_ver == VER:
            lines.append("revision V2.1.3 already present")
        else:
            for ver, person, date, note in reversed(entries):
                # 已有 V2.1.2 则只插 V2.1.3
                if ver == "V2.1.2":
                    # scan first 5 data rows
                    exists = False
                    for r in range(2, min(tbl.Rows.Count, 8) + 1):
                        v = tbl.Cell(r, 1).Range.Text.strip().replace("\r", "").replace("\x07", "")
                        if v == "V2.1.2":
                            exists = True
                            break
                    if exists:
                        continue
                tbl.Rows(2).Select()
                word.Selection.InsertRowsAbove(1)
                row = tbl.Rows(2)
                vals = [ver, person, date, note]
                for c, val in enumerate(vals, start=1):
                    if c <= row.Cells.Count:
                        row.Cells(c).Range.Text = val
            lines.append("inserted revision rows via Word")
        doc.Save()
        doc.Close(False)
    except Exception as e:
        lines.append(f"Word revision insert FAIL: {e}")
        try:
            word.Quit()
        except Exception:
            pass
        raise
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def bump_meta(doc: Document) -> None:
    meta = doc.tables[0]
    for row in meta.rows:
        key = row.cells[0].text.strip()
        if key == "版本号" and len(row.cells) > 1:
            set_cell_text(row.cells[1], VER)
        if key == "更新日期" and len(row.cells) > 1:
            set_cell_text(row.cells[1], TODAY)


def word_open_ok(path: Path, lines: list[str]) -> bool:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(path), ReadOnly=True)
        lines.append(f"WORD OK {path.name} paras={doc.Paragraphs.Count}")
        txt = doc.Content.Text
        lines.append(
            f"  API_section_left={('6.2.9 API' in txt or 'API 业务契约' in txt and '删除' not in txt[:2000])}"
        )
        # more precise: heading-like leftover
        api_hit = ("API 业务契约" in txt) and ("删除 6.2.9「API 业务契约」" not in txt or txt.count("API 业务契约") > 1)
        lines.append(f"  API_string_count_gt1_or_present={api_hit} count={txt.count('API 业务契约')}")
        lines.append(
            f"  mount_old={('施工现场管理 → 人员' in txt or '施工现场管理 · 人员' in txt)}"
        )
        lines.append(
            f"  mount_new={('智慧工地监管 → 人员' in txt or '智慧工地监管 · 人员' in txt)}"
        )
        doc.Close(False)
        return True
    except Exception as e:
        lines.append(f"WORD FAIL {path.name}: {e}")
        return False
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def sanitize_via_word(src: Path, dst: Path, lines: list[str]) -> bool:
    """用 Word 另存一遍，剔除潜在损坏。"""
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(src), ReadOnly=False)
        # wdFormatXMLDocument = 12
        if dst.exists():
            dst.unlink()
        doc.SaveAs2(str(dst), FileFormat=12)
        doc.Close(False)
        lines.append(f"Word SaveAs sanitize -> {dst.name}")
        return True
    except Exception as e:
        lines.append(f"Word sanitize FAIL: {e}")
        return False
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def main() -> None:
    lines: list[str] = []
    good = find_good_base()
    log(lines, f"base: {good.name}")

    # quarantine broken
    if CUR.exists():
        broken = BASE / f"人员实名制_PRD_待评审_损坏不可打开_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        shutil.copy2(CUR, broken)
        log(lines, f"quarantine -> {broken.name}")

    # copy good -> temp work, sanitize with Word first
    work = BASE / "_labor_prd_restore_work.docx"
    shutil.copy2(good, work)
    sanitized = BASE / "_labor_prd_sanitized.docx"
    if not sanitize_via_word(work, sanitized, lines):
        sanitized = work
    else:
        # verify sanitized opens
        if not word_open_ok(sanitized, lines):
            sanitized = work

    shutil.copy2(sanitized, CUR)
    log(lines, "restored current from sanitized base")

    doc = Document(str(CUR))
    remove_api_section(doc, lines)

    n = 0
    for para in doc.paragraphs:
        n += replace_in_paragraph(para, MENU_REPLACEMENTS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    n += replace_in_paragraph(para, MENU_REPLACEMENTS)
    log(lines, f"menu replacements: {n}")

    bump_meta(doc)
    doc.save(str(CUR))
    log(lines, "saved after python edits (meta/menu/api)")

    # Word sanitize after structural paragraph deletes
    final_tmp = BASE / "_labor_prd_final_tmp.docx"
    if sanitize_via_word(CUR, final_tmp, lines):
        shutil.copy2(final_tmp, CUR)
        log(lines, "replaced current with Word-sanitized final")

    # 修订行用 Word 插入，避免 OOXML 损坏
    write_revision_rows_via_word(CUR, lines)

    ok = word_open_ok(CUR, lines)

    # cleanup temps
    for p in [work, sanitized, final_tmp]:
        try:
            if p.exists() and p.name.startswith("_labor_prd"):
                p.unlink()
        except Exception:
            pass

    # leftover checks via python-docx
    d2 = Document(str(CUR))
    api_left = [p.text for p in d2.paragraphs if "业务契约" in p.text]
    leftover = []
    for p in d2.paragraphs:
        if "施工现场管理" in p.text and "调整为" not in p.text:
            leftover.append(p.text[:100])
    for t in d2.tables:
        for row in t.rows:
            for cell in row.cells:
                if "施工现场管理" in cell.text and "调整为" not in cell.text:
                    leftover.append(cell.text[:100])
    meta = {r.cells[0].text.strip(): r.cells[1].text.strip() for r in d2.tables[0].rows if len(r.cells) > 1}
    lines.append(f"meta={meta.get('版本号')} {meta.get('更新日期')}")
    lines.append(f"api_left_paras={api_left}")
    lines.append(f"mount_leftover={len(leftover)}")
    lines.append(f"SUCCESS={ok}")

    OUT_LOG.write_text("\n".join(lines), encoding="utf-8")
    print("log ->", OUT_LOG)
    if not ok:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
